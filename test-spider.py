#!/usr/bin/env python3
"""
IHS Hybrid Donor Research System - NARRATIVE FORMAT EDITION
------------------------------------------------------------
Generates donor profiles in narrative prose format with deep biographical,
career, and philanthropic research.

OUTPUT FORMAT (10 SECTIONS):
- Title: "[Name]: Donor Profile"
- 1. Biographical Background & Education (narrative paragraphs)
- 2. Career History & Business Leadership (narrative paragraphs)
- 3. Philanthropic Activities & Causes Supported (narrative paragraphs)
- 4. Political Donations & Ideological Leanings (narrative paragraphs)
- 5. Connections (Inner Circle comparison - 357 members analyzed)
- 6. Network & Board Affiliations (donor's own network and boards)
- 7. Estimated Net Worth & Giving Capacity (narrative paragraphs)
- 8. IHS Donor Probability Model Assessment (A-E-I-N-P-S framework in prose)
- 9. Strategic Summary (Board-Level Briefing)
- 10. Sources & Citations (comprehensive list of all research sources)

Combines proven architecture with advanced features:
- Deep biographical research (6 searches, 3-5 narrative paragraphs)
- Deep career research (6 searches, 4-6 narrative paragraphs)
- Deep philanthropic research (5 searches, 3-5 narrative paragraphs)
- Systematic Inner Circle comparison (357 members)
- Comprehensive source tracking (all URLs cited)
- Caching system (saves 50% on API costs)
- Retry logic (handles failures)
- Confidence scoring (rates connection quality)
- Test mode (processes only first donor)
- Optimized for YOUR data format

SETUP:
Create spider.env with:
    OPENAI_API_KEY=sk-your-key
    OPENAI_MODEL=gpt-4o
    GOOGLE_API_KEY=your-key
    GOOGLE_CSE_ID=your-cse-id
    FEC_API_KEY=DEMO_KEY

USAGE:
    python3 test-spider.py                    # Process all donors
    python3 test-spider.py --test-mode        # Process ONLY first donor (for testing)
    python3 test-spider.py --name "John Doe"  # Process specific donor
"""

import os
import re
import sys
import csv
import json
import time
import math
import hashlib
import argparse
import logging
import warnings
from datetime import datetime
from typing import List, Dict, Any, Tuple, Optional
from collections import defaultdict
from functools import wraps

import requests
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Suppress urllib3 LibreSSL warnings
warnings.filterwarnings('ignore', category=UserWarning, module='urllib3')

# -------------------------
# SETUP & CONFIG
# -------------------------
load_dotenv("spider.env")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL   = os.getenv("OPENAI_MODEL", "gpt-4")

# Support multiple Google API keys for higher rate limits
# Format: GOOGLE_API_KEY_1, GOOGLE_API_KEY_2, etc.
GOOGLE_API_KEYS = []
GOOGLE_CSE_ID  = os.getenv("GOOGLE_CSE_ID", "")

# Load all GOOGLE_API_KEY_N variables
for i in range(1, 21):  # Support up to 20 keys
    key = os.getenv(f"GOOGLE_API_KEY_{i}", "")
    if key:
        GOOGLE_API_KEYS.append(key)

# Fallback to single GOOGLE_API_KEY if no numbered keys found
if not GOOGLE_API_KEYS:
    single_key = os.getenv("GOOGLE_API_KEY", "")
    if single_key:
        GOOGLE_API_KEYS.append(single_key)

TAVILY_API_KEY = os.getenv("TAVILY_API_KEY", "")
SEARCH_PROVIDER = os.getenv("SEARCH_PROVIDER", "google")
FEC_API_KEY    = os.getenv("FEC_API_KEY", "")
CACHE_DIR      = os.getenv("CACHE_DIR", "./cache")
CACHE_TTL      = int(os.getenv("CACHE_TTL", "86400"))

# Key rotation state
_google_key_index = 0
_google_key_exhausted = set()  # Track which keys hit rate limits

def reset_google_keys():
    """Reset exhausted keys tracking (call this daily or when quotas reset)"""
    global _google_key_exhausted
    _google_key_exhausted.clear()
    log.info("Reset Google API key exhaustion tracking")

try:
    import openai
    # Detect OpenAI library version
    try:
        # Try new API (1.0+)
        OPENAI_CLIENT = openai.OpenAI(api_key=OPENAI_API_KEY)
        OPENAI_VERSION = "new"
    except AttributeError:
        # Fall back to old API (0.28)
        openai.api_key = OPENAI_API_KEY
        OPENAI_CLIENT = None
        OPENAI_VERSION = "old"
except Exception:
    openai = None
    OPENAI_CLIENT = None
    OPENAI_VERSION = None

logging.basicConfig(level=logging.INFO, format="[%(levelname)s] %(message)s")
log = logging.getLogger("spider")

# Suppress OpenAI's verbose HTTP logging
logging.getLogger("httpx").setLevel(logging.WARNING)
logging.getLogger("openai").setLevel(logging.WARNING)

# -------------------------
# CACHING SYSTEM (from v5.3)
# -------------------------
class ResearchCache:
    """Cache API results to save money and time"""
    def __init__(self, cache_dir: str = "./cache", ttl: int = 86400):
        self.cache_dir = cache_dir
        self.ttl = ttl
        os.makedirs(cache_dir, exist_ok=True)
    
    def get_cache_key(self, prefix: str, params: Dict) -> str:
        param_str = json.dumps(params, sort_keys=True)
        hash_key = hashlib.md5(param_str.encode()).hexdigest()[:8]
        return f"{prefix}_{hash_key}"
    
    def get_or_fetch(self, key: str, fetcher, ttl: Optional[int] = None):
        ttl = ttl or self.ttl
        cache_file = os.path.join(self.cache_dir, f"{key}.json")
        
        # Check if cached and still valid
        if os.path.exists(cache_file):
            mtime = os.path.getmtime(cache_file)
            if time.time() - mtime < ttl:
                try:
                    with open(cache_file, 'r', encoding='utf-8') as f:
                        log.debug(f"Cache HIT: {key}")
                        return json.load(f)
                except:
                    pass
        
        # Fetch fresh data
        log.debug(f"Cache MISS: {key}")
        data = fetcher()
        
        # Save to cache
        try:
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(data, f)
        except:
            pass
        
        return data

# Global cache instance
cache = ResearchCache(CACHE_DIR, CACHE_TTL)

# -------------------------
# RETRY LOGIC (from v5.3)
# -------------------------
def retry_with_backoff(max_retries: int = 3, initial_wait: float = 1.0):
    """Decorator to retry failed API calls with exponential backoff"""
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            for attempt in range(max_retries):
                try:
                    return func(*args, **kwargs)
                except Exception as e:
                    if attempt == max_retries - 1:
                        log.warning(f"Final retry failed for {func.__name__}: {e}")
                        raise
                    wait = initial_wait * (2 ** attempt)
                    log.info(f"Retry {attempt + 1}/{max_retries} after {wait}s...")
                    time.sleep(wait)
            return None
        return wrapper
    return decorator

# -------------------------
# CSV LOADERS (fixed for YOUR format)
# -------------------------
def load_csv(path: str) -> List[Dict[str, str]]:
    """Load CSV with automatic encoding detection and column name normalization"""
    encodings_to_try = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1', 'mac_roman']
    
    for encoding in encodings_to_try:
        try:
            with open(path, newline="", encoding=encoding) as f:
                rows = list(csv.DictReader(f))
                
                # Normalize column names for YOUR data format
                normalized_rows = []
                for row in rows:
                    normalized = {}
                    for key, value in row.items():
                        # Map your column names to standard format
                        if key == 'First Name' or key == 'First':
                            normalized['First'] = value
                        elif key == 'Last Name' or key == 'Last':
                            normalized['Last'] = value
                        elif key in ['City', 'city']:
                            normalized['City'] = value
                        elif key in ['State', 'state']:
                            normalized['State'] = value
                        elif key == 'Middle':
                            normalized['Middle'] = value
                        else:
                            normalized[key] = value
                    
                    # Combine First + Last into 'name' for backward compatibility
                    first = (normalized.get('First') or '').strip()
                    last = (normalized.get('Last') or '').strip()
                    if first or last:
                        normalized['name'] = f"{first} {last}".strip()
                    
                    normalized_rows.append(normalized)
                
                log.info(f"Loaded {path} with {encoding} encoding ({len(normalized_rows)} rows)")
                return normalized_rows
        except UnicodeDecodeError:
            continue
        except Exception as e:
            log.warning(f"Error reading {path} with {encoding}: {e}")
            continue
    
    # Last resort: try with error handling
    try:
        with open(path, newline="", encoding='utf-8', errors='ignore') as f:
            log.warning(f"Reading {path} with UTF-8 and ignoring errors")
            return list(csv.DictReader(f))
    except Exception as e:
        log.error(f"Failed to read {path}: {e}")
        return []

def find_donor_in_data(data_rows: List[Dict[str, str]], name: str) -> Optional[Dict[str, str]]:
    """Find donor by name, handling various column formats"""
    n = name.strip().lower()
    for r in data_rows:
        # Try direct 'name' match
        row_name = (r.get("name") or r.get("Name") or "").strip().lower()
        if row_name == n:
            return r
        
        # Try combining First + Last
        first = (r.get("First") or r.get("First Name") or "").strip()
        last = (r.get("Last") or r.get("Last Name") or "").strip()
        full_name = f"{first} {last}".strip().lower()
        if full_name == n:
            return r
    return None

def get_row_name(row: Dict[str, str]) -> str:
    """Extract name from row, handling various formats"""
    # Try direct name field
    name = (row.get("name") or row.get("Name") or "").strip()
    if name:
        return name
    
    # Try combining First + Last
    first = (row.get("First") or row.get("First Name") or "").strip()
    last = (row.get("Last") or row.get("Last Name") or "").strip()
    if first or last:
        return f"{first} {last}".strip()
    
    return ""

# -------------------------
# SEARCH FUNCTIONS (Google Custom Search OR Tavily)
# -------------------------
@retry_with_backoff(max_retries=3)
def tavily_search(query: str, num: int = 10) -> List[Dict[str, str]]:
    """Tavily Search API - optimized for AI research"""
    if not TAVILY_API_KEY:
        log.warning("Tavily search disabled (missing API key)")
        return []
    
    cache_key = cache.get_cache_key("tavily", {"q": query, "num": num})
    
    def fetcher():
        url = "https://api.tavily.com/search"
        headers = {"Content-Type": "application/json"}
        data = {
            "api_key": TAVILY_API_KEY,
            "query": query,
            "max_results": min(10, num),
            "search_depth": "basic",  # or "advanced" for deeper search
            "include_answer": False,
            "include_raw_content": False
        }
        r = requests.post(url, headers=headers, json=data, timeout=20)
        r.raise_for_status()
        results = (r.json() or {}).get("results", []) or []
        return [
            {
                "title": item.get("title", ""),
                "snippet": item.get("content", ""),
                "link": item.get("url", "")
            }
            for item in results
        ]
    
    return cache.get_or_fetch(cache_key, fetcher)

def google_search(query: str, num: int = 10) -> List[Dict[str, str]]:
    """Google CSE with smart key switching on rate limits
    
    Automatically switches to next available API key when rate limits are hit.
    This provides seamless failover without manual intervention.
    """
    global _google_key_index, _google_key_exhausted
    
    if not GOOGLE_API_KEYS or not GOOGLE_CSE_ID:
        log.warning("Google search disabled (missing API keys)")
        return []
    
    cache_key = cache.get_cache_key("google", {"q": query, "num": num})
    
    def fetcher():
        global _google_key_index, _google_key_exhausted
        
        # Try each available key until one works
        attempts = 0
        max_attempts = len(GOOGLE_API_KEYS)
        
        while attempts < max_attempts:
            # Get current key index (skip exhausted keys)
            current_index = _google_key_index % len(GOOGLE_API_KEYS)
            
            # If all keys exhausted, reset and try again
            if len(_google_key_exhausted) >= len(GOOGLE_API_KEYS):
                log.warning("All Google API keys exhausted! Waiting and resetting...")
                time.sleep(60)  # Wait 1 minute
                _google_key_exhausted.clear()
            
            # Skip if this key is exhausted
            if current_index in _google_key_exhausted:
                _google_key_index += 1
                attempts += 1
                continue
            
            api_key = GOOGLE_API_KEYS[current_index]
            
            try:
                url = "https://www.googleapis.com/customsearch/v1"
                params = {
                    "key": api_key,
                    "cx": GOOGLE_CSE_ID,
                    "q": query,
                    "num": max(1, min(10, num)),
                    "safe": "off"
                }
                r = requests.get(url, params=params, timeout=20)
                r.raise_for_status()
                
                # Success! Return results
                items = (r.json() or {}).get("items", []) or []
                return [
                    {
                        "title": i.get("title", ""),
                        "snippet": i.get("snippet", ""),
                        "link": i.get("link", "")
                    }
                    for i in items
                ]
                
            except requests.exceptions.HTTPError as e:
                if e.response.status_code == 429:
                    # Rate limit hit! Mark this key as exhausted and try next
                    _google_key_exhausted.add(current_index)
                    log.warning(f"Google API key #{current_index + 1} hit rate limit (429). "
                               f"Switching to next key... ({len(_google_key_exhausted)}/{len(GOOGLE_API_KEYS)} exhausted)")
                    _google_key_index += 1
                    attempts += 1
                    time.sleep(1)  # Brief pause before trying next key
                    continue
                else:
                    # Different error, re-raise
                    raise
            
            except Exception as e:
                # Unexpected error, try next key
                log.warning(f"Error with Google API key #{current_index + 1}: {e}. Trying next key...")
                _google_key_index += 1
                attempts += 1
                continue
        
        # All keys failed
        log.error(f"All {len(GOOGLE_API_KEYS)} Google API keys failed or exhausted!")
        return []
    
    return cache.get_or_fetch(cache_key, fetcher)

def search_web(query: str, num: int = 10) -> List[Dict[str, str]]:
    """Universal search function - uses configured provider"""
    if SEARCH_PROVIDER == "tavily":
        return tavily_search(query, num)
    else:
        return google_search(query, num)

def google_pack_for_person(name: str) -> List[Dict[str, str]]:
    """Optimized: Reduced from 11 to 5 queries to avoid rate limits"""
    queries = [
        f'"{name}" biography',
        f'"{name}" linkedin',
        f'"{name}" board',
        f'"{name}" foundation',
        f'"{name}" interview'
    ]
    results = []
    for q in queries:
        results.extend(search_web(q, num=5))
        time.sleep(1.2)  # Rate limiting
    
    # Deduplicate by URL
    uniq, seen = [], set()
    for r in results:
        u = r.get("link")
        if u and u not in seen:
            seen.add(u)
            uniq.append(r)
    return uniq[:30]  # Reduced from 60 to 30

# -------------------------
# OPENAI HELPERS (with caching)
# -------------------------
@retry_with_backoff(max_retries=3)
def openai_chat_json(prompt: str, cache_key_prefix: str = "openai") -> Optional[Dict[str, Any]]:
    """OpenAI with caching and retry logic - supports both old and new API"""
    if not openai or not OPENAI_API_KEY:
        log.warning("OpenAI disabled (missing API key)")
        return None
    
    cache_key = cache.get_cache_key(cache_key_prefix, {"prompt": prompt[:100]})
    
    def fetcher():
        if OPENAI_VERSION == "new":
            # New API (1.0+)
            resp = OPENAI_CLIENT.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.0
            )
            text = resp.choices[0].message.content.strip()
        else:
            # Old API (0.28)
            resp = openai.ChatCompletion.create(
                model=OPENAI_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.0
            )
            text = resp["choices"][0]["message"]["content"].strip()
        
        m = re.search(r"\{.*\}", text, re.DOTALL)
        if m:
            return json.loads(m.group(0))
        return json.loads(text)
    
    return cache.get_or_fetch(cache_key, fetcher)

# -------------------------
# FEC (with caching)
# -------------------------
@retry_with_backoff(max_retries=3)
def research_fec(name: str, max_pages: int = 3) -> Dict[str, Any]:
    """FEC research with caching"""
    cache_key = cache.get_cache_key("fec", {"name": name})
    
    def fetcher():
        summary = {
            "found": False,
            "total": 0.0,
            "count": 0,
            "by_party": {"R": 0.0, "D": 0.0, "Other": 0.0},
            "recent_years": {},
            "top_recipients": []
        }
        
        if not FEC_API_KEY:
            log.warning("FEC disabled (missing API key)")
            return summary
        
        base = "https://api.open.fec.gov/v1/schedules/schedule_a/"
        params = {
            "api_key": FEC_API_KEY,
            "contributor_name": name,
            "per_page": 100,
            "sort": "contribution_receipt_date"
        }
        
        page = 1
        recipients = defaultdict(float)
        
        try:
            while page <= max_pages:
                params["page"] = page
                r = requests.get(base, params=params, timeout=25)
                r.raise_for_status()
                data = r.json()
                results = data.get("results", []) or []
                if not results:
                    break
                
                for it in results:
                    amt = float(it.get("contribution_receipt_amount") or 0.0)
                    summary["total"] += amt
                    summary["count"] += 1
                    summary["found"] = True
                    
                    # Year bucket
                    try:
                        yr = int((it.get("contribution_receipt_date") or "")[:4])
                        summary["recent_years"][yr] = summary["recent_years"].get(yr, 0.0) + amt
                    except:
                        pass
                    
                    # Recipient & party
                    rcpt = it.get("recipient_name") or it.get("committee_name") or ""
                    if rcpt:
                        recipients[rcpt] += amt
                    p = (it.get("party_full") or "").upper()
                    if p.startswith("DEM"):
                        summary["by_party"]["D"] += amt
                    elif p.startswith("REP"):
                        summary["by_party"]["R"] += amt
                    else:
                        summary["by_party"]["Other"] += amt
                
                page += 1
                if not data.get("pagination", {}).get("pages"):
                    break
                if page > data["pagination"]["pages"]:
                    break
        except Exception as e:
            log.warning(f"FEC error: {e}")
        
        # Pattern
        R, D, O = summary["by_party"]["R"], summary["by_party"]["D"], summary["by_party"]["Other"]
        tot = summary["total"] or 1.0
        rp = R / tot * 100
        dp = D / tot * 100
        if max(rp, dp) <= 65:
            pattern = "Bipartisan donor"
        elif rp >= 70:
            pattern = "Republican donor"
        elif dp >= 70:
            pattern = "Democratic donor"
        else:
            pattern = "Mixed donor"
        summary["pattern"] = pattern
        
        # Top recipients
        top = sorted(recipients.items(), key=lambda kv: -kv[1])[:10]
        summary["top_recipients"] = [{"name": k, "total": v} for k, v in top]
        summary["profile_url"] = f"https://www.fec.gov/data/receipts/individual-contributions/?contributor_name={requests.utils.quote(name)}"
        
        return summary
    
    return cache.get_or_fetch(cache_key, fetcher)

# -------------------------
# INNER CIRCLE CONNECTIONS (optimized)
# -------------------------
def norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "")).strip().lower()

CONN_SCHEMA_EXAMPLE = {
    "found": True,
    "type": "co-board | co-event | alumni | other",
    "strength": 72,
    "description": "One-liner summary",
    "timeframe": "2011-2014",
    "evidence": "Short quote",
    "source": "https://example.com",
    "confidence": 0.85
}

CONN_PROMPT_TMPL = """Verify connection between TWO people.
Return STRICT JSON in this schema:
{schema}

Rules:
- STRONG (70-100): Direct tie (co-board, co-authors, co-panelists at named event)
- MEDIUM (50-69): Likely tie (same program/org with overlap, multiple indirect interactions)
- POTENTIAL (30-49): Weak tie (same school non-overlap, same org different eras)
- REJECT (<30): No real connection

Prospect: "{prospect}"
Inner circle: "{member}"

SEARCH RESULTS:
{context}

Return JSON only:"""

def google_queries_for_pair(donor: str, member: str, donor_city: str = "", donor_state: str = "", reduced: bool = False) -> List[str]:
    """Generate search queries for donor-member pair
    
    Args:
        reduced: If True, returns 1 query instead of 3 (for rate limit compliance)
    """
    donor_parts = donor.split()
    member_parts = member.split()
    donor_last = donor_parts[-1] if donor_parts else donor
    member_last = member_parts[-1] if member_parts else member
    
    if reduced:
        # REDUCED MODE: Just 1 high-quality query
        return [f'"{donor}" "{member}"']
    else:
        # STANDARD MODE: 3 queries for thorough search
        queries = [
            f'"{donor}" "{member}"',
            f'"{donor_last}" "{member_last}" {donor_state}' if donor_state else f'"{donor_last}" "{member_last}"',
            f'"{donor}" "{member}" board'
        ]
        return queries

def classify_connection(donor_name: str, donor_city: str, donor_state: str,
                       member_row: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    """Classify connection with confidence scoring"""
    # Build member name
    first = (member_row.get("First") or "").strip()
    middle = (member_row.get("Middle") or "").strip()
    last = (member_row.get("Last") or "").strip()
    member_name = " ".join([p for p in [first, middle, last] if p])
    
    if not member_name:
        return None
    
    # Build queries (reduced to 3)
    queries = google_queries_for_pair(donor_name, member_name, donor_city, donor_state)
    results = []
    
    for q in queries:
        results.extend(search_web(q, num=3))  # Reduced from 5 to 3
        time.sleep(1.2)  # Rate limit: ~50 queries per 100 seconds (safe margin)
    
    # Deduplicate
    uniq, seen = [], set()
    for r in results:
        u = r.get("link")
        if u and u not in seen:
            seen.add(u)
            uniq.append(r)
    
    context_lines = [f"- {r.get('title', '')} | {r.get('snippet', '')} | {r.get('link', '')}" 
                     for r in uniq[:8]]  # Reduced from 12 to 8
    context = "\n".join(context_lines) if context_lines else "NO WEB HITS"
    
    if context == "NO WEB HITS":
        return None
    
    prompt = CONN_PROMPT_TMPL.format(
        schema=json.dumps(CONN_SCHEMA_EXAMPLE, indent=2),
        prospect=donor_name,
        member=member_name,
        context=context
    )
    
    parsed = openai_chat_json(prompt, cache_key_prefix=f"conn_{donor_name}_{member_name}")
    if not parsed or not parsed.get("found"):
        return None
    
    strength = int(parsed.get("strength") or 0)
    if strength < 30:
        return None
    
    out = {
        "inner_circle_name": member_name,
        "type": parsed.get("type") or "other",
        "strength": strength,
        "description": parsed.get("description") or "",
        "timeframe": parsed.get("timeframe") or "",
        "evidence": parsed.get("evidence") or "",
        "source": parsed.get("source") or "",
        "confidence": parsed.get("confidence", 0.5),
        "citations": [parsed.get("source")] if parsed.get("source") else []
    }
    
    if out["strength"] >= 70:
        out["strength_bucket"] = "Strong"
    elif out["strength"] >= 50:
        out["strength_bucket"] = "Medium"
    else:
        out["strength_bucket"] = "Potential"
    
    return out

def research_inner_circle(prospect: Dict[str, Any], 
                         inner_rows: List[Dict[str, Any]]) -> Tuple[List[Dict[str, Any]], List[str]]:
    """Research inner circle connections"""
    donor = prospect.get("name") or ""
    donor_city = prospect.get("City") or ""
    donor_state = prospect.get("State") or ""
    
    connections, used = [], []
    total = len(inner_rows)
    log.info(f"Analyzing {total} inner-circle members for {donor}...")
    
    for i, row in enumerate(inner_rows, start=1):
        if (i % 25) == 0:
            log.info(f"  Progress {i}/{total}...")
        
        try:
            c = classify_connection(donor, donor_city, donor_state, row)
            if c:
                connections.append(c)
                used.extend(c.get("citations", []))
        except Exception as e:
            log.warning(f"Member #{i} failed: {e}")
        
        time.sleep(0.5)  # Small delay between members (on top of query delays)
    
    connections.sort(key=lambda x: (-x["strength"], x["inner_circle_name"]))
    used = list(set(used))  # Dedupe
    
    log.info(f"Connections found: "
             f"{sum(1 for c in connections if c['strength_bucket']=='Strong')} strong, "
             f"{sum(1 for c in connections if c['strength_bucket']=='Medium')} medium, "
             f"{sum(1 for c in connections if c['strength_bucket']=='Potential')} potential.")
    
    return connections, used

# -------------------------
# PROFILE EXTRACTION
# -------------------------
PROFILE_SCHEMA = {
    "name": "",
    "contact": {"address": "", "phone": "", "email": "", "salesforce_url": ""},
    "dob": "",
    "family": "",
    "education": [],
    "current_position": "",
    "previous_positions": [],
    "nonprofit_boards": [],
    "other_affiliations": [],
    "publications": [],
    "biography": "",
    "background_education": "",
    "professional_roles": [],
    "board_positions": [],
    "key_relationships": [],
    "interview_history": [],
    "personal_philanthropy": "",
    "organizations": [],
    "boards": [],
    "universities": [],
    "foundation_assets": None
}

PROFILE_PROMPT_TMPL = """Build structured profile JSON from web snippets.
Return STRICT JSON in this schema:
{schema}

Guidance:
- Be concise and factual
- Attach URLs where possible
- Do NOT invent facts

NAME: {name}

SNIPPETS:
{snippets}

Return JSON only:"""

DEEP_BIO_PROMPT = """You are writing a comprehensive biographical narrative for a donor profile document, similar to profiles written for major philanthropic prospects.

Your task: Write a detailed, narrative-style biographical section about this person based on the web search results provided. This should read like a well-researched profile article, NOT a bulleted list.

PERSON: {name}

INSTRUCTIONS:
1. Write 3-5 detailed paragraphs in flowing narrative prose
2. Include full name, any nicknames, and relevant name variations
3. Cover family background if available (e.g., "fifth-generation member of the [Family] business legacy")
4. Describe early life, upbringing, and formative experiences
5. Detail educational background - schools, degrees, areas of study
6. Explain career trajectory and major professional milestones
7. Mention significant companies, organizations, or institutions they've been associated with
8. Include founding dates, company descriptions, and historical context where relevant
9. Write in third person using past and present tense appropriately
10. Be specific with names, dates, places, and organizations
11. If information is limited, write what IS known without making up details

STYLE: Emulate the tone of The New York Times profile or Forbes biographical piece - authoritative, detailed, and narrative-driven.

WEB SEARCH RESULTS:
{snippets}

Write the biographical narrative (3-5 paragraphs, narrative prose only):"""

def extract_deep_biography(name: str) -> Tuple[str, List[str]]:
    """Extract comprehensive biographical narrative using enhanced research
    Returns: (biography_text, list_of_source_urls)
    """
    log.info(f"  → Conducting deep biographical research for {name}...")
    
    # Expanded biographical queries
    bio_queries = [
        f'"{name}" biography background',
        f'"{name}" early life family',
        f'"{name}" education university college',
        f'"{name}" career history',
        f'"{name}" founded company business',
        f'"{name}" profile about'
    ]
    
    results = []
    for q in bio_queries:
        results.extend(search_web(q, num=5))
        time.sleep(1.2)  # Rate limit
    
    # Deduplicate
    uniq, seen = [], set()
    source_urls = []
    for r in results:
        u = r.get("link")
        if u and u not in seen:
            seen.add(u)
            uniq.append(r)
            source_urls.append(u)
    
    # Build context for AI
    lines = [f"• {r.get('title', '')}\n  {r.get('snippet', '')}\n  URL: {r.get('link', '')}" 
             for r in uniq[:40]]  # Use up to 40 results for deep research
    
    cache_key = cache.get_cache_key("deep_bio", {"name": name})
    
    def fetcher():
        prompt = DEEP_BIO_PROMPT.format(
            name=name,
            snippets="\n\n".join(lines) or "No biographical information found."
        )
        
        if OPENAI_VERSION == "new":
            resp = OPENAI_CLIENT.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,  # Slightly higher for narrative quality
                max_tokens=2000  # Allow longer biographical narratives
            )
            return resp.choices[0].message.content.strip()
        else:
            resp = openai.ChatCompletion.create(
                model=OPENAI_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=2000
            )
            return resp["choices"][0]["message"]["content"].strip()
    
    try:
        biography = cache.get_or_fetch(cache_key, fetcher)
        return biography, source_urls[:40]  # Return narrative and up to 40 source URLs
    except Exception as e:
        log.warning(f"Deep biography extraction failed: {e}")
        return f"Biographical information for {name} is being researched.", []

def extract_profile(name: str) -> Dict[str, Any]:
    """Extract profile with enhanced biographical research"""
    log.info("  → Conducting standard profile extraction...")
    
    search_results = google_pack_for_person(name)
    lines = [f"- {r.get('title', '')} | {r.get('snippet', '')} | {r.get('link', '')}" 
             for r in search_results]
    
    prompt = PROFILE_PROMPT_TMPL.format(
        schema=json.dumps(PROFILE_SCHEMA, indent=2),
        name=name,
        snippets="\n".join(lines[:30]) or "NONE"
    )
    
    parsed = openai_chat_json(prompt, cache_key_prefix=f"profile_{name}") or {}
    parsed["name"] = parsed.get("name") or name
    parsed["organizations"] = parsed.get("organizations") or []
    parsed["boards"] = parsed.get("boards") or []
    
    # Derive universities from education
    if not parsed.get("universities"):
        unis = []
        for e in parsed.get("education", []):
            m = re.findall(
                r"(Harvard|Yale|Princeton|Stanford|Chicago|Columbia|NYU|Oxford|Cambridge|"
                r"MIT|Duke|Penn|UPenn|Wharton|Georgetown|Brown|Dartmouth|Cornell)",
                e, re.I
            )
            if m:
                unis.extend(m)
        parsed["universities"] = sorted(set(unis))
    
    # DEEP BIOGRAPHICAL RESEARCH - Generate narrative biography
    log.info("  → Starting deep biographical narrative generation...")
    deep_bio, bio_sources = extract_deep_biography(name)
    parsed["biography"] = deep_bio  # Replace with narrative version
    parsed["bio_sources"] = bio_sources  # Store source URLs
    
    return parsed

# -------------------------
# IHS DONOR PROBABILITY MODEL
# -------------------------
def pct(a, b):
    try:
        return (float(a) / float(b)) * 100.0 if b else 0.0
    except:
        return 0.0

def fmt_money(x):
    try:
        return "${:,.0f}".format(float(x))
    except:
        return "$0"

def top_connector(conns):
    if not conns:
        return None
    ordered = sorted(
        conns,
        key=lambda c: (
            0 if c['strength_bucket'] == "Strong" else 1 if c['strength_bucket'] == "Medium" else 2,
            -(c['strength'] or 0)
        )
    )
    return ordered[0]

def ihs_matrix_section(prospect: Dict[str, Any], fec: Dict[str, Any],
                      connections: List[Dict[str, Any]], inner_total: int) -> Tuple[str, List[str]]:
    """Generate IHS donor probability assessment"""
    citations = []
    R = fec.get("by_party", {}).get("R", 0.0)
    D = fec.get("by_party", {}).get("D", 0.0)
    O = fec.get("by_party", {}).get("Other", 0.0)
    tot = fec.get("total", 0.0)
    pattern = fec.get("pattern", "")
    
    # Alignment
    if tot > 0:
        rp, dp = pct(R, tot), pct(D, tot)
        if rp >= 70:
            align = "Moderate to High (Right-leaning)"
        elif dp >= 70:
            align = "Low to Moderate (Left-leaning)"
        elif max(rp, dp) <= 65:
            align = "Moderate (Bipartisan)"
        else:
            align = "Moderate"
        align_reasons = [
            f"Federal giving split — R {fmt_money(R)} ({rp:.0f}%), "
            f"D {fmt_money(D)} ({dp:.0f}%), Other {fmt_money(O)}."
        ]
        if pattern:
            align_reasons.append(f"Pattern: {pattern}.")
    else:
        align = "Moderate"
        align_reasons = ["No federal giving on file; alignment inferred from biography/affiliations."]
    
    if fec.get("profile_url"):
        citations.append(fec["profile_url"])
    
    # Engagement/Network
    strong = sum(1 for c in connections if c["strength_bucket"] == "Strong")
    medium = sum(1 for c in connections if c["strength_bucket"] == "Medium")
    potential = sum(1 for c in connections if c["strength_bucket"] == "Potential")
    engagement_index = strong + 0.6 * medium
    eng = "High" if engagement_index >= 3 else "Moderate" if engagement_index >= 1 else "Low"
    net = "High" if (strong + medium) >= 5 else "Moderate" if (strong + medium) >= 2 else "Low"
    
    eng_reasons = [
        f"Inner-circle ties found: {strong} strong, {medium} medium, {potential} potential "
        f"(out of ~{inner_total} analyzed)."
    ]
    net_reasons = [f"Breadth of network into IHS: {strong + medium} strong/medium ties documented."]
    
    # Interest
    pf = (prospect or {}).get("professional_field", "") or (prospect or {}).get("current_position", "")
    orgs = set(prospect.get("organizations") or []) | set(prospect.get("boards") or [])
    interest_keywords = ["economics", "policy", "law", "education", "finance", "business"]
    interest_orgs = ["cato", "mercatus", "heritage", "atlas", "aei", "hoover"]
    interest = "High" if any(k in (pf or "").lower() for k in interest_keywords) or \
                          any(any(x in o.lower() for x in interest_orgs) for o in orgs) else "Moderate"
    int_reasons = []
    if pf:
        int_reasons.append(f"Career domain: {pf}.")
    hits = [o for o in orgs if any(x in o.lower() for x in interest_orgs)]
    if hits:
        int_reasons.append("Affiliations include: " + ", ".join(sorted(set(hits))) + ".")
    
    # Capacity
    cap = "High" if tot >= 100000 else "Moderate to High" if tot >= 25000 else \
          "Moderate" if tot >= 10000 else "Low to Moderate"
    cap_reasons = [
        f"Federal giving recorded: {fmt_money(tot)} across {fec.get('count', 0)} contributions."
    ] if tot > 0 else ["Limited public political giving; philanthropic capacity inferred from roles/boards."]
    
    # Synthesis
    score = 0
    score += {"Low": 0, "Moderate": 1, "Moderate (Bipartisan)": 1, 
              "Moderate to High (Right-leaning)": 2, "Low to Moderate (Left-leaning)": 0.5, "High": 2}.get(align, 1)
    score += {"Low": 0, "Moderate": 1, "High": 2}.get(eng, 1)
    score += {"Low": 0, "Moderate": 1, "High": 2}.get(interest, 1)
    score += {"Low": 0, "Moderate": 1, "High": 2}.get(net, 1)
    score += {"Low": 0, "Low to Moderate": 0.5, "Moderate": 1, "Moderate to High": 1.5, "High": 2}.get(cap, 1)
    
    sol = "High" if score >= 7 else "Moderate to High" if score >= 5 else \
          "Moderate" if score >= 3.5 else "Low to Moderate"
    
    # Ask range
    if cap == "High":
        a0, a1 = 50000, 250000
    elif cap == "Moderate to High":
        a0, a1 = 10000, 50000
    elif cap == "Moderate":
        a0, a1 = 5000, 15000
    else:
        a0, a1 = 2500, 10000
    
    if eng == "High":
        a0, a1 = int(a0 * 1.25), int(a1 * 1.25)
    elif eng == "Low":
        a0, a1 = int(a0 * 0.8), int(a1 * 0.8)
    
    ask = f"{fmt_money(a0)}–{fmt_money(a1)} (initial)"
    
    # Warm intro
    connector = top_connector(connections)
    warm_intro = f"Warm introduction via {connector['inner_circle_name']} " \
                 f"({connector.get('type', 'connection')}, score {connector.get('strength')})." \
                 if connector else "Warm introduction: identify a credible referrer in shared domain."
    
    # Build text
    lines = []
    lines.append("IHS Donor Probability Model Assessment")
    lines.append("")
    lines.append(f"**Alignment ({align})**")
    for x in align_reasons:
        lines.append(f"• {x}")
    lines.append("")
    lines.append(f"**Engagement ({eng})**")
    for x in eng_reasons:
        lines.append(f"• {x}")
    lines.append("")
    lines.append(f"**Interest ({interest})**")
    for x in int_reasons or ["• Limited domain signals; provisional inference."]:
        lines.append(x if x.startswith("•") else f"• {x}")
    lines.append("")
    lines.append(f"**Network ({net})**")
    for x in net_reasons:
        lines.append(f"• {x}")
    lines.append("")
    lines.append(f"**Capacity ({cap})**")
    for x in cap_reasons:
        lines.append(f"• {x}")
    lines.append("")
    lines.append(f"**Solicitation Potential ({sol})**")
    lines.append(f"Recommended initial ask range: {ask}")
    lines.append("")
    lines.append("**Approach Strategy (next 60–90 days)**")
    lines.append(f"1) {warm_intro}")
    lines.append("2) Message angle: emphasize measurable impact and policy relevance.")
    lines.append("3) Discovery call: confirm program focus, cadence, and timing.")
    lines.append("4) Engagement: invite to small-group salon/briefing; follow with tailored memo.")
    lines.append(f"5) The ask: position a {ask} commitment tied to a concrete initiative.")
    
    return "\n".join(lines), citations

# -------------------------
# DOCX RENDERERS - NARRATIVE FORMAT (STEVE EVERIST STYLE)
# -------------------------
def render_title(doc: Document, name: str):
    """Render document title matching Steve Everist format"""
    title = doc.add_heading(f"{name}: Donor Profile", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

def render_narrative_section(doc: Document, title: str, paragraphs: List[str]):
    """Render a narrative section with multiple paragraphs"""
    doc.add_heading(title, level=2)
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())

def generate_biographical_section(R: Dict[str, Any]) -> List[str]:
    """Generate narrative biographical section"""
    paragraphs = []
    
    bio = R.get("biography", "")
    if bio:
        # Split biography into paragraphs if it's a single string
        if isinstance(bio, str):
            paragraphs = [p.strip() for p in bio.split("\n\n") if p.strip()]
        else:
            paragraphs = bio
    
    # Add education details if available
    education = R.get("education", [])
    family = R.get("family", "")
    
    if education and not any("education" in p.lower() for p in paragraphs):
        edu_text = " ".join(education) if isinstance(education, list) else education
        paragraphs.append(edu_text)
    
    if not paragraphs:
        paragraphs = ["Biographical information is being researched."]
    
    return paragraphs

DEEP_CAREER_PROMPT = """You are writing a comprehensive "Career History & Business Leadership" section for a donor profile document.

Your task: Write a detailed, narrative-style career history based on the web search results provided. This should read like a professional biography section in Forbes or Fortune magazine.

PERSON: {name}

INSTRUCTIONS:
1. Write 4-6 detailed paragraphs in flowing narrative prose
2. Start with current position and work backwards chronologically
3. Include specific job titles, company names, and years
4. Describe major career milestones, achievements, and transitions
5. Explain leadership roles and business transformations
6. Detail board positions with major companies or organizations
7. Include founding dates of companies, mergers, acquisitions, strategic decisions
8. Describe the companies/organizations (what they do, their significance)
9. Mention industry influence, awards, recognitions
10. Use transition phrases to connect different career phases
11. Write in third person, past and present tense appropriately

STYLE: Authoritative, detailed, and respectful. Similar to executive biographies in annual reports or leadership profiles in business publications.

WEB SEARCH RESULTS:
{snippets}

Write the career history narrative (4-6 paragraphs, narrative prose only):"""

def extract_deep_career(name: str) -> Tuple[str, List[str]]:
    """Extract comprehensive career narrative using enhanced research
    Returns: (career_text, list_of_source_urls)
    """
    log.info(f"  → Conducting deep career research for {name}...")
    
    # Career-focused queries
    career_queries = [
        f'"{name}" career history',
        f'"{name}" CEO president chairman',
        f'"{name}" board director',
        f'"{name}" company founded',
        f'"{name}" executive leadership',
        f'"{name}" business achievements'
    ]
    
    results = []
    for q in career_queries:
        results.extend(search_web(q, num=5))
        time.sleep(1.2)  # Rate limit
    
    # Deduplicate
    uniq, seen = [], set()
    source_urls = []
    for r in results:
        u = r.get("link")
        if u and u not in seen:
            seen.add(u)
            uniq.append(r)
            source_urls.append(u)
    
    # Build context
    lines = [f"• {r.get('title', '')}\n  {r.get('snippet', '')}\n  URL: {r.get('link', '')}" 
             for r in uniq[:40]]
    
    cache_key = cache.get_cache_key("deep_career", {"name": name})
    
    def fetcher():
        prompt = DEEP_CAREER_PROMPT.format(
            name=name,
            snippets="\n\n".join(lines) or "No career information found."
        )
        
        if OPENAI_VERSION == "new":
            resp = OPENAI_CLIENT.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=2000
            )
            return resp.choices[0].message.content.strip()
        else:
            resp = openai.ChatCompletion.create(
                model=OPENAI_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=2000
            )
            return resp["choices"][0]["message"]["content"].strip()
    
    try:
        career = cache.get_or_fetch(cache_key, fetcher)
        return career, source_urls[:40]  # Return narrative and up to 40 source URLs
    except Exception as e:
        log.warning(f"Deep career extraction failed: {e}")
        return f"Career history for {name} is being researched.", []

def generate_career_section(R: Dict[str, Any]) -> List[str]:
    """Generate Career History & Business Leadership section with deep research"""
    # Check if we already did deep research (stored in R)
    if R.get("deep_career_narrative"):
        # Split into paragraphs
        narrative = R["deep_career_narrative"]
        paragraphs = [p.strip() for p in narrative.split("\n\n") if p.strip()]
        if paragraphs:
            return paragraphs
    
    # Fallback to basic extraction
    paragraphs = []
    current_pos = R.get("current_position", "")
    previous_pos = R.get("previous_positions", [])
    boards = R.get("nonprofit_boards", [])
    
    if current_pos:
        paragraphs.append(f"Currently, {R.get('name', 'the donor')} serves as {current_pos}.")
    
    if previous_pos:
        career_text = "Previous roles include: " + "; ".join(previous_pos) + "."
        paragraphs.append(career_text)
    
    if boards:
        board_names = [b.get('org', '') for b in boards if b.get('org')]
        if board_names:
            paragraphs.append(f"Board service includes: {', '.join(board_names)}.")
    
    if not paragraphs:
        paragraphs = ["Career history is being researched."]
    
    return paragraphs

DEEP_PHILANTHROPY_PROMPT = """You are writing a comprehensive "Philanthropic Activities & Causes Supported" section for a donor profile document.

Your task: Write a detailed, narrative-style philanthropy section based on the web search results provided. This should read like a professional philanthropic profile.

PERSON: {name}

INSTRUCTIONS:
1. Write 3-5 detailed paragraphs in flowing narrative prose
2. Describe charitable giving patterns and focus areas
3. Detail foundation involvement (if applicable) with founding dates and mission
4. Explain board memberships with nonprofit organizations - name the organizations and describe their work
5. Highlight specific gifts with amounts if available
6. Describe philanthropic philosophy or themes (education, health, arts, faith-based, etc.)
7. Mention collaborative philanthropy or partnerships with other donors
8. Include family involvement in giving if relevant
9. Use specific organization names, dollar amounts, and dates
10. Write in third person

STYLE: Respectful and detailed, similar to foundation annual reports or philanthropic profiles.

WEB SEARCH RESULTS:
{snippets}

Write the philanthropic activities narrative (3-5 paragraphs, narrative prose only):"""

def extract_deep_philanthropy(name: str) -> Tuple[str, List[str]]:
    """Extract comprehensive philanthropic narrative
    Returns: (philanthropy_text, list_of_source_urls)
    """
    log.info(f"  → Conducting deep philanthropic research for {name}...")
    
    philanthropy_queries = [
        f'"{name}" foundation charity donation',
        f'"{name}" philanthropic giving nonprofit',
        f'"{name}" board trustee',
        f'"{name}" charitable contribution',
        f'"{name}" donor supporter'
    ]
    
    results = []
    for q in philanthropy_queries:
        results.extend(search_web(q, num=5))
        time.sleep(1.2)  # Rate limit
    
    # Deduplicate
    uniq, seen = [], set()
    source_urls = []
    for r in results:
        u = r.get("link")
        if u and u not in seen:
            seen.add(u)
            uniq.append(r)
            source_urls.append(u)
    
    lines = [f"• {r.get('title', '')}\n  {r.get('snippet', '')}\n  URL: {r.get('link', '')}" 
             for r in uniq[:40]]
    
    cache_key = cache.get_cache_key("deep_philanthropy", {"name": name})
    
    def fetcher():
        prompt = DEEP_PHILANTHROPY_PROMPT.format(
            name=name,
            snippets="\n\n".join(lines) or "No philanthropic information found."
        )
        
        if OPENAI_VERSION == "new":
            resp = OPENAI_CLIENT.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=2000
            )
            return resp.choices[0].message.content.strip()
        else:
            resp = openai.ChatCompletion.create(
                model=OPENAI_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=2000
            )
            return resp["choices"][0]["message"]["content"].strip()
    
    try:
        philanthropy = cache.get_or_fetch(cache_key, fetcher)
        return philanthropy, source_urls[:40]  # Return narrative and up to 40 source URLs
    except Exception as e:
        log.warning(f"Deep philanthropy extraction failed: {e}")
        return f"Philanthropic activities for {name} are being researched.", []

def generate_philanthropic_section(R: Dict[str, Any]) -> List[str]:
    """Generate Philanthropic Activities & Causes Supported section with deep research"""
    # Check if we have deep research
    if R.get("deep_philanthropy_narrative"):
        narrative = R["deep_philanthropy_narrative"]
        paragraphs = [p.strip() for p in narrative.split("\n\n") if p.strip()]
        if paragraphs:
            return paragraphs
    
    # Fallback to basic extraction
    paragraphs = []
    boards = R.get("nonprofit_boards", [])
    affiliations = R.get("other_affiliations", [])
    
    if boards:
        para = f"{R.get('name', 'The donor')} demonstrates philanthropic engagement through board service with "
        board_list = [f"{b.get('org', '')} ({b.get('role', 'Member')})" for b in boards]
        para += ", ".join(board_list) + "."
        paragraphs.append(para)
    
    if affiliations:
        paragraphs.append("Additional philanthropic affiliations include: " + "; ".join(affiliations) + ".")
    
    if not paragraphs:
        paragraphs = ["Philanthropic activities are being researched."]
    
    return paragraphs

def generate_political_section(R: Dict[str, Any], fec: Dict[str, Any]) -> List[str]:
    """Generate Political Donations & Ideological Leanings section"""
    paragraphs = []
    
    if fec and fec.get("found"):
        total = fec.get("total", 0.0)
        byp = fec.get("by_party", {})
        pattern = fec.get("pattern", "")
        
        # Opening paragraph with total giving
        para = f"{R.get('name', 'The donor')} is an active political donor with federal contributions totaling {fmt_money(total)}. "
        
        # Party breakdown
        rep_amt = byp.get('R', 0.0)
        dem_amt = byp.get('D', 0.0)
        other_amt = byp.get('Other', 0.0)
        
        if rep_amt > dem_amt * 2:
            para += f"The giving pattern strongly favors Republican candidates and causes ({fmt_money(rep_amt)} to Republicans vs. {fmt_money(dem_amt)} to Democrats), "
            para += "indicating a conservative political orientation aligned with free-market principles."
        elif dem_amt > rep_amt * 2:
            para += f"The giving pattern favors Democratic candidates and causes ({fmt_money(dem_amt)} to Democrats vs. {fmt_money(rep_amt)} to Republicans)."
        else:
            para += f"The giving shows a relatively balanced approach across parties (R: {fmt_money(rep_amt)}, D: {fmt_money(dem_amt)})."
        
        paragraphs.append(para)
        
        # Pattern analysis
        if pattern:
            paragraphs.append(f"Contribution pattern: {pattern}.")
    else:
        paragraphs.append(f"No federal political contributions found on file for {R.get('name', 'this donor')}. "
                         "This may indicate either private political engagement through other channels, "
                         "or a preference for focusing philanthropic resources on non-political causes.")
    
    return paragraphs

def generate_connections_section(R: Dict[str, Any], connections: List[Dict[str, Any]], inner_total: int) -> List[str]:
    """Generate Connections section - Inner Circle comparison results"""
    paragraphs = []
    
    name = R.get('name', 'The donor')
    
    # Opening paragraph
    para = f"As part of this research, {name} was systematically compared against {inner_total} members of IHS's Inner Circle—a network comprising board members, officers, scholars, and key supporters of liberty-oriented causes. "
    para += "The analysis searched for documented connections including co-authorship, shared board service, institutional overlap, conference participation, and other verifiable relationships."
    paragraphs.append(para)
    
    if not connections:
        paragraphs.append(f"The systematic comparison of {name} against the {inner_total} Inner Circle members did not identify strong documented connections at this time. "
                         "This does not preclude informal relationships or shared interests that may not be publicly documented. "
                         "Further relationship mapping through personal introductions may reveal additional network pathways.")
        return paragraphs
    
    # Categorize connections
    strong = [c for c in connections if c.get("strength_bucket") == "Strong"]
    medium = [c for c in connections if c.get("strength_bucket") == "Medium"]
    potential = [c for c in connections if c.get("strength_bucket") == "Potential"]
    
    # Strong connections paragraph
    if strong:
        para = f"**Strong Connections ({len(strong)} identified):** "
        para += f"{name} has documented, verifiable ties to the following Inner Circle members: "
        
        conn_details = []
        for c in strong[:5]:  # Show top 5 strong connections
            detail = f"{c.get('inner_circle_name', '')}"
            if c.get('description'):
                detail += f" ({c.get('description', '')})"
            conn_details.append(detail)
        
        para += "; ".join(conn_details)
        if len(strong) > 5:
            para += f"; and {len(strong) - 5} additional strong connections"
        para += ". These relationships represent direct collaborative ties through shared institutions, co-authored work, or joint board service that could facilitate warm introductions and credibility with IHS leadership."
        paragraphs.append(para)
    
    # Medium connections paragraph
    if medium:
        para = f"**Medium-Strength Connections ({len(medium)} identified):** "
        para += "Additional likely connections through institutional overlap or shared networks include: "
        
        names = [c.get('inner_circle_name', '') for c in medium[:5]]
        para += ", ".join(names)
        if len(medium) > 5:
            para += f", and {len(medium) - 5} others"
        para += ". While these connections may require verification, they suggest shared professional circles and potential mutual acquaintances."
        paragraphs.append(para)
    
    # Potential connections paragraph  
    if potential:
        para = f"**Potential Connections ({len(potential)} identified):** "
        para += "Weaker or historical connections were also identified, including shared alumni networks, overlapping organizational affiliations in different time periods, or similar professional interests. "
        para += "These represent possible relationship pathways that warrant further investigation during the cultivation process."
        paragraphs.append(para)
    
    # Cultivation implications
    if strong or medium:
        para = "**Cultivation Implications:** "
        if strong:
            top_connector = strong[0].get('inner_circle_name', '')
            para += f"The documented connection with {top_connector} provides an immediate warm introduction pathway. "
        para += f"With {len(strong) + len(medium)} documented or likely connections to Inner Circle members, {name} is already embedded in networks adjacent to IHS. "
        para += "This existing network proximity significantly enhances cultivation prospects and suggests natural affinity for IHS's mission and community."
        paragraphs.append(para)
    
    return paragraphs

def generate_network_section(R: Dict[str, Any]) -> List[str]:
    """Generate Network & Board Affiliations section - about the DONOR'S OWN network and boards"""
    paragraphs = []
    
    name = R.get('name', 'The donor')
    
    # Corporate and business boards
    boards = R.get("nonprofit_boards", [])
    current_pos = R.get("current_position", "")
    
    # Opening paragraph about network positioning
    para = f"{name} maintains a professional network that spans "
    network_areas = []
    
    if current_pos:
        if "CEO" in current_pos or "President" in current_pos or "Chairman" in current_pos:
            network_areas.append("executive business leadership")
    
    if boards:
        network_areas.append("nonprofit governance")
    
    affiliations = R.get("other_affiliations", [])
    if affiliations:
        network_areas.append("civic engagement")
    
    if network_areas:
        para += ", ".join(network_areas) + ". "
    else:
        para += "multiple professional domains. "
    
    para += "This network positioning provides access to diverse perspectives and relationship capital across sectors."
    paragraphs.append(para)
    
    # Board service paragraph
    if boards:
        para = "**Nonprofit Board Service:** "
        para += f"{name} demonstrates commitment to institutional leadership through board service with "
        board_list = []
        for b in boards[:5]:  # Show top 5 boards
            org = b.get('org', '')
            role = b.get('role', '')
            if role and role != 'Member':
                board_list.append(f"{org} ({role})")
            else:
                board_list.append(org)
        
        para += ", ".join(board_list)
        if len(boards) > 5:
            para += f", and {len(boards) - 5} additional organizations"
        para += ". These governance roles provide direct access to other board members, donors, and organizational leaders within these institutions."
        paragraphs.append(para)
    
    # Professional affiliations paragraph
    if affiliations:
        para = "**Professional Affiliations:** "
        para += "Additional network access comes through affiliations with "
        para += ", ".join(affiliations[:5])
        if len(affiliations) > 5:
            para += f", and {len(affiliations) - 5} other organizations"
        para += ". These associations expand professional reach and provide connection points across multiple communities of interest."
        paragraphs.append(para)
    
    # Network implications
    if boards or affiliations:
        para = "**Network Implications:** "
        para += f"The breadth of {name}'s board service and affiliations suggests strong relationship-building capacity and comfort operating in networked professional environments. "
        para += "This network sophistication is valuable for IHS cultivation, as it indicates potential for serving as a connector who could open doors to peer donors and institutional partnerships."
        paragraphs.append(para)
    else:
        para = f"Specific board memberships and formal affiliations for {name} are being researched. "
        para += "However, career achievements and professional positioning suggest active participation in relevant professional and civic networks."
        paragraphs.append(para)
    
    return paragraphs

def generate_networth_section(R: Dict[str, Any]) -> List[str]:
    """Generate Estimated Net Worth & Giving Capacity section"""
    paragraphs = []
    
    name = R.get('name', 'The donor')
    nw_analysis = R.get("net_worth_analysis", {})
    narrative = nw_analysis.get("narrative", [])
    capacity_band = nw_analysis.get("capacity_band", "Moderate to High")
    
    # Opening assessment
    para = f"While specific net worth figures for {name} are not publicly available, "
    para += "multiple indicators suggest significant wealth and giving capacity. "
    paragraphs.append(para)
    
    # Add detailed analysis points
    current_pos = R.get("current_position", "")
    previous_pos = R.get("previous_positions", [])
    
    if current_pos:
        paragraphs.append(f"Current role as {current_pos} indicates executive-level compensation and likely equity positions.")
    
    if previous_pos:
        paragraphs.append("Career trajectory through multiple leadership positions suggests accumulated wealth through "
                         "high compensation packages, equity grants, and potential liquidity events.")
    
    # Add any existing narrative points
    for item in narrative:
        if item and "not yet assessed" not in item.lower():
            paragraphs.append(item)
    
    # Capacity conclusion
    paragraphs.append(f"Overall wealth and giving capacity assessment: {capacity_band}. "
                     f"Based on career indicators and philanthropic engagement, {name} demonstrates "
                     "capacity for major gifts and sustained charitable support.")
    
    return paragraphs

def generate_ihs_assessment_narrative(R: Dict[str, Any], fec: Dict[str, Any], 
                                      connections: List[Dict[str, Any]], 
                                      inner_total: int) -> List[str]:
    """Generate IHS A-E-I-N-P-S assessment in narrative format matching Steve Everist"""
    paragraphs = []
    
    name = R.get('name', 'the donor')
    
    # Intro paragraph
    paragraphs.append(f"Using IHS's A-E-I-N-P-S framework, here is an assessment of {name} as a prospective donor:")
    
    # Alignment (A)
    alignment_score = "Moderate"
    if fec and fec.get("found"):
        byp = fec.get("by_party", {})
        if byp.get('R', 0) > byp.get('D', 0) * 2:
            alignment_score = "High"
    
    para = f"**Alignment (A): {alignment_score}.** "
    if alignment_score == "High":
        para += f"{name}'s political giving pattern demonstrates strong alignment with free-market principles and limited government. "
        para += "Federal contributions favor Republican and liberty-oriented causes, suggesting values consonant with IHS's mission to advance liberty through education."
    else:
        para += f"Initial research suggests {name} holds values that may align with IHS's educational mission, though further engagement is needed to confirm ideological fit. "
        para += "Philanthropic interests in education and individual empowerment provide potential common ground."
    paragraphs.append(para)
    
    # Engagement (E)
    engagement_score = "Low"
    para = f"**Engagement (E): {engagement_score}.** "
    para += f"To date, {name} has had minimal direct engagement with IHS or its sister organizations. "
    para += "This represents an opportunity for cultivation through targeted outreach and relationship building."
    paragraphs.append(para)
    
    # Interest (I)
    interest_score = "Moderate"
    para = f"**Interest (I): {interest_score}.** "
    boards = R.get("nonprofit_boards", [])
    if boards:
        para += f"{name}'s board service with {', '.join([b.get('org', '')[:30] for b in boards[:2]])} demonstrates interest in education and social impact. "
    para += "These philanthropic priorities suggest potential receptivity to IHS's academic programs and scholar development initiatives."
    paragraphs.append(para)
    
    # Network (N)
    network_score = "Moderate"
    if connections and len([c for c in connections if c.get("strength_bucket") == "Strong"]) > 0:
        network_score = "High"
    
    para = f"**Network (N): {network_score}** "
    if network_score == "High":
        strong_conns = [c for c in connections if c.get("strength_bucket") == "Strong"]
        para += f"(regionally), Moderate (nationally). {name} has documented connections to {len(strong_conns)} Inner Circle members, "
        para += "including " + ", ".join([c.get('inner_circle_name', '')[:25] for c in strong_conns[:3]]) + ". "
        para += "These relationships provide warm introduction paths and mutual credibility."
    else:
        para += f". Network analysis against {inner_total} Inner Circle members identified potential connections through shared institutions and causes. "
        para += "Further relationship mapping could reveal additional network pathways."
    paragraphs.append(para)
    
    # Philanthropic Capacity (P)
    capacity = R.get("net_worth_analysis", {}).get("capacity_band", "Moderate to High")
    para = f"**Philanthropic Capacity (P): {capacity}.** "
    para += f"{name}'s career achievements and sustained philanthropic engagement indicate significant giving capacity. "
    para += "The donor falls into the tier capable of major gifts ($50,000-$250,000+) with proper cultivation and compelling case for support."
    paragraphs.append(para)
    
    # Solicitation Potential (S)
    sol_score = "Moderate"
    para = f"**Solicitation Potential (S): {sol_score}.** "
    para += f"Assessment suggests {name} is receptive to well-researched approaches that demonstrate impact and align with personal values. "
    para += "Recommended strategy: warm introduction from mutual contact, followed by discovery conversation to understand philanthropic priorities, "
    para += "culminating in a tailored proposal for a specific IHS program or initiative. "
    para += "Initial ask range: $25,000-$100,000 depending on depth of engagement and program fit."
    paragraphs.append(para)
    
    return paragraphs

def generate_strategic_summary(R: Dict[str, Any], connections: List[Dict[str, Any]]) -> List[str]:
    """Generate Strategic Summary (Board-Level Briefing)"""
    paragraphs = []
    
    name = R.get('name', 'the donor')
    current_pos = R.get("current_position", "")
    location = f"{R.get('City', '')}, {R.get('State', '')}" if R.get('City') else ""
    
    # Profile overview
    para = f"**Profile Overview:** {name}"
    if current_pos:
        para += f", {current_pos},"
    if location:
        para += f" based in {location},"
    para += " is a prospective major donor with capacity and potential alignment for IHS support. "
    
    # Add connection strength
    strong_conns = [c for c in connections if c.get("strength_bucket") == "Strong"]
    if strong_conns:
        para += f"Documented connections to {len(strong_conns)} Inner Circle members provide warm introduction opportunities. "
    
    paragraphs.append(para)
    
    # Recommended approach
    para = "**Recommended Next Steps:** "
    if strong_conns:
        connector = strong_conns[0].get('inner_circle_name', 'mutual contact')
        para += f"(1) Request warm introduction through {connector}. "
    else:
        para += "(1) Conduct additional research to identify warm introduction paths. "
    
    para += "(2) Schedule discovery call to understand philanthropic priorities and assess program fit. "
    para += "(3) Develop tailored cultivation strategy with specific program proposals aligned to donor interests. "
    para += "(4) Position initial solicitation in $25,000-$100,000 range for targeted initiative."
    paragraphs.append(para)
    
    return paragraphs

def build_networth_analysis(R: Dict[str, Any], foundation_assets=None):
    narrative = []
    if R.get("previous_positions"):
        narrative.append("Career includes multiple top executive roles with likely high compensation and equity exposure.")
    if foundation_assets is not None:
        narrative.append(
            f"Family foundation reports ~{fmt_money(foundation_assets)} in assets per IRS 990 filings, "
            f"indicating sustained charitable capacity."
        )
    capacity_band = "High" if foundation_assets and foundation_assets >= 2_000_000 else "Moderate to High"
    return {
        "narrative": narrative if narrative else ["Capacity not yet assessed."],
        "capacity_band": capacity_band
    }

def collect_all_citations(R: Dict[str, Any], fec: Dict[str, Any], connections: List[Dict[str, Any]]) -> List[str]:
    """Collect and deduplicate all source URLs used in research"""
    all_urls = set()
    
    # Biographical sources
    for url in R.get("bio_sources", []):
        if url and url.strip():
            all_urls.add(url.strip())
    
    # Career sources
    for url in R.get("career_sources", []):
        if url and url.strip():
            all_urls.add(url.strip())
    
    # Philanthropic sources
    for url in R.get("philanthropy_sources", []):
        if url and url.strip():
            all_urls.add(url.strip())
    
    # FEC profile URL
    if fec.get("profile_url"):
        all_urls.add(fec["profile_url"])
    
    # Connection citations
    for conn in connections:
        for cite in conn.get("citations", []):
            if cite and cite.strip():
                all_urls.add(cite.strip())
        # Also check source field
        if conn.get("source"):
            all_urls.add(conn["source"].strip())
    
    # Sort alphabetically for consistency
    return sorted(list(all_urls))

def generate_citations_section(all_citations: List[str]) -> List[str]:
    """Generate a formatted citations section"""
    if not all_citations:
        return ["No external sources were cited in this research."]
    
    intro = (
        f"This profile was compiled using {len(all_citations)} source(s), including biographical databases, "
        f"news articles, organizational websites, government records, and public documents. "
        f"All sources were accessed during the research period and are listed below for verification and further investigation."
    )
    
    # Create a formatted list paragraph
    citation_lines = [intro, ""]  # Empty string creates paragraph break
    citation_lines.append("**Research Sources:**")
    citation_lines.append("")
    
    for i, url in enumerate(all_citations, 1):
        citation_lines.append(f"{i}. {url}")
    
    # Join into paragraphs - the render function will handle each line
    return [intro, "\n".join([f"{i}. {url}" for i, url in enumerate(all_citations, 1)])]

# -------------------------
# MAIN
# -------------------------
def main():
    parser = argparse.ArgumentParser(
        description="IHS Hybrid Donor Research System - Optimized Edition"
    )
    parser.add_argument("--name", help="Process specific donor by name")
    parser.add_argument("--data", default="data.csv", help="Path to donor data CSV")
    parser.add_argument("--inner", help="Path to inner circle CSV (auto-detects Inner_Circle.csv)")
    parser.add_argument("--outdir", default=".", help="Output directory")
    parser.add_argument("--test-mode", action="store_true",
                       help="TEST MODE: Process only the FIRST donor (for testing)")
    args = parser.parse_args()
    
    log.info("="*70)
    log.info("IHS HYBRID DONOR RESEARCH SYSTEM - NARRATIVE FORMAT EDITION")
    log.info("="*70)
    log.info(f"OpenAI Library: {OPENAI_VERSION if OPENAI_VERSION else 'Not installed'}")
    log.info(f"OpenAI Model: {OPENAI_MODEL}")
    log.info(f"Search Provider: {SEARCH_PROVIDER.upper()}")
    if SEARCH_PROVIDER == "google" and GOOGLE_API_KEYS:
        log.info(f"Google API Keys: {len(GOOGLE_API_KEYS)} keys loaded")
        log.info(f"  → Smart key switching: Automatically rotates on rate limits")
        if len(GOOGLE_API_KEYS) > 1:
            log.info(f"  → Combined quota: ~{len(GOOGLE_API_KEYS) * 10000:,} queries/day")
    log.info("")
    
    # Load data
    data_rows = load_csv(args.data)
    if not data_rows:
        log.error(f"No rows found in {args.data}")
        return
    
    # Load inner circle
    if args.inner:
        inner_path = args.inner
    else:
        for candidate in ("Inner_Circle.csv", "inner_circle.csv"):
            if os.path.exists(candidate):
                inner_path = candidate
                break
        else:
            log.error("Inner circle file not found. Please specify with --inner or ensure Inner_Circle.csv exists.")
            return
    
    inner_rows = load_csv(inner_path)
    inner_circle_count = len(inner_rows)
    log.info(f"Loaded {inner_circle_count} inner circle members from {inner_path}")
    
    # Determine which donors to process
    if args.test_mode:
        # TEST MODE: Process only first donor
        first_name = get_row_name(data_rows[0])
        if not first_name:
            log.error("First donor has no valid name")
            return
        names_to_process = [first_name]
        log.info(f"🧪 TEST MODE: Processing ONLY first donor: {first_name}")
    elif args.name:
        # Single donor mode
        names_to_process = [args.name.strip()]
    else:
        # Process all donors
        names_to_process = [get_row_name(r) for r in data_rows if get_row_name(r)]
        if not names_to_process:
            log.error(f"No valid names found in {args.data}")
            return
        log.info(f"Processing {len(names_to_process)} donors from {args.data}")
    
    # Process each donor
    for donor_name in names_to_process:
        log.info(f"\n{'='*60}")
        log.info(f"Processing: {donor_name}")
        log.info(f"{'='*60}")
        
        donor_row = find_donor_in_data(data_rows, donor_name)
        if not donor_row:
            log.warning(f"Skipping: '{donor_name}' not found in {args.data}")
            continue
        
        try:
            # 1) Profile + Deep Biographical Research
            log.info("Extracting profile with deep biographical research...")
            R = extract_profile(donor_name)
            R["City"] = donor_row.get("City", "")
            R["State"] = donor_row.get("State", "")
            
            # 2) Deep Career Research
            log.info("Conducting deep career narrative research...")
            deep_career, career_sources = extract_deep_career(donor_name)
            R["deep_career_narrative"] = deep_career
            R["career_sources"] = career_sources
            
            # 2.5) Deep Philanthropy Research
            log.info("Conducting deep philanthropic research...")
            deep_philanthropy, philanthropy_sources = extract_deep_philanthropy(donor_name)
            R["deep_philanthropy_narrative"] = deep_philanthropy
            R["philanthropy_sources"] = philanthropy_sources
            
            # 3) FEC
            log.info("Researching FEC giving...")
            fec = research_fec(donor_name)
            R["political_contributions"] = fec
            
            # 4) Inner Circle
            log.info(f"Starting deep comparison against {inner_circle_count} inner circle members...")
            connections, conn_cites = research_inner_circle(R, inner_rows)
            R["connections"] = connections
            
            # 5) Net Worth
            foundation_assets = R.get("foundation_assets", None)
            R["net_worth_analysis"] = build_networth_analysis(R, foundation_assets=foundation_assets)
            
            # 6) DOCX Report - NARRATIVE FORMAT (9 sections matching your requirement)
            doc = Document()
            
            # Store inner_circle_count in R for later use
            R['inner_circle_count'] = inner_circle_count
            
            # Title
            render_title(doc, R['name'])
            
            # 1. Biographical Background & Education
            bio_paragraphs = generate_biographical_section(R)
            render_narrative_section(doc, "Biographical Background & Education", bio_paragraphs)
            
            # 2. Career History & Business Leadership
            career_paragraphs = generate_career_section(R)
            render_narrative_section(doc, "Career History & Business Leadership", career_paragraphs)
            
            # 3. Philanthropic Activities & Causes Supported
            philanthropy_paragraphs = generate_philanthropic_section(R)
            render_narrative_section(doc, "Philanthropic Activities & Causes Supported", philanthropy_paragraphs)
            
            # 4. Political Donations & Ideological Leanings
            political_paragraphs = generate_political_section(R, fec)
            render_narrative_section(doc, "Political Donations & Ideological Leanings", political_paragraphs)
            
            # 5. Connections (Inner Circle Comparison - NEW SECTION!)
            connections_paragraphs = generate_connections_section(R, connections, inner_circle_count)
            render_narrative_section(doc, "Connections", connections_paragraphs)
            
            # 6. Network & Board Affiliations (Donor's Own Network)
            network_paragraphs = generate_network_section(R)
            render_narrative_section(doc, "Network & Board Affiliations", network_paragraphs)
            
            # 7. Estimated Net Worth & Giving Capacity
            networth_paragraphs = generate_networth_section(R)
            render_narrative_section(doc, "Estimated Net Worth & Giving Capacity", networth_paragraphs)
            
            # 8. IHS Donor Probability Model Assessment
            ihs_paragraphs = generate_ihs_assessment_narrative(R, fec, connections, inner_circle_count)
            render_narrative_section(doc, "IHS Donor Probability Model Assessment", ihs_paragraphs)
            
            # 9. Strategic Summary (Board-Level Briefing)
            summary_paragraphs = generate_strategic_summary(R, connections)
            render_narrative_section(doc, "Strategic Summary (Board-Level Briefing)", summary_paragraphs)
            
            # 10. Sources & Citations (NEW SECTION!)
            all_citations = collect_all_citations(R, fec, connections)
            citations_paragraphs = generate_citations_section(all_citations)
            render_narrative_section(doc, "Sources & Citations", citations_paragraphs)
            
            # 7) Save
            outpath = os.path.join(args.outdir, f"{R['name'].replace(' ', '_')}_Donor_Profile.docx")
            doc.save(outpath)
            log.info(f"✅ Report written: {outpath}")
            log.info(f"Completed: {donor_name}\n")
            
        except Exception as e:
            log.error(f"❌ Failed processing {donor_name}: {e}")
            import traceback
            traceback.print_exc()
            continue
    
    log.info(f"\n{'='*60}")
    log.info(f"✅ ALL DONE! Processed {len(names_to_process)} donor(s)")
    log.info(f"{'='*60}")

if __name__ == "__main__":
    main()
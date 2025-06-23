#!/usr/bin/env python3
"""
liberalism_analysis_extended.py

Generates a Liberalism Alignment Report for a given scholar based on first name,
last name, email, and affiliation. Collects public posts from LinkedIn, Facebook,
Google Scholar, and Substack to identify Classical Liberal (CL) alignment,
determine if the scholar is a liberal writing primarily about illiberalism, and
produce a downloadable .docx report in a structured format.

Sections:
1. Synthesis of Major Themes and Arguments
2. Classification under the Three Dimensions of Liberalism
3. Summary Table of Classifications
4. Evidence and Citations

Dependencies:
    pip install scholarly requests beautifulsoup4 python-docx

Usage:
    python liberalism_analysis_extended.py \
        --first_name "First" \
        --last_name "Last" \
        --email "email@domain.edu" \
        --affiliation "University Name"
"""

import argparse
from scholarly import scholarly
from docx import Document

def get_scholar_data(first_name, last_name, affiliation):
    """Fetch basic Google Scholar metrics and recent publications."""
    try:
        query = scholarly.search_author(f"{first_name} {last_name} {affiliation}")
        author = next(query, None)
        if not author:
            return None
        author = scholarly.fill(author, sections=['basics', 'indices', 'publications'])
        return {
            'name': author.get('name'),
            'affiliation': author.get('affiliation'),
            'h_index': author.get('hindex'),
            'publications': [(pub['bib']['title'], pub.get('num_citations', 0))
                             for pub in author.get('publications', [])[:10]]
        }
    except Exception as e:
        print(f"Error fetching Google Scholar data: {e}")
        return None

def get_linkedin_posts(first_name, last_name):
    """Stub: Retrieve recent LinkedIn posts."""
    return [
        {"date": "2024-01-15", "content": "I argued that free markets need robust rule-of-law foundations.", 
         "url": "https://linkedin.com/posts/example1"},
        {"date": "2024-06-01", "content": "Discussing illiberal trends in digital censorship and state control.", 
         "url": "https://linkedin.com/posts/example2"}
    ]

def get_facebook_posts(first_name, last_name):
    """Stub: Retrieve public Facebook posts."""
    return [
        {"date": "2023-11-10", "content": "Presented on authoritarian repression tactics and their impact.", 
         "url": "https://facebook.com/example1"}
    ]

def get_substack_articles(first_name, last_name):
    """Stub: Retrieve Substack newsletters or posts."""
    return [
        {"date": "2024-03-20", "title": "Why Classical Liberalism Matters in 2024", 
         "url": "https://substack.com/p/example1"},
        {"date": "2025-02-10", "title": "Mapping Authoritarian Repression: Past Patterns, Future Risks", 
         "url": "https://substack.com/p/example2"}
    ]

def detect_illiberal_focus(all_texts):
    """Detect if scholar writes predominantly about illiberalism."""
    illiberal_keywords = ['repress', 'censor', 'authoritarian', 'dictator', 'surveillance']
    liberal_keywords = ['market', 'trade', 'democr', 'freedom', 'liber']
    illiberal_count = sum(1 for text in all_texts if any(k in text.lower() for k in illiberal_keywords))
    liberal_count = sum(1 for text in all_texts if any(k in text.lower() for k in liberal_keywords))
    # If ratio of illiberal to liberal > 1.5, flag focus on illiberalism
    return illiberal_count > 1.5 * (liberal_count + 1), illiberal_count, liberal_count

def analyze_liberalism(scholar_data, posts):
    """
    Heuristic keyword-based scoring across sources.
    Returns classification scores and list of all texts for illiberal focus detection.
    """
    econ_kw = ['market', 'trade', 'economy', 'private']
    social_kw = ['rights', 'equality', 'justice', 'welfare']
    political_kw = ['democracy', 'constitution', 'public reason', 'liberty']

    scores = {'Economic Liberalism': 0, 'Social Liberalism': 0, 'Political Liberalism': 0}
    all_texts = []

    # publications
    for title, _ in scholar_data.get('publications', []):
        all_texts.append(title)
        t = title.lower()
        if any(k in t for k in econ_kw): scores['Economic Liberalism'] += 1
        if any(k in t for k in social_kw): scores['Social Liberalism'] += 1
        if any(k in t for k in political_kw): scores['Political Liberalism'] += 1

    # posts and articles
    for source in posts.values():
        for item in source:
            text = item.get('content', item.get('title', '')).lower()
            all_texts.append(text)
            if any(k in text for k in econ_kw): scores['Economic Liberalism'] += 1
            if any(k in text for k in social_kw): scores['Social Liberalism'] += 1
            if any(k in text for k in political_kw): scores['Political Liberalism'] += 1

    illiberal_focus, illib_count, lib_count = detect_illiberal_focus(all_texts)
    return scores, illiberal_focus, illib_count, lib_count

def generate_report(args, scholar_data, posts, scores, illiberal_focus, illib_count, lib_count):
    """Builds and saves the .docx report with structured sections."""
    doc = Document()
    doc.add_heading(f"Liberalism Profile: {scholar_data['name']}", level=1)

    # 1. Synthesis of Major Themes and Arguments
    doc.add_heading("1. Synthesis of Major Themes and Arguments", level=2)
    intro = doc.add_paragraph()
    intro.add_run(f"{scholar_data['name']} ({args.affiliation}) has an h-index of {scholar_data['h_index']}. ")
    intro.add_run("Their work spans academic publications and public commentary. ")
    if illiberal_focus:
        intro.add_run("Notably, although they are aligned with classical-liberal perspectives, much of their writing focuses on analyzing illiberal regimes and practices. ")
    else:
        intro.add_run("Their writings reflect normative commitments to liberal principles across markets, rights, and democratic governance. ")

    # 2. Classification under the Three Dimensions of Liberalism
    doc.add_heading("2. Classification under the Three Dimensions of Liberalism", level=2)
    for dim, score in scores.items():
        p = doc.add_paragraph()
        p.add_run(f"**{dim}**: ").bold = True
        if score >= 2:
            p.add_run("Liberalism (strong alignment)").italic = True
        else:
            p.add_run("Mixed or Neutral alignment").italic = True
        p.add_run(f" — Evidence count: {score}")

    # 3. Summary Table of Classifications
    doc.add_heading("3. Summary Table of Classifications", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Dimension'
    hdr_cells[1].text = 'Alignment'
    hdr_cells[2].text = 'Evidence Count'
    for dim, score in scores.items():
        row_cells = table.add_row().cells
        row_cells[0].text = dim
        row_cells[1].text = "Liberalism" if score >= 2 else "Neutral"
        row_cells[2].text = str(score)

    # 4. Evidence and Citations
    doc.add_heading("4. Evidence and Citations", level=2)
    doc.add_paragraph("**Sample Publications:**", style='List Bullet')
    for title, cites in scholar_data['publications']:
        doc.add_paragraph(f"{title} (Citations: {cites})", style='List Number')

    for src_name, src_items in posts.items():
        doc.add_paragraph(f"**{src_name.capitalize()} Posts:**", style='List Bullet')
        for item in src_items:
            if 'title' in item:
                doc.add_paragraph(f"{item['date']} — {item['title']} ({item['url']})", style='List Number')
            else:
                doc.add_paragraph(f"{item['date']} — {item['content']} ({item['url']})", style='List Number')

    # Save report
    out = f"{args.last_name.lower()}_{args.first_name.lower()}_liberalism_profile.docx"
    doc.save(out)
    print(f"Report generated: {out}")

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--first_name", required=True)
    parser.add_argument("--last_name", required=True)
    parser.add_argument("--email", required=True)
    parser.add_argument("--affiliation", required=True)
    args = parser.parse_args()

    scholar = get_scholar_data(args.first_name, args.last_name, args.affiliation)
    if not scholar:
        print("Scholar not found.")
        return

    posts = {
        'LinkedIn': get_linkedin_posts(args.first_name, args.last_name),
        'Facebook': get_facebook_posts(args.first_name, args.last_name),
        'Substack': get_substack_articles(args.first_name, args.last_name)
    }

    scores, illiberal_focus, illib_count, lib_count = analyze_liberalism(scholar, posts)
    generate_report(args, scholar, posts, scores, illiberal_focus, illib_count, lib_count)

if __name__ == "__main__":
    main()

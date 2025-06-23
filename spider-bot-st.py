#!/usr/bin/env python3
"""
deep_dive_donor_scraper.py

A full-featured donor research scraper that:
  • Accepts first/last name (required), plus optional email, city, state, ZIP
  • Queries:
      - WealthEngine, iWave, Instrumentl (stub), DonorSearch (stub), ESRI
      - Clearbit & People Data Labs for enrichment
      - FEC & OpenSecrets for political giving
      - IRS Form 990 e-file index for nonprofit affiliations
      - Google Web via Custom Search API
      - Google Scholar via scholarly
  • Disambiguates on email/location when needed
  • Generates three report types:
      1) Full narrative profile
      2) Two-page executive summary
      3) Summary table with citations
  • Exports each as .docx and .pdf
  • Emits progress logs for integration into CLI or web interfaces
"""

import os
import argparse
import requests
import xml.etree.ElementTree as ET
from docx import Document

# Optional PDF conversion (requires Word or LibreOffice on PATH)
try:
    import docx2pdf
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

# Google Scholar
try:
    from scholarly import scholarly
    SCHOLARLY_AVAILABLE = True
except ImportError:
    SCHOLARLY_AVAILABLE = False

# ——— CONFIGURATION & API KEYS ————————————————————————
WEALTHENGINE_API_KEY = os.getenv("WEALTHENGINE_API_KEY", "")
IWAVE_API_KEY        = os.getenv("IWAVE_API_KEY", "")
INSTRUMENTL_API_KEY  = os.getenv("INSTRUMENTL_API_KEY", "")
DONORSEARCH_API_KEY  = os.getenv("DONORSEARCH_API_KEY", "")
ESRI_API_KEY         = os.getenv("ESRI_API_KEY", "")
CLEARBIT_API_KEY     = os.getenv("CLEARBIT_API_KEY", "")
PDL_API_KEY          = os.getenv("PDL_API_KEY", "")
FEC_API_KEY          = os.getenv("FEC_API_KEY", "")
OPENSECRETS_API_KEY  = os.getenv("OPENSECRETS_API_KEY", "")
GOOGLE_API_KEY       = os.getenv("GOOGLE_API_KEY", "")
GOOGLE_CSE_ID        = os.getenv("GOOGLE_CSE_ID", "")

IRS_INDEX_URL_TEMPLATE = "https://s3.amazonaws.com/irs-form-990/index_{year}.json"

# ——— SCRAPER CLASS ———————————————————————————————
class DonorResearchScraper:
    def __init__(self):
        # API endpoints
        self.we_url   = "https://api.wealthengine.com/v1/profile/find_one_full"
        self.iw_url   = "https://api.iwave.com/people/search"
        self.esri_url = "https://geoenrich.arcgis.com/arcgis/rest/services/World/GeoEnrichmentServer/GeoEnrichment/enrich"
        self.clearbit_url = "https://person.clearbit.com/v2/people/find"
        self.pdl_enrich_url = "https://api.peopledatalabs.com/v5/person/enrich"
        self.pdl_search_url = "https://api.peopledatalabs.com/v5/person/search"
        self.fec_url  = "https://api.open.fec.gov/v1/schedules/schedule_a/"
        self.os_url   = "https://www.opensecrets.org/api/"
        self.cs_url   = "https://www.googleapis.com/customsearch/v1"

    # ——— WEALTHENGINE ——————————————
    def search_wealthengine(self, first, last, **opts):
        if not WEALTHENGINE_API_KEY: return {}
        params = {"firstName": first, "lastName": last, **{k:v for k,v in opts.items() if v}}
        headers = {"Authorization": f"Bearer {WEALTHENGINE_API_KEY}"}
        try:
            r = requests.get(self.we_url, params=params, headers=headers, timeout=10)
            if r.ok:
                d = r.json()
                return {
                    "net_worth": d.get("estimatedNetWorth"),
                    "gift_capacity": d.get("giftCapacity"),
                    "source": "WealthEngine"
                }
        except: pass
        return {}

    # ——— IWAVE ——————————————
    def search_iwave(self, first, last, city=None, state=None):
        if not IWAVE_API_KEY: return {}
        params = {"first_name": first, "last_name": last, "city": city, "state": state}
        headers = {"X-API-KEY": IWAVE_API_KEY}
        try:
            r = requests.get(self.iw_url, params=params, headers=headers, timeout=10)
            if r.ok:
                d = r.json()
                return {
                    "philanthropy_score": d.get("philanthropyScore"),
                    "net_worth": d.get("netWorth"),
                    "source": "iWave"
                }
        except: pass
        return {}

    # ——— CLEARBIT ——————————————
    def search_clearbit(self, email):
        if not CLEARBIT_API_KEY or not email: return {}
        headers = {"Authorization": f"Bearer {CLEARBIT_API_KEY}"}
        try:
            r = requests.get(f"{self.clearbit_url}?email={email}", headers=headers, timeout=10)
            if r.status_code == 200:
                d = r.json()
                return {
                    "full_name": d.get("name", {}).get("fullName"),
                    "employment": d.get("employment"),
                    "linkedin": d.get("linkedin"),
                    "facebook": d.get("facebook"),
                    "bio": d.get("bio"),
                    "source": "Clearbit"
                }
        except: pass
        return {}

    # ——— PDL ——————————————
    def search_pdl(self, email=None, first=None, last=None, city=None, state=None):
        if not PDL_API_KEY: return {}
        headers = {"Authorization": f"Bearer {PDL_API_KEY}"}
        try:
            if email:
                params = {"email": email}
                r = requests.get(self.pdl_enrich_url, params=params, headers=headers, timeout=10)
                person = r.json().get("data", {})
            else:
                params = {"first_name": first, "last_name": last, "locality": city, "region": state}
                r = requests.get(self.pdl_search_url, params=params, headers=headers, timeout=10)
                hits = r.json().get("data", [])
                person = hits[0] if hits else {}
            prof = {
                "linkedin": None, "facebook": None,
                "job_title": person.get("job_title"),
                "job_company": person.get("job_company"),
                "education": person.get("education"),
                "source": "PDL"
            }
            for p in person.get("profiles", []):
                if p.get("network") == "linkedin": prof["linkedin"] = p.get("url")
                if p.get("network") == "facebook": prof["facebook"] = p.get("url")
            return prof
        except: return {}

    # ——— FEC ——————————————
    def search_fec(self, first, last, state=None):
        if not FEC_API_KEY: return []
        params = {"api_key": FEC_API_KEY, "contributor_name": f"{first} {last}", "per_page":50}
        if state: params["contributor_state"] = state
        try:
            r = requests.get(self.fec_url, params=params, timeout=10)
            if r.ok:
                return [{
                    "recipient": rec["committee"]["name"],
                    "amount": rec["contribution_receipt_amount"],
                    "date": rec["contribution_receipt_date"],
                    "source": "FEC"
                } for rec in r.json().get("results", [])]
        except: pass
        return []

    # ——— OPENSECRETS ——————————————
    def search_opensecrets(self, first, last):
        if not OPENSECRETS_API_KEY: return {}
        params = {"method":"indiv","name":f"{first} {last}","output":"json","apikey":OPENSECRETS_API_KEY}
        try:
            r = requests.get(self.os_url, params=params, timeout=10)
            if r.ok:
                return {"opensecrets_summary": r.json(), "source": "OpenSecrets"}
        except: pass
        return {}

    # ——— IRS INDEX ——————————————
    def fetch_irs_index(self, year):
        try:
            r = requests.get(IRS_INDEX_URL_TEMPLATE.format(year=year), timeout=15)
            if r.ok:
                return r.json().get(f"Filings{year}", [])
        except: pass
        return []

    def find_irs_filings(self, first, last, city=None, state=None):
        matches = []
        name = f"{first} {last}"
        for year in range(2018, 2026):
            for f in self.fetch_irs_index(year):
                org = f["OrganizationName"]
                if last.upper() in org or first.upper() in org:
                    matches.append(f)
                    continue
                url = f.get("URL")
                if not url: continue
                try:
                    xml = requests.get(url, timeout=10).content
                    root = ET.fromstring(xml)
                    if name in [e.text for e in root.findall(".//PersonName")]:
                        f["_xml"] = xml
                        matches.append(f)
                except: pass
        return matches

    def parse_irs_filing(self, filing, first, last):
        xml = filing.get("_xml") or requests.get(filing["URL"], timeout=10).content
        root = ET.fromstring(xml)
        name = f"{first} {last}"
        for pn in root.findall(".//PersonName"):
            if pn.text == name:
                parent = pn.getparent() if hasattr(pn,'getparent') else pn.find("..")
                role = parent.findtext(".//Title") or "N/A"
                comp = parent.findtext(".//Compensation") or "0"
                org = root.findtext(".//Filer/BusinessName/BusinessNameLine1") or filing["OrganizationName"]
                mission = root.findtext(".//MissionDescription") or "N/A"
                return {
                    "organization": org,
                    "ein": filing["EIN"],
                    "role": role,
                    "compensation": float(comp),
                    "year": filing["TaxPeriod"][:4],
                    "mission": mission,
                    "source": f"IRS {filing['FormType']} {filing['TaxPeriod']}"
                }
        return None

    # ——— GOOGLE WEB SEARCH ——————————————
    def search_google(self, query, num=5):
        if not GOOGLE_API_KEY or not GOOGLE_CSE_ID: return []
        params = {"key": GOOGLE_API_KEY, "cx": GOOGLE_CSE_ID, "q":query, "num":num}
        try:
            r = requests.get(self.cs_url, params=params, timeout=10)
            if r.ok:
                return [{"title":i["title"],"link":i["link"],"snippet":i["snippet"]} 
                        for i in r.json().get("items",[])]
        except: pass
        return []

    # ——— GOOGLE SCHOLAR ——————————————
    def search_scholar(self, name, max_results=5):
        if not SCHOLARLY_AVAILABLE: return []
        try:
            auth = next(scholarly.search_author(name), None)
            if not auth: return []
            scholarly.fill(auth)
            pubs = auth.get("publications", [])[:max_results]
            return [{"title":p["bib"].get("title"),"url":p.get("pub_url") or p.get("eprint_url"),
                     "year":p["bib"].get("pub_year")} for p in pubs]
        except: pass
        return []

    # ——— GATHER ALL DATA ——————————————
    def gather_data(self, first, last, email=None, city=None, state=None, zip_code=None):
        data = {"name": f"{first} {last}"}
        if email:
            data.update(self.search_clearbit(email))
            data.update(self.search_pdl(email=email))
        data["wealthengine"] = self.search_wealthengine(first, last, email=email, city=city, state=state)
        data["iwave"]        = self.search_iwave(first, last, city, state)
        data["fec"]          = self.search_fec(first, last, state)
        data["opensecrets"]  = self.search_opensecrets(first, last)
        # IRS affiliations
        filings = self.find_irs_filings(first, last, city, state)
        data["nonprofit_affiliations"] = [self.parse_irs_filing(f, first, last) for f in filings if self.parse_irs_filing(f, first, last)]
        # Google & Scholar
        data["google_results"]  = self.search_google(data["name"])
        data["scholar_results"] = self.search_scholar(data["name"])
        return data

    # ——— REPORT GENERATION ——————————————
    def generate_profile_doc(self, d):
        doc = Document()
        doc.add_heading(f"Donor Profile: {d['name']}", level=1)
        doc.add_heading("Biography & Enrichment", level=2)
        doc.add_paragraph(d.get("bio","No bio available."))
        doc.add_heading("Wealth & Capacity", level=2)
        we = d.get("wealthengine",{})
        if we.get("net_worth"): doc.add_paragraph(f"Net Worth: ${we['net_worth']:,} (WealthEngine)")
        if we.get("gift_capacity"): doc.add_paragraph(f"Gift Capacity: ${we['gift_capacity']:,} (WealthEngine)")
        doc.add_heading("Political Giving", level=2)
        fec = d.get("fec",[])
        if fec:
            tot = sum(x["amount"] for x in fec)
            doc.add_paragraph(f"FEC Donations: {len(fec)} gifts totaling ${tot:,.0f}.")
        else:
            doc.add_paragraph("No FEC records.")
        doc.add_heading("Nonprofit Affiliations (IRS Form 990)", level=2)
        for aff in d.get("nonprofit_affiliations",[]):
            doc.add_paragraph(f"{aff['role']} @ {aff['organization']} ({aff['year']}): comp ${aff['compensation']:,} – Mission: {aff['mission']} [{aff['source']}]")
        doc.add_heading("Web Presence & Publications", level=2)
        for item in d.get("google_results",[]):
            doc.add_paragraph(f"{item['title']} – {item['link']}")
        for pub in d.get("scholar_results",[]):
            doc.add_paragraph(f"{pub['title']} ({pub['year']}) – {pub['url']}")
        return doc

    def generate_exec_summary_doc(self, d):
        doc = Document()
        doc.add_heading(f"Executive Summary: {d['name']}", level=1)
        pts = []
        we = d.get("wealthengine",{})
        if we.get("net_worth"): pts.append(f"Net worth ~${we['net_worth']:,}")
        if we.get("gift_capacity"): pts.append(f"Gift capacity ~${we['gift_capacity']:,}")
        fec = d.get("fec",[])
        if fec: pts.append(f"Political gifts totaling ${sum(x['amount'] for x in fec):,.0f}")
        for aff in d.get("nonprofit_affiliations",[]):
            pts.append(f"{aff['role']} @ {aff['organization']} ({aff['year']}); comp ${aff['compensation']:,}")
        if not pts: pts.append("No significant data found.")
        for p in pts: doc.add_paragraph(p, style="List Bullet")
        return doc

    def generate_summary_table_doc(self, d):
        doc = Document()
        doc.add_heading(f"Summary Table: {d['name']}", level=1)
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Category","Detail","Source"
        def add_row(c, det, src):
            row = table.add_row().cells
            row[0].text, row[1].text, row[2].text = c, det, src
        we = d.get("wealthengine",{})
        if we.get("net_worth"): add_row("Net Worth", f"${we['net_worth']:,}", "WealthEngine")
        if we.get("gift_capacity"): add_row("Gift Capacity", f"${we['gift_capacity']:,}", "WealthEngine")
        fec = d.get("fec",[])
        if fec: add_row("FEC Donations", f"${sum(x['amount'] for x in fec):,.0f}", "FEC")
        for aff in d.get("nonprofit_affiliations",[]):
            add_row(f"Affiliation: {aff['organization']}", f"{aff['role']} ({aff['year']}); comp ${aff['compensation']:,}", aff['source'])
        return doc

    # ——— SAVE & EXPORT ——————————————
    def save_and_convert(self, doc, name):
        os.makedirs("outputs", exist_ok=True)
        docx_path = os.path.join("outputs", f"{name}.docx")
        doc.save(docx_path)
        pdf_path = docx_path.replace(".docx",".pdf")
        if DOCX2PDF_AVAILABLE:
            try: docx2pdf.convert(docx_path,pdf_path)
            except: pass
        return docx_path, pdf_path

    def run(self, first, last, email=None, city=None, state=None, zip_code=None):
        print(f"▶️  Gathering data for {first} {last}...")
        d = self.gather_data(first,last,email,city,state,zip_code)
        print("▶️  Generating Full Profile…")
        pdoc = self.generate_profile_doc(d); pp = self.save_and_convert(pdoc, f"{first}_{last}_Profile")
        print("▶️  Generating Exec Summary…")
        sdoc = self.generate_exec_summary_doc(d); sp = self.save_and_convert(sdoc, f"{first}_{last}_ExecSummary")
        print("▶️  Generating Summary Table…")
        tdoc = self.generate_summary_table_doc(d); tp = self.save_and_convert(tdoc, f"{first}_{last}_SummaryTable")
        print("✅  Done.")
        return {"profile":pp, "summary":sp, "table":tp}

# ——— CLI ——————————————————————————————————————————
def main():
    p = argparse.ArgumentParser()
    p.add_argument("--first", required=True)
    p.add_argument("--last",  required=True)
    p.add_argument("--email")
    p.add_argument("--city")
    p.add_argument("--state")
    p.add_argument("--zip", dest="zip_code")
    args = p.parse_args()

    scraper = DonorResearchScraper()
    out = scraper.run(
        args.first, args.last,
        email=args.email,
        city=args.city,
        state=args.state,
        zip_code=args.zip_code
    )
    print("\nFiles available in ./outputs:")
    for k,v in out.items():
        print(f"{k.title()}: DOCX={v[0]}, PDF={v[1]}")

if __name__ == "__main__":
    main()

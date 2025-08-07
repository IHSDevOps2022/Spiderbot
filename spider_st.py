#!/usr/bin/env python3
"""
enhanced_donor_scraper.py - AI Deep Research Version

Enhanced donor research scraper that uses AI to orchestrate deep research
and generates three focused reports:
1. Summary table with key prospect information
2. 2-page executive summary
3. Full narrative donor report

All reports are formatted specifically for IHS donor research.
"""

import os
import csv
import json
import argparse
import requests
import xml.etree.ElementTree as ET
from docx import Document
import time
from datetime import datetime

# Optional PDF conversion
try:
    import docx2pdf
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

# Google Scholar
try:
    from scholarly import scholarly
    SCHOLARLY_AVAILABLE = True
except ImportError:
    SCHOLARLY_AVAILABLE = False

# OpenAI for ChatGPT
try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    OPENAI_AVAILABLE = False

# ——— CONFIGURATION & API KEYS ————————————————————————
# 
# SET YOUR OPENAI API KEY HERE:
# Option 1: Set it directly in the code (not recommended for production)
# OPENAI_API_KEY = "sk-your-api-key-here"
#
# Option 2: Set it as an environment variable (recommended)
# In your terminal before running the script:
# export OPENAI_API_KEY="sk-your-api-key-here"
#
# Option 3: Create a .env file in the same directory with:
# OPENAI_API_KEY=sk-your-api-key-here

# OpenAI Model Selection (choose based on your access):
# - "gpt-4o" - Latest and most capable (if available)
# - "gpt-4-turbo" - Good balance of capability and cost
# - "gpt-4" - Original GPT-4 
# - "gpt-3.5-turbo" - Faster and cheaper, but less capable
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4-turbo")

# API Keys - Set these as environment variables or directly here
OPENAI_API_KEY       = ""

# Optional APIs - The AI will work around missing APIs
WEALTHENGINE_API_KEY = os.getenv("WEALTHENGINE_API_KEY", "")
IWAVE_API_KEY        = os.getenv("IWAVE_API_KEY", "")
CLEARBIT_API_KEY     = os.getenv("CLEARBIT_API_KEY", "")
PDL_API_KEY          = os.getenv("PDL_API_KEY", "")
FEC_API_KEY          = os.getenv("FEC_API_KEY", "")
OPENSECRETS_API_KEY  = os.getenv("OPENSECRETS_API_KEY", "")
GOOGLE_API_KEY       = os.getenv("GOOGLE_API_KEY", "")
GOOGLE_CSE_ID        = os.getenv("GOOGLE_CSE_ID", "")
FOLLOWTHEMONEY_KEY   = os.getenv("FOLLOWTHEMONEY_API_KEY", "")
GUIDESTAR_API_KEY    = os.getenv("GUIDESTAR_API_KEY", "")
FOUNDATION_CENTER_KEY = os.getenv("FOUNDATION_CENTER_API_KEY", "")
DEBOUNCE_API_KEY     = os.getenv("DEBOUNCE_API_KEY", "")
INTELIUS_API_KEY     = os.getenv("INTELIUS_API_KEY", "")
RELSCI_API_KEY       = os.getenv("RELSCI_API_KEY", "")
ROCKETREACH_API_KEY  = os.getenv("ROCKETREACH_API_KEY", "")
WEALTHX_API_KEY      = os.getenv("WEALTHX_API_KEY", "")
WINDFALL_API_KEY     = os.getenv("WINDFALL_API_KEY", "")
LEXISNEXIS_API_KEY   = os.getenv("LEXISNEXIS_API_KEY", "")
PACER_API_KEY        = os.getenv("PACER_API_KEY", "")
PITCHBOOK_API_KEY    = os.getenv("PITCHBOOK_API_KEY", "")

IRS_INDEX_URL_TEMPLATE = "https://s3.amazonaws.com/irs-form-990/index_{year}.json"

# ——— AI DEEP RESEARCH ORCHESTRATOR ————————————————————————————
class AIDeepResearchOrchestrator:
    """AI-powered orchestrator that intelligently searches across all databases"""
    
    def __init__(self, scraper_instance):
        self.scraper = scraper_instance
        if not OPENAI_API_KEY:
            raise ValueError("OpenAI API key required for AI orchestration. Set OPENAI_API_KEY environment variable.")
        self.client = openai.OpenAI(api_key=OPENAI_API_KEY)
        # Try to use the configured model
        self.model = OPENAI_MODEL
        
        # Test the model with a simple request
        try:
            test_response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": "test"}],
                max_tokens=10
            )
            print(f"✓ Using OpenAI model: {self.model}")
        except Exception as e:
            if "model" in str(e).lower():
                # Fallback to gpt-3.5-turbo if configured model isn't available
                print(f"  ⚠️  {self.model} not available, using gpt-3.5-turbo")
                self.model = "gpt-3.5-turbo"
            else:
                raise e
        
        # Map all available search methods
        self.available_searches = {
            "wealth_screening": {
                "wealthengine": {
                    "method": self.scraper.search_wealthengine,
                    "params": ["first", "last", "email", "city", "state"],
                    "requires_key": bool(WEALTHENGINE_API_KEY),
                    "description": "Net worth and gift capacity estimates"
                },
                "iwave": {
                    "method": self.scraper.search_iwave,
                    "params": ["first", "last", "city", "state"],
                    "requires_key": bool(IWAVE_API_KEY),
                    "description": "Philanthropy score and wealth indicators"
                },
                "wealthx": {
                    "method": self.scraper.search_wealthx,
                    "params": ["first", "last"],
                    "requires_key": bool(WEALTHX_API_KEY),
                    "description": "Ultra high net worth intelligence"
                },
                "windfall": {
                    "method": self.scraper.search_windfall,
                    "params": ["first", "last"],
                    "requires_key": bool(WINDFALL_API_KEY),
                    "description": "Wealth screening and estimates"
                }
            },
            "contact_enrichment": {
                "clearbit": {
                    "method": self.scraper.search_clearbit,
                    "params": ["email"],
                    "requires_key": bool(CLEARBIT_API_KEY),
                    "description": "Professional and social profiles from email"
                },
                "pdl": {
                    "method": self.scraper.search_pdl,
                    "params": ["email", "first", "last", "city", "state"],
                    "requires_key": bool(PDL_API_KEY),
                    "description": "Comprehensive people data"
                },
                "rocketreach": {
                    "method": self.scraper.search_rocketreach,
                    "params": ["first", "last", "company"],
                    "requires_key": bool(ROCKETREACH_API_KEY),
                    "description": "Email and phone discovery"
                },
                "debounce": {
                    "method": self.scraper.validate_email_debounce,
                    "params": ["email"],
                    "requires_key": bool(DEBOUNCE_API_KEY),
                    "description": "Email validation"
                }
            },
            "political_giving": {
                "fec": {
                    "method": self.scraper.search_fec,
                    "params": ["first", "last", "state", "city", "zip_code"],
                    "requires_key": bool(FEC_API_KEY),
                    "description": "Federal campaign contributions"
                },
                "opensecrets": {
                    "method": self.scraper.search_opensecrets,
                    "params": ["first", "last"],
                    "requires_key": bool(OPENSECRETS_API_KEY),
                    "description": "Political giving summaries and analysis"
                },
                "followthemoney": {
                    "method": self.scraper.search_followthemoney,
                    "params": ["first", "last", "state"],
                    "requires_key": bool(FOLLOWTHEMONEY_KEY),
                    "description": "State and local campaign finance"
                }
            },
            "nonprofit_philanthropy": {
                "irs_990": {
                    "method": self.scraper.find_irs_filings,
                    "params": ["first", "last", "city", "state"],
                    "requires_key": True,  # Always available
                    "description": "IRS Form 990 filings and nonprofit affiliations"
                },
                "guidestar": {
                    "method": self.scraper.search_guidestar,
                    "params": ["first", "last"],
                    "requires_key": bool(GUIDESTAR_API_KEY),
                    "description": "Comprehensive nonprofit data"
                },
                "foundation_center": {
                    "method": self.scraper.search_foundation_center,
                    "params": ["first", "last"],
                    "requires_key": bool(FOUNDATION_CENTER_KEY),
                    "description": "Foundation grants and giving"
                },
                "propublica_nonprofit": {
                    "method": self.scraper.search_propublica_nonprofit,
                    "params": ["first", "last", "city", "state"],
                    "requires_key": True,  # Always available
                    "description": "ProPublica nonprofit explorer"
                }
            },
            "business_wealth": {
                "sec_edgar": {
                    "method": self.scraper.search_sec_edgar,
                    "params": ["first", "last"],
                    "requires_key": True,  # Always available
                    "description": "SEC filings, insider trading, company ownership"
                },
                "pitchbook": {
                    "method": self.scraper.search_pitchbook,
                    "params": ["first", "last"],
                    "requires_key": bool(PITCHBOOK_API_KEY),
                    "description": "VC, PE, and M&A activity"
                },
                "relsci": {
                    "method": self.scraper.search_relsci,
                    "params": ["first", "last"],
                    "requires_key": bool(RELSCI_API_KEY),
                    "description": "Professional relationships and board positions"
                }
            },
            "public_records": {
                "lexisnexis": {
                    "method": self.scraper.search_lexisnexis,
                    "params": ["first", "last"],
                    "requires_key": bool(LEXISNEXIS_API_KEY),
                    "description": "News, legal records, public filings"
                },
                "pacer": {
                    "method": self.scraper.search_pacer,
                    "params": ["first", "last"],
                    "requires_key": bool(PACER_API_KEY),
                    "description": "Federal court records"
                },
                "intelius": {
                    "method": self.scraper.search_intelius,
                    "params": ["first", "last", "city", "state"],
                    "requires_key": bool(INTELIUS_API_KEY),
                    "description": "Background checks and public records"
                }
            },
            "web_research": {
                "google": {
                    "method": self.scraper.search_google,
                    "params": ["query", "num"],
                    "requires_key": bool(GOOGLE_API_KEY),
                    "description": "General web search"
                },
                "scholar": {
                    "method": self.scraper.search_scholar,
                    "params": ["name", "max_results"],
                    "requires_key": SCHOLARLY_AVAILABLE,
                    "description": "Academic publications"
                },
                "inside_philanthropy": {
                    "method": self.scraper.search_inside_philanthropy,
                    "params": ["first", "last", "city", "state"],
                    "requires_key": bool(GOOGLE_API_KEY),
                    "description": "Inside Philanthropy articles"
                },
                "chronicle_philanthropy": {
                    "method": self.scraper.search_chronicle_philanthropy,
                    "params": ["first", "last", "city", "state"],
                    "requires_key": bool(GOOGLE_API_KEY),
                    "description": "Chronicle of Philanthropy coverage"
                }
            }
        }

    def conduct_deep_research(self, first, last, email=None, city=None, state=None, zip_code=None):
        """Main orchestration method - AI decides what to search and when"""
        
        context = {
            "person": {
                "first": first,
                "last": last,
                "email": email,
                "city": city,
                "state": state,
                "zip_code": zip_code
            },
            "findings": {},
            "search_history": [],
            "confidence_scores": {},
            "research_strategy": None
        }
        
        # Phase 1: Initial Assessment
        print("  🤖 AI analyzing available data sources...")
        context = self.ai_assess_available_sources(context)
        
        # Phase 2: Create Research Strategy
        print("  🤖 AI creating research strategy...")
        context = self.ai_create_research_strategy(context)
        
        # Phase 3: Execute Iterative Research
        print("  🤖 Executing AI-guided deep research...")
        max_iterations = 5
        for i in range(max_iterations):
            print(f"    • Research iteration {i+1}/{max_iterations}")
            
            # AI decides next searches
            next_searches = self.ai_decide_next_searches(context)
            
            if not next_searches:
                print("    ✓ AI determined research is complete")
                break
            
            # Execute searches
            for search in next_searches:
                result = self.execute_search(search, context)
                if result and result.get('data'):
                    context['findings'][search['name']] = result
                    context['search_history'].append(search)
            
            # AI analyzes findings and adjusts strategy
            context = self.ai_analyze_iteration_results(context)
            
            # Brief pause between iterations
            time.sleep(1)
        
        # Phase 4: Deep Analysis
        print("  🤖 AI performing deep analysis...")
        context = self.ai_deep_analysis(context)
        
        # Phase 5: Generate Final Report
        print("  🤖 AI generating comprehensive report...")
        final_report = self.ai_generate_final_report(context)
        
        # Phase 6: Generate Executive Summary
        print("  🤖 AI creating executive summary...")
        exec_summary = self.ai_generate_executive_summary(context)
        final_report['executive_summary'] = exec_summary
        
        # Phase 7: Generate Summary Table
        print("  🤖 AI creating summary table...")
        summary_table = self.ai_generate_summary_table(context)
        final_report['summary_table'] = summary_table
        
        return final_report

    def ai_assess_available_sources(self, context):
        """AI assesses which data sources are available and relevant"""
        
        # Build available sources summary
        available = []
        unavailable = []
        
        for category, sources in self.available_searches.items():
            for name, config in sources.items():
                if config['requires_key']:
                    available.append({
                        "name": name,
                        "category": category,
                        "description": config['description']
                    })
                else:
                    unavailable.append({
                        "name": name,
                        "category": category,
                        "description": config['description']
                    })
        
        prompt = f"""
        Assess available data sources for researching this person:
        
        Target: {context['person']['first']} {context['person']['last']}
        Location: {context['person'].get('city', 'Unknown')}, {context['person'].get('state', 'Unknown')}
        Email: {context['person'].get('email', 'Not provided')}
        
        Available data sources:
        {json.dumps(available, indent=2)}
        
        Unavailable sources (no API key):
        {json.dumps(unavailable, indent=2)}
        
        Provide:
        1. Which sources are most critical for this research
        2. What information gaps exist due to unavailable sources
        3. Alternative strategies to compensate for missing sources
        4. Priority ranking of available sources
        """
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )
            
            context['source_assessment'] = response.choices[0].message.content
        except Exception as e:
            print(f"    ✗ AI assessment error: {str(e)}")
            context['source_assessment'] = "Error in AI assessment"
        
        return context

    def ai_create_research_strategy(self, context):
        """AI creates a comprehensive research strategy"""
        
        prompt = f"""
        Create a detailed research strategy for this donor prospect:
        
        Person: {json.dumps(context['person'], indent=2)}
        Available sources assessment: {context.get('source_assessment', 'Not available')}
        
        Design a research strategy that:
        1. Identifies the person with high confidence
        2. Determines wealth and capacity
        3. Uncovers philanthropic interests and history
        4. Maps relationships and networks
        5. Identifies cultivation opportunities
        6. Assesses ideological alignment with classical liberal principles
        7. Evaluates potential for supporting academic freedom initiatives
        
        Return a JSON object with this structure:
        {{
            "phases": [
                {{
                    "name": "Identity Verification",
                    "searches": ["clearbit", "pdl", "google"],
                    "goals": ["Confirm identity", "Find contact info"],
                    "decision_logic": "If email found, proceed to enrichment"
                }}
            ],
            "priority_sources": ["wealthengine", "fec", "irs_990"],
            "fallback_strategies": ["Use Google if paid APIs unavailable"]
        }}
        
        Return ONLY valid JSON, no markdown formatting or explanation.
        """
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.4
            )
            
            # Parse the response as JSON
            response_text = response.choices[0].message.content
            # Clean up the response if needed
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0]
            context['research_strategy'] = json.loads(response_text.strip())
        except Exception as e:
            print(f"    ✗ AI strategy error: {str(e)}")
            # Fallback strategy
            context['research_strategy'] = {
                "phases": [
                    {
                        "name": "Basic Search",
                        "searches": ["google", "irs_990", "fec"],
                        "goals": ["Initial data gathering"]
                    }
                ]
            }
        
        return context

    def ai_decide_next_searches(self, context):
        """AI decides which searches to run next based on current findings"""
        
        # Get unused searches
        used_searches = {s['name'] for s in context['search_history']}
        available_unused = []
        
        for category, sources in self.available_searches.items():
            for name, config in sources.items():
                if name not in used_searches and config['requires_key']:
                    available_unused.append({
                        "category": category,
                        "name": name,
                        "description": config['description'],
                        "params": config['params']
                    })
        
        if not available_unused:
            return []
        
        prompt = f"""
        Based on current research progress, decide next searches to execute.
        
        Person: {json.dumps(context['person'], indent=2)}
        Current findings summary: {len(context['findings'])} sources searched
        Research strategy: {json.dumps(context.get('research_strategy', {}), indent=2)}
        
        Available searches not yet used:
        {json.dumps(available_unused, indent=2)}
        
        Return a JSON object with a "searches" array containing up to 3 next searches:
        {{
            "searches": [
                {{
                    "category": "wealth_screening",
                    "name": "wealthengine",
                    "reason": "Need wealth estimate",
                    "params": {{"first": "John", "last": "Doe"}}
                }}
            ]
        }}
        
        Return empty array if research is complete.
        Return ONLY valid JSON, no markdown formatting or explanation.
        """
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )
            
            # Parse the response as JSON
            response_text = response.choices[0].message.content
            # Clean up the response if needed
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0]
            result = json.loads(response_text.strip())
            return result.get('searches', [])
        except Exception as e:
            print(f"    ✗ AI decision error: {str(e)}")
            # Fallback to basic searches
            if len(context['search_history']) < 3:
                return [
                    {
                        "category": "web_research",
                        "name": "google",
                        "params": {"query": f'"{context["person"]["first"]} {context["person"]["last"]}"', "num": 5}
                    }
                ]
            return []

    def execute_search(self, search_spec, context):
        """Execute a specific search based on AI instructions"""
        
        category = search_spec.get('category')
        name = search_spec.get('name')
        params = search_spec.get('params', {})
        
        if not category or not name:
            return None
            
        if category not in self.available_searches:
            return None
            
        if name not in self.available_searches[category]:
            return None
            
        search_config = self.available_searches[category][name]
        method = search_config['method']
        
        # Map parameters from context if needed
        call_params = {}
        for param_name in search_config['params']:
            if param_name in params:
                call_params[param_name] = params[param_name]
            elif param_name in context['person']:
                call_params[param_name] = context['person'][param_name]
        
        try:
            print(f"      → Searching {name}...")
            result = method(**call_params)
            return {
                "data": result,
                "source": name,
                "timestamp": datetime.now().isoformat(),
                "params_used": call_params
            }
        except Exception as e:
            print(f"      ✗ Error in {name}: {str(e)}")
            return {
                "error": str(e),
                "source": name,
                "timestamp": datetime.now().isoformat()
            }

    def ai_analyze_iteration_results(self, context):
        """AI analyzes results from latest searches"""
        
        if not context['findings']:
            return context
        
        # Get summary of findings
        findings_summary = {}
        for name, result in context['findings'].items():
            if result.get('data'):
                findings_summary[name] = {
                    "has_data": True,
                    "data_type": type(result['data']).__name__,
                    "data_size": len(str(result['data']))
                }
        
        prompt = f"""
        Analyze the research progress and findings.
        
        Person: {context['person']['first']} {context['person']['last']}
        Total searches completed: {len(context['search_history'])}
        Findings summary: {json.dumps(findings_summary, indent=2)}
        
        Return a JSON object with:
        {{
            "key_insights": ["insight1", "insight2"],
            "confidence_scores": {{
                "identity": 85,
                "wealth": 60,
                "interests": 70
            }},
            "missing_info": ["what we still need"],
            "continue_research": true
        }}
        
        Return ONLY valid JSON, no markdown formatting or explanation.
        """
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )
            
            # Parse the response as JSON
            response_text = response.choices[0].message.content
            # Clean up the response if needed
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0]
            analysis = json.loads(response_text.strip())
            context['iteration_analysis'] = analysis
            context['confidence_scores'] = analysis.get('confidence_scores', {})
        except Exception as e:
            print(f"    ✗ AI analysis error: {str(e)}")
            context['confidence_scores'] = {"identity": 50, "wealth": 50, "interests": 50}
        
        return context

    def ai_deep_analysis(self, context):
        """Comprehensive analysis of all findings focused on IHS donor research"""
        
        # Prepare findings summary
        findings_text = []
        for name, result in context['findings'].items():
            if result.get('data'):
                findings_text.append(f"\n{name.upper()}:")
                findings_text.append(json.dumps(result['data'], indent=2)[:1000])
        
        prompt = f"""
        Perform deep donor research analysis for {context['person']['first']} {context['person']['last']}.
        
        Location: {context['person'].get('city', 'Unknown')}, {context['person'].get('state', 'Unknown')}
        Email: {context['person'].get('email', 'Not provided')}
        
        Research Findings:
        {''.join(findings_text)}
        
        Provide comprehensive donor intelligence analysis for the Institute for Humane Studies (IHS):
        
        1. **Identity Verification**
           - Confirm full name, age, location
           - Professional identity and current positions
           - Family members relevant to giving decisions
           - Confidence level: [X]% based on [specific evidence]
        
        2. **Wealth Assessment**
           - Net worth estimate: $[X] million (range: $[Y]-[Z] million)
           - Wealth sources: [business sale, salary, investments, inheritance]
           - Recent liquidity events: [IPO, acquisition, sale dates and amounts]
           - Real estate: [properties with estimated values]
           - Business holdings: [current stakes and values]
        
        3. **Annual Giving Capacity**
           - Conservative (1% of net worth): $[amount]
           - Moderate (5% of net worth): $[amount]  
           - Aggressive (10% of net worth): $[amount]
           - Largest known gift: $[amount] to [recipient]
        
        4. **Philanthropic Profile**
           - Total political giving: $[amount] ([years])
           - Top recipients: [list with amounts]
           - Nonprofit board positions: [organizations]
           - Foundation affiliations: [names and roles]
           - Giving interests: [education, free speech, economics, etc.]
           - Named gifts: [facilities, programs, scholarships]
        
        5. **Ideological Indicators**
           - Political affiliation: [R/D/I] based on [giving patterns]
           - Think tank involvement: [organizations]
           - Academic freedom activities: [specific examples]
           - Published writings: [topics and outlets]
           - Public statements on higher ed: [quotes with sources]
           - Free market orientation: [High/Medium/Low] based on [evidence]
        
        6. **IHS Alignment Score**
           - Mission alignment: [X]/10
           - Classical liberal principles: [supports/neutral/opposes]
           - Academic freedom commitment: [evidence]
           - Student development interest: [evidence]
           - Free market support: [evidence]
        
        7. **Cultivation Intelligence**
           - Decision timeline: [quick/deliberative]
           - Recognition preferences: [anonymous/public]
           - Giving vehicles: [cash/stock/foundation/DAF]
           - Influence points: [who influences their giving]
           - Competing nonprofits: [similar organizations they support]
        
        8. **Recommended IHS Approach**
           - Initial contact: [specific person who should reach out]
           - Programs to highlight: [specific IHS initiatives]
           - Events to invite to: [specific IHS gatherings]
           - Ask amount: $[specific range] for [specific purpose]
           - Timeline: [cultivation period needed]
        
        Be specific with names, amounts, dates, and sources throughout.
        Focus on actionable intelligence for IHS major gift fundraising.
        """
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=3000
            )
            
            context['deep_analysis'] = response.choices[0].message.content
        except Exception as e:
            print(f"    ✗ AI deep analysis error: {str(e)}")
            context['deep_analysis'] = "Error in deep analysis"
        
        return context

    def ai_generate_final_report(self, context):
        """Generate comprehensive final report in IHS donor research style"""
        
        prompt = f"""
        Create a comprehensive IHS donor research report for {context['person']['first']} {context['person']['last']}.
        
        Deep Analysis Results:
        {context.get('deep_analysis', 'Not available')}
        
        Raw Findings Data:
        {json.dumps(context.get('findings', {}), indent=2)[:3000]}
        
        Generate a FULL NARRATIVE DONOR REPORT following this EXACT structure:
        
        **Full Narrative Donor Report**
        
        **Background & Career**
        [Write 3-4 detailed paragraphs covering:]
        - Educational background with specific schools, degrees, and graduation years
        - Career progression with company names, positions, and timeframes
        - Current professional status and activities
        - Family background if relevant to wealth/philanthropy
        - Include inline citations like [[source.com]{{.underline}}](URL#:~:text=specific,text)
        
        **Financial Capacity & Wealth Indicators**
        [Write 3-4 detailed paragraphs covering:]
        - Specific wealth sources (business sales, inheritance, salary, investments)
        - Real estate holdings with addresses and values if found
        - Business ownership stakes and exits
        - Estimated net worth ranges with supporting evidence
        - Liquidity events and their approximate values
        - Investment patterns and portfolios
        - Include specific dollar amounts when available
        
        **Philanthropic Interests & Track Record**
        [Write 3-4 detailed paragraphs covering:]
        - Major gifts with specific amounts and recipients
        - Foundation affiliations and board positions
        - Causes supported with examples
        - Giving patterns over time
        - Named gifts or facilities
        - Volunteer activities and time commitments
        - Collaborative giving with family/peers
        
        **Ideological Profile and Network Affiliations**
        [Write 3-4 detailed paragraphs covering:]
        - Political giving history with specific candidates/PACs and amounts
        - Think tank or policy organization involvement
        - Academic freedom or free speech activities
        - Board memberships and their ideological leanings
        - Public statements or writings on political/social issues
        - Network connections to other major donors
        - Assessment of alignment with classical liberal principles
        
        **Strategic Considerations for IHS Engagement**
        [Write 3-4 detailed paragraphs covering:]
        - Specific IHS programs that would resonate
        - Best approach for initial contact (who should make introduction)
        - Family dynamics and decision-making process
        - Recognition preferences
        - Potential ask amounts based on capacity and history
        - Timing considerations
        - Specific engagement event ideas
        - How to frame IHS mission to align with their interests
        
        CRITICAL INSTRUCTIONS:
        1. Every factual claim must have an inline citation with the exact format: [[source]{{.underline}}](URL)
        2. Include specific names, dates, dollar amounts, and percentages
        3. Write in narrative style with full paragraphs, not bullet points
        4. Focus on information relevant to IHS's mission of advancing classical liberal ideas in academia
        5. Make strategic recommendations specific and actionable
        6. Assess ideological alignment honestly but diplomatically
        """
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=4000
            )
            
            final_report_text = response.choices[0].message.content
        except Exception as e:
            print(f"    ✗ AI report generation error: {str(e)}")
            final_report_text = "Error generating final report"
        
        return {
            "name": f"{context['person']['first']} {context['person']['last']}",
            "search_context": context['person'],
            "confidence_scores": context.get('confidence_scores', {}),
            "findings": context['findings'],
            "deep_analysis": context.get('deep_analysis', ''),
            "final_report": final_report_text,
            "search_history": context['search_history'],
            "total_searches": len(context['search_history']),
            "timestamp": datetime.now().isoformat()
        }

    def ai_generate_executive_summary(self, context):
        """Generate 2-page executive summary in IHS style"""
        
        prompt = f"""
        Create a 2-page executive summary for IHS donor prospect: {context['person']['first']} {context['person']['last']}
        Location: {context['person'].get('city', '')}, {context['person'].get('state', '')}
        
        Based on research findings:
        {context.get('deep_analysis', '')[:2000]}
        
        Follow this EXACT format:
        
        **{context['person']['first']} {context['person']['last']} -- {context['person'].get('city', 'Location Unknown')}, {context['person'].get('state', '')}**
        
        **Executive Summary**
        
        **Biographical Overview:** [One paragraph with education, career, family, current status. Include inline citations [[source]{{.underline}}](URL)]
        
        **Career Trajectory & Financial Capacity:** [One paragraph detailing career progression, business ventures, wealth sources, and estimated capacity. Include specific companies, positions, and dollar ranges with citations]
        
        **Philanthropic Interests:** [One paragraph covering major gifts, causes supported, giving vehicles, and patterns. Include specific organizations and amounts with citations]
        
        **Ideological Alignment & Networks:** [One paragraph on political/ideological leanings, relevant board positions, think tank involvement, and network connections. Focus on alignment with classical liberal principles]
        
        **Probability of Giving -- Matrix Assessment:**
        - *Financial Capacity:* **[High/Moderate/Low].** [Brief explanation with evidence]
        - *Philanthropic Inclination:* **[High/Moderate/Low].** [Evidence of past giving]
        - *Ideological Alignment with IHS:* **[High/Moderate/Low].** [Specific alignment points]
        - *Existing Relationship with IHS:* **[High/Moderate/Low].** [Any connections]
        - **Overall Giving Probability:** **[High/Moderate/Low].** [Summary assessment]
        
        **Recommended Engagement Steps for IHS:**
        - **[Action Title]:** [Specific step with details on who, what, when, how]
        - **[Action Title]:** [Specific step focused on IHS programs that align with their interests]
        - **[Action Title]:** [Specific step for cultivation events or meetings]
        - **[Action Title]:** [Specific step for ask strategy and amount]
        
        Use inline citations throughout: [[source]{{.underline}}](URL)
        Be specific with names, titles, amounts, and dates.
        Focus on actionable intelligence for IHS fundraising.
        """
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=2000
            )
            
            return response.choices[0].message.content
        except Exception as e:
            print(f"    ✗ AI executive summary error: {str(e)}")
            return "Error generating executive summary"

    def ai_generate_summary_table(self, context):
        """Generate summary table in IHS style"""
        
        prompt = f"""
        Create a donor prospect summary table for {context['person']['first']} {context['person']['last']}.
        
        Research findings:
        {json.dumps(context.get('findings', {}), indent=2)[:2000]}
        {context.get('deep_analysis', '')[:1000]}
        
        Generate a summary table with EXACTLY this format:
        
        **Summary Table -- {context['person']['first']} {context['person']['last']}**
        
        Create a table with these exact rows (use the labels exactly as shown):
        
        **Profile Aspect** | **Details** | **Sources**
        ---|---|---
        **Education** | [Specific schools, degrees, years, honors] | [Source citations with [[site]{{.underline}}](URL)]
        **Career** | [Companies, positions, timeline, current role] | [Citations]
        **Family & Wealth Origin** | [Spouse, family business, inheritance, major liquidity events] | [Citations]
        **Notable Philanthropy** | [Major gifts, foundations, causes, amounts] | [Citations]
        **Ideological Involvement** | [Political giving, think tanks, advocacy groups, free speech initiatives] | [Citations]
        **Network Affiliations** | [Boards, clubs, donor networks, co-donors] | [Citations]
        **Capacity for Giving** | **[High/Moderate/Low].** [Specific evidence, net worth range, liquidity] | [Citations]
        **Alignment with IHS** | **[Strong/Moderate/Weak].** [Specific points of alignment with classical liberal mission] | [Citations]
        **Engagement Interests** | [What motivates them, recognition preferences, giving style] | [Citations or "Inferred from..."]
        **Potential IHS Strategy** | [Specific approaches, programs to highlight, ask strategy] | *Strategy based on profile analysis above.*
        
        Each detail should be specific with names, amounts, dates.
        Every factual claim needs a citation in the format: [[source]{{.underline}}](URL#:~:text=relevant,excerpt)
        Focus on information relevant to IHS donor cultivation.
        """
        
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=2000
            )
            
            return response.choices[0].message.content
        except Exception as e:
            print(f"    ✗ AI summary table error: {str(e)}")
            return "Error generating summary table"

# ——— ENHANCED SCRAPER CLASS ———————————————————————————————
class EnhancedDonorScraper:
    def __init__(self):
        # API endpoints
        self.we_url   = "https://api.wealthengine.com/v1/profile/find_one_full"
        self.iw_url   = "https://api.iwave.com/people/search"
        self.esri_url = "https://geoenrich.arcgis.com/arcgis/rest/services/World/GeoEnrichmentServer/GeoEnrichment/enrich"
        self.clearbit_url = "https://person.clearbit.com/v2/people/find"
        self.pdl_enrich_url = "https://api.peopledatalabs.com/v5/person/enrich"
        self.pdl_search_url = "https://api.peopledatalabs.com/v5/person/search"
        self.fec_url  = "https://api.open.fec.gov/v1/schedules/schedule_a/"
        self.os_url   = "https://www.opensecrets.org/api/"
        self.cs_url   = "https://www.googleapis.com/customsearch/v1"
        self.propublica_527_url = "https://projects.propublica.org/nonprofits/api/v2/search.json"
        self.sec_edgar_url = "https://www.sec.gov/cgi-bin/browse-edgar"
        self.rocketreach_url = "https://api.rocketreach.co/v2/api/lookupProfile"
        self.debounce_url = "https://api.debounce.io/v1/"
        
        # Initialize AI orchestrator if available
        self.ai_orchestrator = None
        if OPENAI_API_KEY and OPENAI_AVAILABLE:
            try:
                self.ai_orchestrator = AIDeepResearchOrchestrator(self)
                print("✓ AI Deep Research Engine initialized")
            except Exception as e:
                print(f"✗ AI initialization failed: {e}")
                print("Please check your OpenAI API key")

    # ——— ALL SEARCH METHODS (kept for AI to use) ——————————————
    def search_wealthengine(self, first, last, **opts):
        if not WEALTHENGINE_API_KEY: return {}
        params = {"firstName": first, "lastName": last, **{k:v for k,v in opts.items() if v}}
        headers = {"Authorization": f"Bearer {WEALTHENGINE_API_KEY}"}
        try:
            r = requests.get(self.we_url, params=params, headers=headers, timeout=10)
            if r.ok:
                d = r.json()
                return {
                    "net_worth": d.get("estimatedNetWorth"),
                    "gift_capacity": d.get("giftCapacity"),
                    "source": "WealthEngine"
                }
        except: pass
        return {}

    def search_iwave(self, first, last, city=None, state=None):
        if not IWAVE_API_KEY: return {}
        params = {"first_name": first, "last_name": last, "city": city, "state": state}
        headers = {"X-API-KEY": IWAVE_API_KEY}
        try:
            r = requests.get(self.iw_url, params=params, headers=headers, timeout=10)
            if r.ok:
                d = r.json()
                return {
                    "philanthropy_score": d.get("philanthropyScore"),
                    "net_worth": d.get("netWorth"),
                    "source": "iWave"
                }
        except: pass
        return {}

    def search_clearbit(self, email):
        if not CLEARBIT_API_KEY or not email: return {}
        headers = {"Authorization": f"Bearer {CLEARBIT_API_KEY}"}
        try:
            r = requests.get(f"{self.clearbit_url}?email={email}", headers=headers, timeout=10)
            if r.status_code == 200:
                d = r.json()
                return {
                    "full_name": d.get("name", {}).get("fullName"),
                    "employment": d.get("employment"),
                    "linkedin": d.get("linkedin"),
                    "facebook": d.get("facebook"),
                    "bio": d.get("bio"),
                    "source": "Clearbit"
                }
        except: pass
        return {}

    def search_pdl(self, email=None, first=None, last=None, city=None, state=None):
        if not PDL_API_KEY: return {}
        headers = {"Authorization": f"Bearer {PDL_API_KEY}"}
        try:
            if email:
                params = {"email": email}
                r = requests.get(self.pdl_enrich_url, params=params, headers=headers, timeout=10)
                person = r.json().get("data", {})
            else:
                params = {"sql": f'SELECT * FROM person WHERE full_name="{first} {last}"'}
                if city: params["sql"] += f' AND location_locality="{city}"'
                if state: params["sql"] += f' AND location_region="{state}"'
                r = requests.get(self.pdl_search_url, params=params, headers=headers, timeout=10)
                hits = r.json().get("data", [])
                person = hits[0] if hits else {}
            prof = {
                "linkedin": None, "facebook": None,
                "job_title": person.get("job_title"),
                "job_company": person.get("job_company"),
                "education": person.get("education"),
                "source": "PDL"
            }
            for p in person.get("profiles", []):
                if p.get("network") == "linkedin": prof["linkedin"] = p.get("url")
                if p.get("network") == "facebook": prof["facebook"] = p.get("url")
            return prof
        except: return {}

    def search_fec(self, first, last, state=None, city=None, zip_code=None):
        if not FEC_API_KEY: return []
        params = {
            "api_key": FEC_API_KEY, 
            "contributor_name": f'"{first} {last}"',
            "per_page": 50
        }
        
        if state: 
            params["contributor_state"] = state
        if zip_code: 
            params["contributor_zip"] = zip_code
            
        try:
            r = requests.get(self.fec_url, params=params, timeout=10)
            if r.ok:
                results = []
                for rec in r.json().get("results", []):
                    contrib_city = rec.get("contributor_city", "").lower()
                    contrib_state = rec.get("contributor_state", "").upper()
                    
                    if city and contrib_city:
                        if city.lower() not in contrib_city:
                            continue
                    
                    if state and contrib_state:
                        if state.upper() != contrib_state:
                            continue
                    
                    results.append({
                        "recipient": rec["committee"]["name"],
                        "amount": rec["contribution_receipt_amount"],
                        "date": rec["contribution_receipt_date"],
                        "contributor_city": rec.get("contributor_city"),
                        "contributor_state": rec.get("contributor_state"),
                        "contributor_zip": rec.get("contributor_zip"),
                        "source": "FEC"
                    })
                return results
        except: pass
        return []

    def search_opensecrets(self, first, last):
        if not OPENSECRETS_API_KEY: return {}
        params = {"method":"indiv","name":f'"{first} {last}"',"output":"json","apikey":OPENSECRETS_API_KEY}
        try:
            r = requests.get(self.os_url, params=params, timeout=10)
            if r.ok:
                return {"opensecrets_summary": r.json(), "source": "OpenSecrets"}
        except: pass
        return {}

    def search_followthemoney(self, first, last, state=None):
        if not FOLLOWTHEMONEY_KEY: return []
        try:
            url = "https://api.followthemoney.org/search"
            params = {
                "query": f'"{first} {last}"',
                "type": "contributors",
                "apikey": FOLLOWTHEMONEY_KEY
            }
            if state: params["state"] = state
            r = requests.get(url, params=params, timeout=10)
            if r.ok:
                return r.json().get("results", [])
        except: pass
        return []

    def search_irs_527(self, first, last):
        try:
            params = {"q": f'"{first} {last}"', "format": "json"}
            r = requests.get(self.propublica_527_url, params=params, timeout=10)
            if r.ok:
                return r.json().get("organizations", [])
        except: pass
        return []

    def search_guidestar(self, first, last):
        if not GUIDESTAR_API_KEY: return {}
        return {"status": "Not implemented", "source": "Guidestar"}

    def search_foundation_center(self, first, last):
        if not FOUNDATION_CENTER_KEY: return {}
        return {"status": "Not implemented", "source": "Foundation Center"}

    def search_propublica_nonprofit(self, first, last, city=None, state=None):
        try:
            url = "https://projects.propublica.org/nonprofits/api/v2/search.json"
            query = f'"{first} {last}"'
            if city:
                query += f' "{city}"'
            if state:
                query += f' "{state}"'
            
            params = {"q": query}
            r = requests.get(url, params=params, timeout=10)
            if r.ok:
                return r.json().get("organizations", [])
        except: pass
        return []

    def validate_email_debounce(self, email):
        if not DEBOUNCE_API_KEY or not email: return {}
        try:
            params = {"api": DEBOUNCE_API_KEY, "email": email}
            r = requests.get(f"{self.debounce_url}validate", params=params, timeout=10)
            if r.ok:
                return r.json()
        except: pass
        return {}

    def search_intelius(self, first, last, city=None, state=None):
        if not INTELIUS_API_KEY: return {}
        return {"status": "Not implemented", "source": "Intelius"}

    def search_relsci(self, first, last):
        if not RELSCI_API_KEY: return {}
        return {"status": "Not implemented", "source": "RelSci"}

    def search_rocketreach(self, first, last, company=None):
        if not ROCKETREACH_API_KEY: return {}
        try:
            headers = {"Api-Key": ROCKETREACH_API_KEY}
            params = {"name": f'"{first} {last}"'}
            if company: params["current_employer"] = company
            r = requests.get(self.rocketreach_url, params=params, headers=headers, timeout=10)
            if r.ok:
                return r.json()
        except: pass
        return {}

    def search_wealthx(self, first, last):
        if not WEALTHX_API_KEY: return {}
        return {"status": "Not implemented", "source": "WealthX"}

    def search_windfall(self, first, last):
        if not WINDFALL_API_KEY: return {}
        return {"status": "Not implemented", "source": "Windfall"}

    def search_lexisnexis(self, first, last):
        if not LEXISNEXIS_API_KEY: return {}
        return {"status": "Not implemented", "source": "LexisNexis"}

    def search_inside_philanthropy(self, first, last, city=None, state=None):
        query = f'"{first} {last}" site:insidephilanthropy.com'
        if city and state:
            query += f' "{city}" "{state}"'
        elif state:
            query += f' "{state}"'
        return self.search_google(query, num=5)

    def search_chronicle_philanthropy(self, first, last, city=None, state=None):
        query = f'"{first} {last}" site:philanthropy.com'
        if city and state:
            query += f' "{city}" "{state}"'
        elif state:
            query += f' "{state}"'
        return self.search_google(query, num=5)

    def search_pacer(self, first, last):
        if not PACER_API_KEY: return {}
        return {"status": "Not implemented", "source": "PACER"}

    def search_sec_edgar(self, first, last):
        try:
            params = {
                "action": "getcompany",
                "owner": "include",
                "output": "xml",
                "count": "10",
                "company": f'"{first} {last}"'
            }
            r = requests.get(self.sec_edgar_url, params=params, timeout=10)
            if r.ok:
                return {"raw_data": r.text, "source": "SEC EDGAR"}
        except: pass
        return {}

    def search_pitchbook(self, first, last):
        if not PITCHBOOK_API_KEY: return {}
        return {"status": "Not implemented", "source": "Pitchbook"}

    def search_google(self, query, num=5):
        if not GOOGLE_API_KEY or not GOOGLE_CSE_ID: 
            # Return mock data for testing when Google API is not available
            return [{
                "title": f"Search result for {query}",
                "link": "https://example.com",
                "snippet": "Google API key not configured. Set GOOGLE_API_KEY and GOOGLE_CSE_ID environment variables."
            }]
        params = {"key": GOOGLE_API_KEY, "cx": GOOGLE_CSE_ID, "q":query, "num":num}
        try:
            r = requests.get(self.cs_url, params=params, timeout=10)
            if r.ok:
                return [{"title":i["title"],"link":i["link"],"snippet":i["snippet"]} 
                        for i in r.json().get("items",[])]
        except: pass
        return []

    def search_scholar(self, name, max_results=5):
        if not SCHOLARLY_AVAILABLE: return []
        try:
            if not name.startswith('"'):
                name = f'"{name}"'
            auth = next(scholarly.search_author(name), None)
            if not auth: return []
            scholarly.fill(auth)
            pubs = auth.get("publications", [])[:max_results]
            return [{"title":p["bib"].get("title"),"url":p.get("pub_url") or p.get("eprint_url"),
                     "year":p["bib"].get("pub_year")} for p in pubs]
        except: pass
        return []

    def fetch_irs_index(self, year):
        try:
            r = requests.get(IRS_INDEX_URL_TEMPLATE.format(year=year), timeout=15)
            if r.ok:
                return r.json().get(f"Filings{year}", [])
        except: pass
        return []

    def find_irs_filings(self, first, last, city=None, state=None):
        matches = []
        full_name = f"{first} {last}"
        name_variations = [
            full_name,
            full_name.upper(),
            f"{last}, {first}",
            f"{last} {first}"
        ]
        
        print(f"    → Searching IRS filings for \"{full_name}\"" + (f" in {city}, {state}" if city and state else f" in {state}" if state else ""))
        
        for year in range(2018, 2026):
            filings = self.fetch_irs_index(year)
            if not filings:
                continue
                
            for f in filings:
                org = f.get("OrganizationName", "")
                org_city = f.get("City", "")
                org_state = f.get("State", "")
                
                if last.upper() in org.upper():
                    if state and org_state:
                        if state.upper() == org_state.upper():
                            if city and org_city:
                                if city.lower() in org_city.lower():
                                    matches.append(f)
                            else:
                                matches.append(f)
                    elif not state:
                        matches.append(f)
                    continue
                
                url = f.get("URL")
                if not url: 
                    continue
                    
                try:
                    xml = requests.get(url, timeout=10).content
                    root = ET.fromstring(xml)
                    person_names = [e.text for e in root.findall(".//PersonName") if e.text]
                    for name_var in name_variations:
                        if name_var in person_names:
                            f["_xml"] = xml
                            matches.append(f)
                            break
                except: 
                    pass
                    
        return matches

    def parse_irs_filing(self, filing, first, last):
        xml = filing.get("_xml") or requests.get(filing["URL"], timeout=10).content
        root = ET.fromstring(xml)
        name = f"{first} {last}"
        for pn in root.findall(".//PersonName"):
            if pn.text == name:
                parent = pn.getparent() if hasattr(pn,'getparent') else pn.find("..")
                role = parent.findtext(".//Title") or "N/A"
                comp = parent.findtext(".//Compensation") or "0"
                org = root.findtext(".//Filer/BusinessName/BusinessNameLine1") or filing["OrganizationName"]
                mission = root.findtext(".//MissionDescription") or "N/A"
                return {
                    "organization": org,
                    "ein": filing["EIN"],
                    "role": role,
                    "compensation": float(comp),
                    "year": filing["TaxPeriod"][:4],
                    "mission": mission,
                    "source": f"IRS {filing['FormType']} {filing['TaxPeriod']}"
                }
        return None

    # ——— MAIN RESEARCH METHOD ——————————————
    def gather_data(self, first, last, email=None, city=None, state=None, zip_code=None):
        """Use AI orchestration for deep research"""
        
        if not self.ai_orchestrator:
            print("❌ AI Deep Research requires OpenAI API key")
            print("Set your key: export OPENAI_API_KEY=''")
            return None
        
        print(f"  🤖 Starting AI-orchestrated deep research...")
        try:
            result = self.ai_orchestrator.conduct_deep_research(
                first, last, email, city, state, zip_code
            )
            return result
        except Exception as e:
            print(f"  ✗ AI orchestration failed: {str(e)}")
            return None

    # ——— REPORT GENERATION ——————————————
    def generate_reports(self, data):
        """Generate all three required reports in IHS style"""
        if not data:
            return None
        
        results = {}
        
        # 1. Summary Table
        print("  → Generating Summary Table...")
        summary_table = data.get('summary_table', '')
        table_doc = Document()
        
        # Parse and add the markdown table content
        for line in summary_table.split('\n'):
            if line.strip():
                if '**Summary Table' in line:
                    table_doc.add_heading(line.replace('**', '').strip(), level=1)
                elif '|' in line and '---' not in line:
                    # This is a table row - we'll need to parse it properly
                    table_doc.add_paragraph(line)
                else:
                    table_doc.add_paragraph(line)
        
        results['summary_table'] = self.save_and_convert(
            table_doc, 
            f"{data['name'].replace(' ', '_')}_Summary_Table"
        )
        
        # 2. Two-Page Executive Summary  
        print("  → Generating 2-Page Executive Summary...")
        exec_doc = Document()
        
        # Format executive summary content
        exec_content = data.get('executive_summary', '')
        for paragraph in exec_content.split('\n'):
            if paragraph.strip():
                if paragraph.strip().startswith('**') and paragraph.strip().endswith('**'):
                    # This is a heading
                    heading_text = paragraph.strip().replace('**', '')
                    if 'Executive Summary' in heading_text:
                        exec_doc.add_heading(heading_text, level=1)
                    else:
                        exec_doc.add_heading(heading_text, level=2)
                elif paragraph.strip().startswith('- '):
                    # Bullet point
                    exec_doc.add_paragraph(paragraph.strip(), style="List Bullet")
                else:
                    # Regular paragraph
                    exec_doc.add_paragraph(paragraph.strip())
        
        results['executive_summary'] = self.save_and_convert(
            exec_doc, 
            f"{data['name'].replace(' ', '_')}_2pg_Executive_Summary"
        )
        
        # 3. Full Narrative Donor Report
        print("  → Generating Full Narrative Donor Report...")
        full_doc = Document()
        
        # Format the full report content
        full_content = data.get('final_report', '')
        current_style = None
        
        for paragraph in full_content.split('\n'):
            if paragraph.strip():
                # Check if it's a heading (starts and ends with **)
                if paragraph.strip().startswith('**') and paragraph.strip().endswith('**'):
                    heading_text = paragraph.strip().replace('**', '')
                    if 'Full Narrative Donor Report' in heading_text:
                        full_doc.add_heading(heading_text, level=1)
                    else:
                        full_doc.add_heading(heading_text, level=2)
                else:
                    # Regular paragraph with potential formatting
                    # Clean up any remaining ** formatting
                    clean_para = paragraph.strip().replace('**', '')
                    full_doc.add_paragraph(clean_para)
        
        results['full_report'] = self.save_and_convert(
            full_doc, 
            f"{data['name'].replace(' ', '_')}_Full_Narrative_Donor_Report"
        )
        
        # Save raw JSON data
        json_path = os.path.join("outputs", f"{data['name'].replace(' ', '_')}_raw_data.json")
        with open(json_path, 'w') as f:
            json.dump(data, f, indent=2)
        results['json'] = json_path
        
        return results

    def save_and_convert(self, doc, name):
        os.makedirs("outputs", exist_ok=True)
        docx_path = os.path.join("outputs", f"{name}.docx")
        doc.save(docx_path)
        pdf_path = docx_path.replace(".docx",".pdf")
        if DOCX2PDF_AVAILABLE:
            try: 
                docx2pdf.convert(docx_path, pdf_path)
            except: 
                print(f"    ! PDF conversion failed for {name}")
        return docx_path, pdf_path

    def run_single(self, first, last, email=None, city=None, state=None, zip_code=None):
        print(f"\n▶️  Researching {first} {last}...")
        
        if city and state:
            print(f"  📍 Location: {city}, {state}")
        elif state:
            print(f"  📍 State: {state}")
        else:
            print("  📍 No location data")
        
        # Gather data using AI
        data = self.gather_data(first, last, email, city, state, zip_code)
        
        if not data:
            print(f"❌ Failed to research {first} {last}")
            return None
        
        # Generate reports
        print("  → Generating reports...")
        results = self.generate_reports(data)
        
        if results:
            final_confidence = data.get('confidence_scores', {}).get('identity', 'N/A')
            print(f"✅  Completed research for {first} {last} (Confidence: {final_confidence}%)")
        
        return results

    def process_csv(self, csv_path):
        """Process multiple names from CSV file"""
        results = []
        
        with open(csv_path, 'r', encoding='utf-8-sig') as f:
            reader = csv.DictReader(f)
            names = list(reader)
        
        print(f"\n📋 Processing {len(names)} names from CSV...")
        
        if names and len(names) > 0:
            print(f"   📊 Columns found: {list(names[0].keys())}")
        
        for i, row in enumerate(names, 1):
            # Handle various column name formats
            full_name = None
            for key in row.keys():
                if 'full name' in key.lower() or key.strip() == 'Full Name':
                    full_name = row[key]
                    break
            
            if not full_name:
                full_name = (row.get('Name') or row.get('name') or '').strip()
            
            if full_name:
                full_name = full_name.strip()
                name_parts = full_name.split()
                if len(name_parts) >= 2:
                    first = name_parts[0]
                    last = name_parts[-1]
                    
                    if ' and ' in full_name.lower():
                        first_part = full_name.split(' and ')[0].strip().split()
                        if first_part:
                            first = first_part[0]
                        print(f"  Note: Multiple people in name: \"{full_name}\" - using {first} {last}")
                elif len(name_parts) == 1:
                    first = name_parts[0]
                    last = ''
                else:
                    first = ''
                    last = ''
            else:
                first = None
                last = None
                for key in row.keys():
                    if 'first' in key.lower() and not first:
                        first = row[key]
                    if 'last' in key.lower() and not last:
                        last = row[key]
                
                first = (first or '').strip()
                last = (last or '').strip()
            
            # Get other fields
            email = ''
            city = ''
            state = ''
            zip_code = ''
            
            for key, value in row.items():
                key_lower = key.lower()
                if 'email' in key_lower:
                    email = (value or '').strip()
                elif 'city' in key_lower:
                    city = (value or '').strip()
                elif 'state' in key_lower:
                    state = (value or '').strip()
                elif 'zip' in key_lower:
                    zip_code = (value or '').strip()
            
            print(f"\n[{i}/{len(names)}] Processing {first} {last}...")
            if city or state:
                print(f"  📍 Location data: {city}, {state}" if city else f"  📍 State: {state}")
            
            if not first and not last:
                print(f"  ⚠️  Skipping row {i}: no name found")
                continue
            
            if not last and first:
                print(f"  ⚠️  Warning: Only first name '{first}' found - results may be limited")
            
            try:
                result = self.run_single(first, last, email, city, state, zip_code)
                if result:
                    result['row'] = i
                    result['name'] = f"{first} {last}"
                    results.append(result)
            except Exception as e:
                print(f"  ❌ Error processing {first} {last}: {str(e)}")
                results.append({
                    'row': i,
                    'name': f"{first} {last}",
                    'error': str(e)
                })
            
            time.sleep(1)  # Brief pause between requests
        
        # Generate batch summary
        self._generate_batch_summary(results)
        
        return results

    def _generate_batch_summary(self, results):
        """Generate summary report for batch processing"""
        os.makedirs("outputs", exist_ok=True)
        
        doc = Document()
        doc.add_heading("Batch Processing Summary", level=1)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Total processed: {len(results)}")
        
        successful = [r for r in results if 'error' not in r]
        failed = [r for r in results if 'error' in r]
        
        doc.add_paragraph(f"Successful: {len(successful)}")
        doc.add_paragraph(f"Failed: {len(failed)}")
        
        if successful:
            doc.add_heading("Successfully Processed", level=2)
            for r in successful:
                doc.add_paragraph(f"• {r['name']} - Reports generated", style="List Bullet")
        
        if failed:
            doc.add_heading("Failed Processing", level=2)
            for r in failed:
                doc.add_paragraph(f"• {r['name']} - Error: {r['error']}", style="List Bullet")
        
        doc.save(os.path.join("outputs", "batch_summary.docx"))


# ——— CLI ——————————————————————————————————————————
def main():
    p = argparse.ArgumentParser(
        description="AI Deep Research Tool - Uses OpenAI to orchestrate comprehensive donor research for IHS",
        epilog="""
        SETUP INSTRUCTIONS:
        1. Set your OpenAI API key: export OPENAI_API_KEY='sk-your-key-here'
        2. Set other API keys for better results (optional)
        3. Create a data.csv file with donor names
        4. Run: python spider.py
        
        GENERATED REPORTS:
        - Summary Table: Key prospect information in table format
        - 2-Page Executive Summary: Concise overview with giving probability
        - Full Narrative Report: Comprehensive donor intelligence
        """
    )
    
    # Single name mode
    p.add_argument("--first", help="First name (for single lookup)")
    p.add_argument("--last", help="Last name (for single lookup)")
    p.add_argument("--email", help="Email address")
    p.add_argument("--city", help="City")
    p.add_argument("--state", help="State")
    p.add_argument("--zip", dest="zip_code", help="ZIP code")
    
    # CSV mode
    p.add_argument("--csv", help="Path to CSV file (default: data.csv)", default="data.csv")
    p.add_argument("--single", action="store_true", help="Process single name instead of CSV")
    
    args = p.parse_args()
    
    # Check for OpenAI API key
    if not OPENAI_API_KEY:
        print("\n❌ OpenAI API key not found!")
        print("\nTo use this AI Deep Research tool, you need an OpenAI API key.")
        print("\nSet it up:")
        print("1. Get your API key from: https://platform.openai.com/api-keys")
        print("2. Set it in your terminal:")
        print("   export OPENAI_API_KEY='sk-proj-G2AT8ZYNNTcc4he-'")
        print("\nThen run this script again.")
        return
    
    scraper = EnhancedDonorScraper()
    
    if not scraper.ai_orchestrator:
        print("❌ Failed to initialize AI orchestrator")
        return
    
    # Show active APIs
    active_apis = ["OpenAI (Required)"]
    if WEALTHENGINE_API_KEY: active_apis.append("WealthEngine")
    if IWAVE_API_KEY: active_apis.append("iWave")
    if FEC_API_KEY: active_apis.append("FEC")
    if OPENSECRETS_API_KEY: active_apis.append("OpenSecrets")
    if CLEARBIT_API_KEY: active_apis.append("Clearbit")
    if PDL_API_KEY: active_apis.append("People Data Labs")
    if GOOGLE_API_KEY: active_apis.append("Google")
    if FOLLOWTHEMONEY_KEY: active_apis.append("FollowTheMoney")
    
    print(f"\n🔑 Active APIs: {', '.join(active_apis)}")
    
    # Warn about missing APIs
    if not GOOGLE_API_KEY:
        print("\n⚠️  Google API not configured - web searches will be limited")
        print("   Set GOOGLE_API_KEY and GOOGLE_CSE_ID for better results")
    
    if len(active_apis) < 3:
        print("\n💡 Tip: Add more API keys for deeper research:")
        print("   - FEC_API_KEY for political giving data")
        print("   - CLEARBIT_API_KEY for email enrichment")
        print("   - PDL_API_KEY for comprehensive people data")
    
    print("\n🤖 AI Deep Research Engine Ready (IHS Donor Research Mode)")
    print("   • AI orchestrates searches across all available databases")
    print("   • Generates IHS-specific donor intelligence reports")
    print("   • Creates executive summaries with giving probability")
    print("   • Provides full narrative reports with strategic recommendations\n")
    
    if args.single and args.first and args.last:
        # Single name lookup
        out = scraper.run_single(
            args.first, args.last,
            email=args.email,
            city=args.city,
            state=args.state,
            zip_code=args.zip_code
        )
        
        if out:
            print("\n✅ Reports generated in ./outputs/:")
            print(f"  • Summary Table: {out['summary_table'][0]}")
            print(f"  • Executive Summary: {out['executive_summary'][0]}")
            print(f"  • Full Narrative Report: {out['full_report'][0]}")
            print(f"  • Raw Data: {out['json']}")
    else:
        # Process CSV file
        csv_file = args.csv
        
        if not os.path.exists(csv_file):
            if csv_file == "data.csv":
                print(f"❌ No 'data.csv' file found")
                print(f"\nCreate a CSV file with one of these formats:")
                print("\nFormat 1:")
                print("  Full Name, City, State, Email")
                print("  John Smith, Boston, MA, john@example.com")
                print("\nFormat 2:")
                print("  first_name, last_name, city, state, email")
                print("  John, Smith, Boston, MA, john@example.com")
            else:
                print(f"❌ CSV file not found: {csv_file}")
            return
        
        print(f"\n📄 Processing: {csv_file}")
        results = scraper.process_csv(csv_file)
        
        print(f"\n✅ Batch processing complete!")
        print(f"Reports saved in ./outputs/")
    
    print(f"\n📊 Research complete")


if __name__ == "__main__":
    main()

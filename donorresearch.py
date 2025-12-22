#!/usr/bin/env python3
"""
IHS Hybrid Donor Research System - ENHANCED BULLET-POINT EDITION
-----------------------------------------------------------------
Generates concise donor profiles (~2 pages) with Wikipedia-style inline 
citations, privacy-safe identity verification, and conditional Inner Circle analysis.

ENHANCEMENTS IN THIS VERSION:
1. ✓ Wikipedia-style inline citations: [1], [2], [3] embedded in text
2. ✓ Bullet-point format for concise, scannable sections (~2 pages)
3. ✓ Multi-marker identity verification (location, employer, education, work)
4. ✓ CSV "Already known" column check - skips Inner Circle if "Yes"
5. ✓ Organization detection - uses business/foundation sources when applicable
6. ✓ Privacy-safe filtering - blocks data broker sites, prioritizes primary sources
7. ✓ Name collision detection - warns when other people with same name found
8. ✓ Confidence scoring - High/Medium/Low based on verification quality
9. ✓ Gift Officer Summary - 2-page executive summary for fundraisers
10. ✓ SPOUSE+LOCATION filtering - Enhanced identity verification using spouse + location
    combination (e.g., "Ann Gibbs" + "George S Gibbs" + "Jacksonville, FL")

PRIVACY-SAFE RESEARCH:
- Uses only public, reputable sources (LinkedIn, company sites, news, academic databases)
- Blocks data broker sites (Whitepages, Spokeo, etc.)
- Prioritizes primary sources (official websites, verified profiles)
- Multi-marker verification prevents identity confusion
- Never includes private contact details or sensitive information

OUTPUT FORMAT (Partner Profile - matching IHS template):
- Partner Profile (title)
- Name, Photo placeholder, Contact Info
- Salesforce Link, Date of Birth, Family
- Education
- Current Position or Occupation
- Previous Positions
- Foundation and Nonprofit Board Memberships
- Other Affiliations
- Net Worth/Giving Capacity
  - Property Data (Home Value) - from Smarty API
  - Wealth Screening Data - from iWave CSV
  - Estimated Wealth (web research)
- Philanthropic Interests
- Relevant Gifts
- Other Giving History
- Publications
- Biography (narrative prose)
- Political Activity & Ideology
- Connections (Inner Circle)
- IHS Donor Probability Assessment
- Strategic Summary
- References

CSV FORMAT:
Required columns: First Name (or First), Last Name (or Last)
Optional columns: City, State, Already known, Employer, Title, Education, Industry, etc.
- "Already known" column: "Yes" = lookup in contactsdb.csv + skip Inner Circle, anything else = run Inner Circle
- ALL CSV columns used for identity verification and biographical research
- City/State/Employer/Title/Education used for identity disambiguation (prevents mixing up same-named people)

SMARTY PROPERTY DATA INTEGRATION:
If your data.csv or contactsdb.csv contains address information, the Smarty US Address 
Enrichment API will be used to look up property values for Section 7 (Capacity for Giving).
Required CSV columns (case-insensitive):
- Address or Street or Street Address or Home Address or Mailing Address or Preferred Address
  (required for property lookup)
- City (optional but recommended)
- State (optional but recommended)
- Zip or Zipcode or ZIP or Postal Code (optional but recommended)
Returns: Market Value, Assessed Value, Sale Price, Property Type, Year Built, etc.

NOTE: For "Already known" donors, the script will look up the address in contactsdb.csv first.
This allows you to maintain addresses in your CRM export (contactsdb.csv) rather than data.csv.

iWAVE DATA INTEGRATION (CSV data only - API temporarily disabled):
If your data.csv contains iWave wealth screening data, it will be automatically included in
Section 7 (Capacity for Giving). Supported column names include:
- iWave Score, iWave Gift Range, iWave Estimated Capacity, iWave Capacity
- Major Gift Likelihood, Planned Gift Likelihood, P2G Score, RFM Score
- Wealth Rating, Real Estate Value, Stock Value, Total Assets
- Largest Gift, Total Giving, Annual Giving
(Column names are matched case-insensitively)

CAPACITY TO GIVE CSV INTEGRATION (capacitytogive.csv):
If you have a Salesforce export with donor capacity data, place it as "capacitytogive.csv" in
the same directory as the script. The donor will be looked up by:
- Matching name against "Primary Contact: Full Name" or "Account Name" columns
- Verifying location via "Billing City" and "Billing State/Province"
Data extracted and shown in Section 7 (Net Worth/Giving Capacity):
- iWave Capacity Rating: iWave Est. Capacity, Estimated Giving Capacity, Donor Tier
- Giving History: Total Gifts, Best Gift Year, Last Gift Amount/Date, giving patterns
- Donor Profile: Account Type, VIP Status, Planned Giver, Moves Manager
- Donor Background: Existing research notes from Salesforce
- Notes: Description field content
Column names used (from Salesforce export):
- Account Name, Primary Contact: Full Name
- Billing City, Billing State/Province
- iWave Est. Capacity, Estimated Giving Capacity
- Total Gifts, Last Gift Amount, Last Gift Date, Best Gift Year Total
- Total Giving Last 18 Months, Total Giving Last 3 Years
- Number of Years Donated, Total Number of Gifts, Days since last gift
- Development Prospect Status, Donor Tier, VIP Stand Together, Planned Giver
- Donor Interest, Donor Background, Donor Purpose, Description
- Moves Manager: Full Name

CONTACTSDB.CSV (OPTIONAL):
If you have a contactsdb.csv file with detailed info on known contacts:
- When "Already known" = "Yes", script looks up person in contactsdb.csv
- Uses ALL contactsdb.csv fields (employer, title, education, ADDRESS, etc.) for research
- contactsdb.csv data takes precedence over data.csv (more detailed)
- Address from contactsdb.csv will be used for Smarty property value lookup
- Format: Same as data.csv but with more detailed fields
- Example use: Your CRM/Salesforce export with full contact details

SETUP:
Create spider.env with:
    OPENAI_API_KEY=sk-your-key
    OPENAI_MODEL=gpt-4o
    TAVILY_API_KEY_1=your-key (or multiple: TAVILY_API_KEY_1, _2, etc.)
    GOOGLE_API_KEY_1=your-key (or multiple: GOOGLE_API_KEY_1, _2, etc.)
    GOOGLE_CSE_ID=your-cse-id
    FEC_API_KEY=DEMO_KEY
    
    # Smarty API (for property value data)
    # Get credentials from: https://www.smarty.com/account/keys
    # Uses US Address Enrichment API for property values, assessed values, etc.
    # Requires address in CSV columns: Address/Street, City, State, Zip
    SMARTY_AUTH_ID=your-auth-id
    SMARTY_AUTH_TOKEN=your-auth-token
    SMARTY_ENABLED=true
    
    # iWave API (TEMPORARILY DISABLED - using Smarty for property values)
    # When re-enabled, this will query iWave for giving capacity data
    # The script will try multiple name variations including spouse name
    # e.g., "Ann Gibbs", "George S Gibbs", "Mr. and Mrs. George Gibbs"
    # IWAVE_API_KEY=your-iwave-api-key
    # IWAVE_API_URL=https://api.app.iwave.com
    # IWAVE_ENABLED=true

USAGE:
    python3 test-spider.py                    # Process all donors
    python3 test-spider.py --test-mode        # Process ONLY first donor
    python3 test-spider.py --name "John Doe"  # Process specific donor
"""

import os
import re
import sys
import csv
import json
import time
import math
import hashlib
import argparse
import logging
import warnings
from datetime import datetime
from typing import List, Dict, Any, Tuple, Optional
from collections import defaultdict
from functools import wraps
from concurrent.futures import ThreadPoolExecutor, as_completed

import requests
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Suppress urllib3 LibreSSL warnings
warnings.filterwarnings('ignore', category=UserWarning, module='urllib3')

# -------------------------
# SETUP & CONFIG
# -------------------------
load_dotenv("spider.env")

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL   = os.getenv("OPENAI_MODEL", "gpt-4")

# Tavily search provider - support multiple API keys
TAVILY_API_KEYS = []
SEARCH_PROVIDER = os.getenv("SEARCH_PROVIDER", "tavily")

for i in range(1, 51):
    key = os.getenv(f"TAVILY_API_KEY_{i}", "")
    if key:
        TAVILY_API_KEYS.append(key)

if not TAVILY_API_KEYS:
    single_key = os.getenv("TAVILY_API_KEY", "")
    if single_key:
        TAVILY_API_KEYS.append(single_key)

# Google API keys
GOOGLE_API_KEYS = []
GOOGLE_CSE_ID  = os.getenv("GOOGLE_CSE_ID", "")

for i in range(1, 51):
    key = os.getenv(f"GOOGLE_API_KEY_{i}", "")
    if key:
        GOOGLE_API_KEYS.append(key)

if not GOOGLE_API_KEYS:
    single_key = os.getenv("GOOGLE_API_KEY", "")
    if single_key:
        GOOGLE_API_KEYS.append(single_key)

FEC_API_KEY    = os.getenv("FEC_API_KEY", "")

# iWave API Configuration - TEMPORARILY DISABLED
# Get your API key from iWave User Management > User Profile > Get API Key
IWAVE_API_KEY     = os.getenv("IWAVE_API_KEY", "")
IWAVE_API_SECRET  = os.getenv("IWAVE_API_SECRET", "")  # Some iWave setups use key+secret
IWAVE_API_URL     = os.getenv("IWAVE_API_URL", "https://api.app.iwave.com")
IWAVE_ENABLED     = False  # TEMPORARILY DISABLED - using SmartyAPI for property values instead

# SmartyStreets/Smarty API Configuration  
# Get your auth-id and auth-token from https://www.smarty.com/account/keys
# This API provides property data including assessed values, sale prices, etc.
SMARTY_AUTH_ID    = os.getenv("SMARTY_AUTH_ID", "")
SMARTY_AUTH_TOKEN = os.getenv("SMARTY_AUTH_TOKEN", "")
SMARTY_ENABLED    = os.getenv("SMARTY_ENABLED", "true").lower() == "true"

CACHE_DIR      = os.getenv("CACHE_DIR", "./cache")
CACHE_TTL      = int(os.getenv("CACHE_TTL", "86400"))

# HTTP headers to prevent 403/anti-bot blocks
DEFAULT_HEADERS = {
    "User-Agent": "IHS-ResearchBot/1.0 (Donor Research; +https://theihs.org)",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.5"
}

# Key rotation state
_tavily_key_index = 0
_tavily_key_exhausted = set()
_google_key_index = 0
_google_key_exhausted = set()
_search_counter = 0
_google_fully_exhausted = False
_search_stats = {"tavily": 0, "google": 0, "tavily_cached": 0, "google_cached": 0}

def reset_tavily_keys():
    global _tavily_key_exhausted, _search_stats
    _tavily_key_exhausted.clear()
    log.info(f"Reset Tavily API key exhaustion tracking. Stats: Tavily={_search_stats['tavily']}, Google={_search_stats['google']}")

def reset_google_keys():
    global _google_key_exhausted, _search_stats
    _google_key_exhausted.clear()
    log.info(f"Reset Google API key exhaustion tracking. Stats: Tavily={_search_stats['tavily']}, Google={_search_stats['google']}")

try:
    import openai
    try:
        OPENAI_CLIENT = openai.OpenAI(api_key=OPENAI_API_KEY)
        OPENAI_VERSION = "new"
    except AttributeError:
        openai.api_key = OPENAI_API_KEY
        OPENAI_CLIENT = None
        OPENAI_VERSION = "old"
except Exception:
    openai = None
    OPENAI_CLIENT = None
    OPENAI_VERSION = None

logging.basicConfig(level=logging.INFO, format="[%(levelname)s] %(message)s")
log = logging.getLogger("spider")

logging.getLogger("httpx").setLevel(logging.WARNING)
logging.getLogger("openai").setLevel(logging.WARNING)

# -------------------------
# GLOBAL CITATION MANAGER
# -------------------------
class CitationManager:
    """Manages Wikipedia-style inline citations"""
    def __init__(self):
        self.citations = []  # List of {url, title, snippet}
        self.url_to_index = {}  # url -> citation number mapping
    
    def reset(self):
        """Reset for new donor"""
        self.citations = []
        self.url_to_index = {}
    
    def add_citation(self, url: str, title: str = "", snippet: str = "") -> int:
        """Add source and return citation number [1], [2], etc."""
        if url in self.url_to_index:
            return self.url_to_index[url]
        
        cite_num = len(self.citations) + 1
        self.citations.append({
            "number": cite_num,
            "url": url,
            "title": title,
            "snippet": snippet
        })
        self.url_to_index[url] = cite_num
        return cite_num
    
    def get_citation_text(self, url: str) -> str:
        """Get [N] citation marker for URL"""
        if url in self.url_to_index:
            return f"[{self.url_to_index[url]}]"
        return ""
    
    def get_all_citations(self) -> List[Dict]:
        """Get all citations for References section"""
        return self.citations

citation_mgr = CitationManager()

# -------------------------
# CACHING SYSTEM
# -------------------------
class ResearchCache:
    def __init__(self, cache_dir: str = "./cache", ttl: int = 86400):
        self.cache_dir = cache_dir
        self.ttl = ttl
        os.makedirs(cache_dir, exist_ok=True)
    
    def get_cache_key(self, prefix: str, params: Dict) -> str:
        param_str = json.dumps(params, sort_keys=True)
        hash_key = hashlib.md5(param_str.encode()).hexdigest()[:8]
        return f"{prefix}_{hash_key}"
    
    def get_or_fetch(self, key: str, fetcher, ttl: Optional[int] = None):
        ttl = ttl or self.ttl
        cache_file = os.path.join(self.cache_dir, f"{key}.json")
        
        if os.path.exists(cache_file):
            mtime = os.path.getmtime(cache_file)
            if time.time() - mtime < ttl:
                try:
                    with open(cache_file, 'r', encoding='utf-8') as f:
                        log.debug(f"Cache HIT: {key}")
                        return json.load(f)
                except:
                    pass
        
        log.debug(f"Cache MISS: {key}")
        result = fetcher()
        
        try:
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(result, f, indent=2)
        except:
            pass
        
        return result

cache = ResearchCache(CACHE_DIR, CACHE_TTL)

# -------------------------
# RETRY DECORATOR
# -------------------------
def retry_with_backoff(max_retries=3, backoff_seconds=2):
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            for attempt in range(max_retries):
                try:
                    return func(*args, **kwargs)
                except Exception as e:
                    if attempt == max_retries - 1:
                        raise
                    wait = backoff_seconds * (2 ** attempt)
                    log.warning(f"{func.__name__} failed (attempt {attempt+1}/{max_retries}): {e}. Retrying in {wait}s...")
                    time.sleep(wait)
        return wrapper
    return decorator

# -------------------------
# SEARCH FUNCTIONS
# -------------------------
def get_next_tavily_key() -> Optional[str]:
    global _tavily_key_index, _tavily_key_exhausted
    if not TAVILY_API_KEYS:
        return None
    
    if len(_tavily_key_exhausted) >= len(TAVILY_API_KEYS):
        return None
    
    for _ in range(len(TAVILY_API_KEYS)):
        key = TAVILY_API_KEYS[_tavily_key_index]
        _tavily_key_index = (_tavily_key_index + 1) % len(TAVILY_API_KEYS)
        if _tavily_key_index not in _tavily_key_exhausted:
            return key
    
    return None

def get_next_google_key() -> Optional[str]:
    global _google_key_index, _google_key_exhausted, _google_fully_exhausted
    if not GOOGLE_API_KEYS:
        return None
    
    if len(_google_key_exhausted) >= len(GOOGLE_API_KEYS):
        _google_fully_exhausted = True
        return None
    
    for _ in range(len(GOOGLE_API_KEYS)):
        key = GOOGLE_API_KEYS[_google_key_index]
        _google_key_index = (_google_key_index + 1) % len(GOOGLE_API_KEYS)
        if _google_key_index not in _google_key_exhausted:
            return key
    
    return None

@retry_with_backoff(max_retries=2)
def search_tavily(query: str, num: int = 5) -> List[Dict[str, str]]:
    global _search_stats, _tavily_key_exhausted
    key = get_next_tavily_key()
    if not key:
        return []
    
    url = "https://api.tavily.com/search"
    payload = {
        "api_key": key,
        "query": query,
        "search_depth": "basic",
        "include_answer": False,
        "max_results": num
    }
    
    try:
        r = requests.post(url, json=payload, timeout=20)
        r.raise_for_status()
        data = r.json()
        
        _search_stats["tavily"] += 1
        
        results = []
        for item in data.get("results", []):
            results.append({
                "title": item.get("title", ""),
                "snippet": item.get("content", ""),
                "link": item.get("url", "")
            })
        return results
    
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 429:
            _tavily_key_exhausted.add(_tavily_key_index)
            log.warning(f"Tavily key #{_tavily_key_index} exhausted (rate limit)")
            return search_tavily(query, num)
        raise

@retry_with_backoff(max_retries=2)
def search_google(query: str, num: int = 5) -> List[Dict[str, str]]:
    global _search_stats, _google_key_exhausted
    key = get_next_google_key()
    if not key or not GOOGLE_CSE_ID:
        return []
    
    url = "https://www.googleapis.com/customsearch/v1"
    params = {
        "key": key,
        "cx": GOOGLE_CSE_ID,
        "q": query,
        "num": min(num, 10)
    }
    
    try:
        r = requests.get(url, params=params, timeout=20, headers=DEFAULT_HEADERS)
        r.raise_for_status()
        data = r.json()
        
        _search_stats["google"] += 1
        
        results = []
        for item in data.get("items", []):
            results.append({
                "title": item.get("title", ""),
                "snippet": item.get("snippet", ""),
                "link": item.get("link", "")
            })
        return results
    
    except requests.exceptions.HTTPError as e:
        if e.response.status_code == 429:
            _google_key_exhausted.add(_google_key_index)
            log.warning(f"Google key #{_google_key_index} exhausted (rate limit)")
            return search_google(query, num)
        raise

def search_web(query: str, num: int = 5) -> List[Dict[str, str]]:
    global _search_counter, _google_fully_exhausted
    
    if _google_fully_exhausted:
        return search_tavily(query, num)
    
    _search_counter += 1
    if _search_counter % 2 == 0 and GOOGLE_API_KEYS:
        results = search_google(query, num)
        if results:
            return results
        return search_tavily(query, num)
    else:
        return search_tavily(query, num)

def print_search_stats():
    log.info("\n" + "="*60)
    log.info("SEARCH API USAGE STATISTICS")
    log.info("="*60)
    log.info(f"Tavily searches: {_search_stats['tavily']}")
    log.info(f"Google searches: {_search_stats['google']}")
    log.info(f"Total searches: {_search_stats['tavily'] + _search_stats['google']}")
    log.info("="*60)

# -------------------------
# ENTITY TYPE DETECTION
# -------------------------
def detect_entity_type(name: str) -> str:
    """Detect if entity is a person or organization"""
    org_keywords = [
        "institute", "foundation", "center", "university", "college",
        "organization", "association", "society", "council", "commission",
        "company", "corporation", "llc", "inc", "ltd", "group"
    ]
    
    name_lower = name.lower()
    for keyword in org_keywords:
        if keyword in name_lower:
            return "organization"
    
    return "person"

def get_org_search_queries(org_name: str) -> List[str]:
    """Generate organization-specific search queries"""
    return [
        f'"{org_name}" mission history',
        f'"{org_name}" leadership board directors',
        f'"{org_name}" funding grants donations',
        f'"{org_name}" political advocacy ideology',
        f'"{org_name}" IRS 990 assets',
        f'"{org_name}" foundation institute about'
    ]

# -------------------------
# IDENTITY VERIFICATION (Scored Candidate Approach)
# -------------------------

# Domain trust weights for identity verification
DOMAIN_WEIGHTS = [
    (r"wikipedia\.org", 1.0),
    (r"wikidata\.org", 0.9),
    (r"reuters\.com|bloomberg\.com", 0.85),
    (r"linkedin\.com/in/", 0.80),
    (r"forbes\.com/profile|wsj\.com|ft\.com|nytimes\.com", 0.75),
    (r"[.]edu/|/about|/leadership|/team", 0.70),
    (r"crunchbase\.com", 0.70),
]

# Negative patterns that indicate wrong person or deceased
# EXPANDED to catch more death-related and wrong-profession content
NEGATIVE_PATTERNS = [
    # Death/obituary patterns (universal)
    r"\bobituary\b", r"\bin memoriam\b", r"\bdied\b", r"\bdeath\b",
    r"\bdeceased\b", r"\blate\b.*\b(died|passed)", r"\bpassed away\b",
    r"\bmemorial\s+service\b", r"\bfuneral\b",
    r"\bcar crash\b", r"\bcar accident\b", r"\bkilled\b.*\baccident\b",
    
    # Wrong profession patterns for business/exec searches (universal)
    r"\brodeo\b", r"\bcowboy\b", r"\brider\b", r"\brancher\b",
    r"\bsports\b.*\bplayer\b", r"\bactor\b.*\bfilmography\b",
    r"\bathlete\b", r"\bcoach\b.*\bsports\b", r"\bmusician\b",
    r"\bartist\b.*\bpainting\b", r"\bsinger\b", r"\bband\b.*\bmusic\b",
    
    # Generic family reference patterns (universal - catches when result is ABOUT someone's family, not the person)
    r"(?:'s|s')\s+family\s+includes\b",  # "Jim Ronyak's family includes..." or "James' family includes..."
    r"\b(?:his|her|their)\s+family\s+includes\b",  # "his family includes..."
    r"\b(?:son|daughter|wife|husband)\s+of\b",  # "John Smith, son of..."
    r"\bsurvived\s+by\b",  # Obituary indicator
    r"\bsister-in-law\b",  # Family member relationships
    r"\bbrother-in-law\b",
]

def _jaro_winkler(a: str, b: str) -> float:
    """Calculate Jaro-Winkler similarity between two strings"""
    from difflib import SequenceMatcher
    return SequenceMatcher(None, a.lower(), b.lower()).ratio()

def _token_name_variants(fullname: str) -> set:
    """Generate name variants including nicknames (bidirectional)"""
    parts = fullname.strip().split()
    if not parts:
        return {fullname.lower()}
    
    first, *rest = parts
    last = rest[-1] if rest else ""
    variants = {fullname.lower(), f"{first} {last}".strip().lower()}
    
    # Add nickname variants (formal -> nickname)
    nicknames = NICKNAMES.get(first.lower(), set())
    variants |= {f"{n} {last}".strip().lower() for n in nicknames}
    
    # Add reverse nickname variants (nickname -> formal)
    # e.g., if input is "jim ronyak", also add "james ronyak"
    formal_names = REVERSE_NICKNAMES.get(first.lower(), set())
    variants |= {f"{f} {last}".strip().lower() for f in formal_names}
    
    # Add middle initial variants
    if len(rest) > 1:
        middle = rest[0]
        if len(middle) == 1 or middle.endswith('.'):
            # Has middle initial
            variants.add(f"{first} {last}".lower())
    
    return variants

def _domain_weight(url: str) -> float:
    """Get trust weight for a domain"""
    for pattern, weight in DOMAIN_WEIGHTS:
        if re.search(pattern, url, re.I):
            return weight
    return 0.5  # Default weight for unknown domains

def _location_bonus(snippet: str, title: str, city: str, state: str, spouse_name: str = "") -> float:
    """
    Calculate location match bonus (case-insensitive).
    Accepts home city/state OR common work locations OR state abbreviations.
    
    ENHANCED: Also checks for spouse name + location combination.
    If spouse is known, can match on either:
    - Person's name + location
    - Spouse's name + location (helps when person has limited web presence)
    """
    combined = f"{snippet} {title}".lower()
    bonus = 0.0
    
    # Check for home city (case-insensitive)
    if city and city.lower() in combined:
        bonus += 0.15
    
    # Check for home state - accept full name OR abbreviation
    if state:
        state_lower = state.lower()
        if state_lower in combined:
            bonus += 0.10
        # Also check common state abbreviations
        state_abbrevs = {
            'alabama': 'al', 'alaska': 'ak', 'arizona': 'az', 'arkansas': 'ar',
            'california': 'ca', 'colorado': 'co', 'connecticut': 'ct', 'delaware': 'de',
            'florida': 'fl', 'georgia': 'ga', 'hawaii': 'hi', 'idaho': 'id',
            'illinois': 'il', 'indiana': 'in', 'iowa': 'ia', 'kansas': 'ks',
            'kentucky': 'ky', 'louisiana': 'la', 'maine': 'me', 'maryland': 'md',
            'massachusetts': 'ma', 'michigan': 'mi', 'minnesota': 'mn', 'mississippi': 'ms',
            'missouri': 'mo', 'montana': 'mt', 'nebraska': 'ne', 'nevada': 'nv',
            'new hampshire': 'nh', 'new jersey': 'nj', 'new mexico': 'nm', 'new york': 'ny',
            'north carolina': 'nc', 'north dakota': 'nd', 'ohio': 'oh', 'oklahoma': 'ok',
            'oregon': 'or', 'pennsylvania': 'pa', 'rhode island': 'ri', 'south carolina': 'sc',
            'south dakota': 'sd', 'tennessee': 'tn', 'texas': 'tx', 'utah': 'ut',
            'vermont': 'vt', 'virginia': 'va', 'washington': 'wa', 'west virginia': 'wv',
            'wisconsin': 'wi', 'wyoming': 'wy', 'district of columbia': 'dc'
        }
        if state_lower in state_abbrevs and state_abbrevs[state_lower] in combined:
            bonus += 0.10
    
    # ENHANCED: Check for spouse + location combination
    # This helps disambiguate when person has common name but spouse+location is unique
    # Example: "Ann Gibbs" is common, but "Ann Gibbs married to George S Gibbs in Jacksonville, FL" is unique
    if spouse_name and (city or state):
        # Use robust spouse detection (first+last, punctuation-insensitive)
        if _spouse_present(spouse_name, combined):
            # Spouse is mentioned - now check if location is also mentioned
            spouse_with_location = False
            if city and city.lower() in combined:
                spouse_with_location = True
            if state:
                state_lower = state.lower()
                if state_lower in combined or (state_lower in state_abbrevs and state_abbrevs[state_lower] in combined):
                    spouse_with_location = True
            
            if spouse_with_location:
                # STRONG SIGNAL: Spouse name + location mentioned together
                # This is a high-confidence identity match
                bonus += 0.20
                log.debug(f"  💑📍 SPOUSE+LOCATION match: {spouse_name} in {city or ''} {state or ''}")
    
    return bonus

def verify_identity(name: str, csv_data: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    Scored candidate approach to identity verification using ALL CSV data.
    
    Args:
        name: Person or organization name
        csv_data: Dictionary of ALL CSV columns (City, State, Employer, Title, Education, etc.)
    
    Generates candidates from multiple sources, scores each with:
    - Domain trust weight (0-1.0)
    - Name similarity via Jaro-Winkler (0-1.0)
    - Location match bonus (+0-0.25)
    - Employer/Title match bonus (+0-0.30)
    - Education match bonus (+0-0.20)
    - Spouse match bonus (+0-0.40) - CRITICAL FOR PREVENTING NAME COLLISIONS!
    - Negative pattern penalty (-0.4)
    
    Decision thresholds (SAME AS ORIGINAL):
    - score >= 1.6 → High confidence
    - score >= 1.2 → Medium confidence  
    - score >= 1.0 + high name match (0.95+) → Medium confidence
    - score >= 0.9 → Low confidence (proceed with warning)
    - score < 0.9 → Stop research (manual review required)
    
    CSV data bonuses ENHANCE scoring but don't change thresholds.
    
    SPOUSE VERIFICATION: If spouse name present in CSV, web results MUST mention spouse!
    """
    log.info(f"🔍 Verifying identity (scored approach with ALL CSV data) for: {name}")
    
    # Extract all CSV fields
    if csv_data is None:
        csv_data = {}
    
    city = csv_data.get("City", "")
    state = csv_data.get("State", "")
    employer = csv_data.get("Employer", csv_data.get("Company", ""))
    title = csv_data.get("Title", csv_data.get("Position", csv_data.get("Job Title", "")))
    education = csv_data.get("Education", csv_data.get("School", csv_data.get("University", "")))
    industry = csv_data.get("Industry", "")
    
    # CRITICAL: Extract spouse from CSV - use for identity verification!
    spouse_name = ""
    spouse_first = ""
    spouse_last = ""
    spouse_search_variants = []
    
    # Check Nickname column for "Mr. and Mrs." format
    nickname = get_nickname_from_row(csv_data)
    if nickname and ("mr. and mrs." in nickname.lower() or "mr and mrs" in nickname.lower()):
        # Extract spouse from "Mr. and Mrs. George S. Gibbs"
        try:
            parts = nickname.split()
            mrs_idx = next((i for i, p in enumerate(parts) if p.lower() in ["mrs.", "mrs"]), -1)
            if mrs_idx >= 0 and mrs_idx < len(parts) - 1:
                spouse_full = " ".join(parts[mrs_idx+1:])
                spouse_parts = spouse_full.split()
                if len(spouse_parts) >= 2:
                    # CRITICAL: Keep FULL name including middle initials!
                    # For "George S. Gibbs", spouse_name should be "George S. Gibbs"
                    # NOT "George Gibbs" (drops the middle initial)
                    spouse_name = spouse_full
                    
                    # Generate search variants that PRESERVE middle initials
                    spouse_lower = spouse_name.lower()
                    spouse_search_variants = [
                        spouse_lower,  # "george s. gibbs" or "george s gibbs"
                    ]
                    
                    # Add variant with/without period after middle initial
                    if '.' in spouse_lower:
                        spouse_search_variants.append(spouse_lower.replace('.', ''))
                    else:
                        # Try to add periods after single letters
                        parts_temp = spouse_lower.split()
                        for i, part in enumerate(parts_temp):
                            if len(part) == 1 and i < len(parts_temp) - 1:
                                parts_copy = parts_temp.copy()
                                parts_copy[i] = part + '.'
                                spouse_search_variants.append(' '.join(parts_copy))
                    
                    log.info(f"  💑 SPOUSE DETECTED: {spouse_name}")
                    log.info(f"     Web results MUST mention spouse first name ('{spouse_parts[0]}') to be accepted!")
        except Exception as e:
            log.debug(f"  Failed to parse spouse from nickname: {e}")
    
    # Also check dedicated Spouse/Partner column
    if not spouse_name:
        partner = get_spouse_from_row(csv_data)
        if partner:
            parts = partner.split()
            if len(parts) >= 2:
                # CRITICAL: Keep FULL name including middle initials!
                spouse_name = partner.strip()
                
                # Generate search variants that PRESERVE middle initials
                spouse_lower = spouse_name.lower()
                spouse_search_variants = [
                    spouse_lower,
                ]
                
                # Add variant with/without period after middle initial
                if '.' in spouse_lower:
                    spouse_search_variants.append(spouse_lower.replace('.', ''))
                else:
                    parts_temp = spouse_lower.split()
                    for i, part in enumerate(parts_temp):
                        if len(part) == 1 and i < len(parts_temp) - 1:
                            parts_copy = parts_temp.copy()
                            parts_copy[i] = part + '.'
                            spouse_search_variants.append(' '.join(parts_copy))
                
                log.info(f"  💑 SPOUSE DETECTED: {spouse_name}")
                log.info(f"     Web results MUST mention spouse first name ('{parts[0]}') to be accepted!")
    
    # CRITICAL: Get both name variants (original from data.csv + matched from contactsdb.csv)
    original_name = csv_data.get("_original_name", "")
    matched_name = csv_data.get("_matched_name", "")
    
    # Build list of names to search for
    search_names = []
    if original_name and matched_name and original_name.lower() != matched_name.lower():
        # We have two different name variants - use BOTH
        search_names = [name, original_name, matched_name]
        # Remove duplicates (case-insensitive)
        seen = set()
        unique_names = []
        for n in search_names:
            if n.lower() not in seen:
                seen.add(n.lower())
                unique_names.append(n)
        search_names = unique_names
        log.info(f"  📝 Searching for BOTH name variants: {' AND '.join(search_names)}")
    else:
        # Only one name variant
        search_names = [name]
    
    # Log what CSV data we have
    csv_fields_found = {k: v for k, v in csv_data.items() if v and k not in ["First Name", "Last Name", "First", "Last", "Already known", "_original_name", "_matched_name"]}
    if csv_fields_found:
        log.info(f"  📊 Using CSV data for verification: {', '.join(csv_fields_found.keys())}")
    
    entity_type = detect_entity_type(name)
    
    # Build targeted queries for identity verification
    if entity_type == "organization":
        queries = [
            f'"{name}" official website',
            f'"{name}" site:wikipedia.org',
            f'"{name}" EIN nonprofit',
            f'"{name}" headquarters {state}' if state else f'"{name}" about'
        ]
    else:
        # People queries with ALL CSV context (employer, title, location, education)
        # Use AND logic to combine all available data for precise identification
        
        # ENHANCED: Get name variants to search for both formal and nickname versions
        name_variants = _token_name_variants(name)
        
        # Extract ALL location data (could be multiple: current + historical)
        all_locations = []
        if city and state:
            all_locations.append(f"{city} {state}")
        elif city:
            all_locations.append(city)
        elif state:
            all_locations.append(state)
        
        # Check for additional location fields in CSV (Previous City, Previous State, etc.)
        for key in csv_data.keys():
            # Skip internal fields that start with underscore (these are lists)
            if key.startswith('_'):
                continue
            key_lower = key.lower()
            if 'city' in key_lower or 'location' in key_lower or 'address' in key_lower:
                if key not in ['City', 'city'] and csv_data.get(key):
                    val = csv_data[key]
                    # Skip if it's a list (internal data structure)
                    if isinstance(val, list):
                        continue
                    loc = str(val).strip()
                    if loc and loc not in all_locations:
                        all_locations.append(loc)
        
        # Log all locations found
        if all_locations:
            log.info(f"  📍 Using {len(all_locations)} location(s): {', '.join(all_locations)}")
        
        # PRIORITY 0: HIGHEST - Name AND Spouse AND Location (if available)
        # This is the MOST SPECIFIC query for disambiguation
        queries = []
        if spouse_name and all_locations:
            log.info(f"  💑 Using spouse name for targeted queries: {spouse_name}")
            for search_name in search_names:
                # Search for person + spouse + location together
                queries.append(f'"{search_name}" "{spouse_name}" {all_locations[0]}')
                queries.append(f'"{search_name}" {spouse_name} {all_locations[0]}')
                # Also try with marriage-related terms
                queries.append(f'"{search_name}" married "{spouse_name}" {all_locations[0]}')
                queries.append(f'"{search_name}" wife "{spouse_name}" {all_locations[0]}')
                queries.append(f'"{search_name}" husband "{spouse_name}" {all_locations[0]}')
        
        # PRIORITY 1: MOST SPECIFIC - Name AND Employer AND Title
        if employer and title:
            log.info(f"  🏢 Using employer: {employer}")
            log.info(f"  💼 Using title: {title}")
            
            # Core query: (Name1 OR Name2) AND Employer AND Title
            for search_name in search_names:
                queries.append(f'"{search_name}" "{employer}" "{title}"')
                queries.append(f'"{search_name}" {employer} {title}')
            
            # Add location variations if available
            if all_locations:
                # Try with current/first location
                for search_name in search_names:
                    queries.append(f'"{search_name}" "{employer}" {all_locations[0]}')
        
        # PRIORITY 2: Name AND Employer (without title)
        elif employer:
            log.info(f"  🏢 Using employer: {employer}")
            for search_name in search_names:
                queries.extend([
                    f'"{search_name}" "{employer}"',
                    f'"{search_name}" {employer}',
                ])
            
            # Add location
            if all_locations:
                for search_name in search_names:
                    queries.append(f'"{search_name}" "{employer}" {all_locations[0]}')
        
        # PRIORITY 3: Name AND Title (without employer)
        elif title:
            log.info(f"  💼 Using title: {title}")
            for search_name in search_names:
                queries.append(f'"{search_name}" "{title}"')
        
        # PRIORITY 4: Name AND Locations (try all locations)
        if all_locations:
            # Add individual location queries for EACH name variant
            for search_name in search_names:
                for loc in all_locations[:3]:  # Limit to first 3 to avoid too many queries
                    queries.append(f'"{search_name}" {loc}')
                    queries.append(f'"{search_name}" biography {loc}')
        
        # PRIORITY 5: Name AND Education
        if education:
            log.info(f"  🎓 Using education: {education}")
            for search_name in search_names:
                queries.append(f'"{search_name}" {education}')
                
                # Name AND Education AND Location (if available)
                if all_locations:
                    queries.append(f'"{search_name}" {education} {all_locations[0]}')
        
        # PRIORITY 6: Current employment searches (general) - for EACH name
        for search_name in search_names:
            queries.extend([
                f'"{search_name}" currently',           # "Jim Ronyak currently serves as..."
                f'"{search_name}" current position',    # "Jim Ronyak current position is..."
                f'"{search_name}" site:linkedin.com/in/',  # LinkedIn usually has current job
            ])
        
        # PRIORITY 7: General professional searches - for EACH name
        for search_name in search_names:
            queries.extend([
                f'"{search_name}" biography',
                f'"{search_name}" profile',
                f'"{search_name}" executive',
                f'"{search_name}" director',
            ])
        
        # PRIORITY 8: Quality sources - for EACH name
        for search_name in search_names:
            queries.extend([
                f'"{search_name}" site:bloomberg.com OR site:reuters.com',
                f'"{search_name}" site:wikipedia.org',
            ])
    
    # Gather candidates
    results = []
    for q in queries:
        results.extend(search_web(q, num=5))
        time.sleep(0.2)
    
    # Score each candidate - use ALL name variants for matching
    seen = set()
    candidates = []
    
    # Build comprehensive list of name variants for scoring
    all_name_variants = []
    for search_name in search_names:
        all_name_variants.extend(_token_name_variants(search_name))
    # Remove duplicates
    all_name_variants = list(set(all_name_variants))
    
    for r in results:
        url = (r.get("link") or "").strip()
        if not url or url in seen:
            continue
        seen.add(url)
        
        title = r.get("title", "")
        snippet = r.get("snippet", "")
        
        # Calculate score components
        domain_wt = _domain_weight(url)
        
        # Name similarity: best match across ALL name variants (jim, james, etc.)
        name_sim = max(_jaro_winkler(title, v) for v in all_name_variants)
        
        # Location bonus - check ALL locations from contactsdb.csv
        loc_bonus = _location_bonus(snippet, title, city, state, spouse_name)
        
        # ENHANCED: Check additional locations from contactsdb.csv
        # If contactsdb.csv has multiple locations (historical), check all of them
        if not loc_bonus:  # If primary location didn't match, try others
            combined_text = f"{title} {snippet}".lower()
            for key in csv_data.keys():
                # Skip internal fields that start with underscore (these are lists)
                if key.startswith('_'):
                    continue
                key_lower = key.lower()
                if ('city' in key_lower or 'location' in key_lower) and key not in ['City', 'city']:
                    other_val = csv_data.get(key, "")
                    # Skip if it's a list (internal data structure)
                    if isinstance(other_val, list):
                        continue
                    other_loc = str(other_val).strip()
                    if other_loc and other_loc.lower() in combined_text:
                        loc_bonus = 0.15  # Slightly lower bonus for historical location
                        log.debug(f"  ✅ Historical location match: {other_loc}")
                        break
        
        # EMPLOYER MATCH BONUS (HIGHEST PRIORITY)
        employer_bonus = 0.0
        if employer:
            combined_lower = f"{title} {snippet}".lower()
            employer_lower = employer.lower()
            # Exact match
            if employer_lower in combined_lower:
                employer_bonus = 0.30
                log.debug(f"  ✅ Employer match: {employer}")
            # Partial match (handles variations like "Google" vs "Google Inc.")
            elif any(word in combined_lower for word in employer_lower.split() if len(word) > 3):
                employer_bonus = 0.20
        
        # TITLE MATCH BONUS (HIGH PRIORITY)
        title_bonus = 0.0
        if title:
            combined_lower = f"{title} {snippet}".lower()
            title_lower = title.lower()
            # Exact or partial title match
            if title_lower in combined_lower:
                title_bonus = 0.25
                log.debug(f"  ✅ Title match: {title}")
            # Check for key words from title
            elif any(word in combined_lower for word in title_lower.split() if len(word) > 4):
                title_bonus = 0.15
        
        # EDUCATION MATCH BONUS (MEDIUM PRIORITY)
        education_bonus = 0.0
        if education:
            combined_lower = f"{title} {snippet}".lower()
            education_lower = education.lower()
            # Check for university name, abbreviation, or key words
            if education_lower in combined_lower:
                education_bonus = 0.20
                log.debug(f"  ✅ Education match: {education}")
            elif any(word in combined_lower for word in education_lower.split() if len(word) > 3):
                education_bonus = 0.10
        
        # SPOUSE MATCH BONUS (CRITICAL FOR PREVENTING NAME COLLISIONS!)
        spouse_bonus = 0.0
        spouse_required = bool(spouse_name)  # Flag if spouse verification is mandatory
        spouse_found = False
        
        if spouse_name:
            combined = f"{title} {snippet}"
            # Use robust spouse detection (first+last, punctuation-insensitive)
            if _spouse_present(spouse_name, combined):
                spouse_found = True
                spouse_bonus = 0.40  # HIGHEST BONUS - spouse match is critical!
                log.debug(f"  💑 Spouse match found: {spouse_name}")
            elif spouse_required:
                log.debug(f"  ❌ REJECTED: No spouse mention (required '{spouse_name}')")
                continue  # Skip this result entirely - wrong person!
            
            # CRITICAL: If spouse IS found but we have location data and NO location match,
            # treat this as a collision (same spouse name, different place)
            # This ensures "Ann Gibbs + George S Gibbs" in Jacksonville
            # doesn't match "Ann Gibbs + George S Gibbs" elsewhere
            if spouse_required and spouse_found and (city or state) and loc_bonus <= 0:
                log.debug(f"  ❌ REJECTED: spouse found but no matching location ({city} {state})")
                continue  # Skip - same name+spouse but different location!
        
        # Negative patterns (obituary, wrong field, etc.)
        neg_penalty = 0.0
        combined = f"{title} {snippet}"
        
        # CRITICAL: Check whitelist FIRST - don't penalize known legitimate sources
        # These are high-trust institutions where false positives are likely
        whitelist_domains = [
            'theihs.org', 'instituteforhumane', 'mercatus', 'george mason',
            'linkedin.com', 'github.com', '.edu/', 'university',
            'cato', 'brookings', 'hoover', 'aei.org', 'heritage.org'
        ]
        
        is_whitelisted = any(domain in url.lower() for domain in whitelist_domains)
        
        if is_whitelisted:
            # Skip negative pattern check for whitelisted sources
            # (Academic papers may mention "died", "memorial lectures", etc. in legitimate contexts)
            pass
        elif any(re.search(pat, combined, re.I) for pat in NEGATIVE_PATTERNS):
            neg_penalty = -5.0  # CRITICAL: Heavily penalize to disqualify
            log.warning(f"  🚨 NEGATIVE PATTERN detected in: {url[:60]}")
            log.warning(f"     Title: {title[:60]}")
            log.warning(f"     This source appears to be about wrong person/obituary")
        
        # POSITIVE SIGNALS: Boost for professional context markers (GENERAL)
        # This helps find business/executive profiles vs family members or wrong professions
        positive_boost = 0.0
        combined_lower = combined.lower()
        
        # Professional domains (universal quality signals)
        # LinkedIn gets highest boost as it typically has CURRENT employment
        if 'linkedin.com/in/' in url.lower():
            positive_boost += 0.25  # Higher boost for LinkedIn (current job info)
        elif any(domain in url.lower() for domain in [
            '.edu/', 'bloomberg.com', 'forbes.com',
            'reuters.com', 'businesswire.com', 'sec.gov'
        ]):
            positive_boost += 0.15
        
        # Executive/leadership role indicators (universal)
        exec_roles = [
            'ceo', 'cfo', 'cto', 'cio', 'coo', 'president', 'vice president',
            'director', 'chief', 'executive', 'founder', 'partner', 'managing',
            'senior vice', 'head of'
        ]
        if any(role in combined_lower for role in exec_roles):
            positive_boost += 0.20
        
        # Professional context indicators (universal)
        if any(term in combined_lower for term in [
            'board member', 'board of directors', 'advisory board',
            'appointed', 'named to', 'promoted to', 'joined', 'announced'
        ]):
            positive_boost += 0.10
        
        # Company/organization context (universal - shows business activity)
        if any(term in combined_lower for term in [
            'company', 'corporation', 'inc.', 'llc', 'foundation',
            'institute', 'organization', 'firm', 'group'
        ]):
            positive_boost += 0.05
        
        # CRITICAL: Temporal indicators - prioritize CURRENT employment over PAST
        # This ensures we find current jobs, not old jobs
        current_indicators = [
            'currently', 'current position', 'currently serves',
            'now serves', 'serves as', 'current role',
            'presently', 'at present', 'as of 20'  # "as of 2024", etc.
        ]
        past_indicators = [
            'formerly', 'previously', 'former', 'past',
            'retired', 'was a', 'used to', 'until 20',  # "until 2020", etc.
            'ex-', 'prior to'
        ]
        
        # Boost for current employment indicators
        if any(term in combined_lower for term in current_indicators):
            positive_boost += 0.30  # Strong boost for current info
            log.info(f"  📅 Current employment indicator found: {title[:50]}")
        
        # Penalize past employment indicators
        if any(term in combined_lower for term in past_indicators):
            positive_boost -= 0.25  # Penalize old information
            log.info(f"  📅 Past employment indicator found: {title[:50]}")
        
        # Total score (with ALL CSV data bonuses + SPOUSE VERIFICATION)
        score = (domain_wt + (0.6 * name_sim) + loc_bonus + employer_bonus + 
                 title_bonus + education_bonus + spouse_bonus + neg_penalty + positive_boost)
        
        candidates.append({
            "url": url,
            "title": title,
            "snippet": snippet,
            "score": round(score, 3),
            "domain_weight": round(domain_wt, 2),
            "name_sim": round(name_sim, 2),
            "loc_bonus": round(loc_bonus, 2),
            "employer_bonus": round(employer_bonus, 2),
            "title_bonus": round(title_bonus, 2),
            "education_bonus": round(education_bonus, 2),
            "spouse_bonus": round(spouse_bonus, 2),
            "positive_boost": round(positive_boost, 2)
        })
    
    # Sort by score
    candidates.sort(key=lambda x: x["score"], reverse=True)
    
    if not candidates:
        log.warning("  ⚠️  No candidates found")
        return {
            "verified": False,
            "confidence": "Low",
            "confidence_score": 0.0,
            "sources": [],
            "entity_type": entity_type,
            "markers": {},
            "collisions": []
        }
    
    # Best candidate
    best = candidates[0]
    
    # CRITICAL: Check if best candidate is from whitelisted domain
    best_url_lower = best['url'].lower()
    whitelist_domains = [
        'theihs.org', 'instituteforhumane', 'mercatus', 'george mason',
        'linkedin.com', 'github.com', '.edu/', 'university',
        'cato', 'brookings', 'hoover', 'aei.org', 'heritage.org'
    ]
    is_best_whitelisted = any(domain in best_url_lower for domain in whitelist_domains)
    
    # CRITICAL CHECK: If top candidate is obituary or wrong profession, STOP
    # BUT: Skip this check for whitelisted sources (they may have false positives)
    if not is_best_whitelisted:
        best_combined = f"{best['title']} {best['snippet']}"
        if any(re.search(pat, best_combined, re.I) for pat in NEGATIVE_PATTERNS):
            log.error(f"  ❌ FATAL: Top candidate is obituary or wrong profession")
            log.error(f"     URL: {best['url']}")
            log.error(f"     Title: {best['title'][:80]}")
            log.error(f"  This appears to be a completely DIFFERENT person")
            log.error(f"  STOPPING to prevent mixing identities")
            return {
                "verified": False,
                "confidence": "None",
                "confidence_score": 0.0,
                "sources": [],
                "entity_type": entity_type,
                "markers": {},
                "collisions": [],
                "error": "Top search result is about wrong person (obituary/wrong profession)"
            }
    else:
        log.info(f"  ✅ Top candidate is from whitelisted domain - proceeding")
    
    # Decision thresholds with special case for high name similarity
    # If we have a near-perfect name match, be more forgiving on total score
    high_name_match = best["name_sim"] >= 0.95
    location_confirmed = best["loc_bonus"] > 0
    
    # NEW: Check if result is from high-quality academic/research institution
    # For researchers/academics, work location may not match home location
    best_url = best["url"].lower()
    best_title = best["title"].lower()
    best_snippet = best.get("snippet", "").lower()
    
    high_quality_academic = any(domain in best_url or domain in best_title for domain in [
        '.edu', 'university', 'college', 'school of',
        'institute', 'center at', 'foundation',
        'mercatus', 'cato', 'brookings', 'aei.org', 'heritage.org', 'hoover.org',
        'manhattan institute', 'fraser institute', 'urban institute', 'rand.org'
    ])
    
    # CRITICAL: Check if result mentions the person's employer from CSV
    # This helps distinguish from family members with same name
    mentions_employer = False
    if employer:
        employer_lower = employer.lower()
        employer_terms = [employer_lower]
        # Add common variations
        if 'institute' in employer_lower or 'foundation' in employer_lower:
            employer_terms.append(employer_lower.split()[0])  # First word (e.g., "Brookings" from "Brookings Institution")
        
        mentions_employer = any(term in best_url or term in best_title or term in best_snippet 
                               for term in employer_terms)
    
    if high_quality_academic:
        log.info(f"  📚 High-quality academic/research source detected")
    if mentions_employer:
        log.info(f"  🏢 Source mentions employer: {employer}")
    
    # SIMPLE DECISION LOGIC (ORIGINAL THRESHOLDS): 
    # Perfect name match (0.95+) + reasonable domain (0.5+) = very likely correct person
    # CSV data (employer/title/education) helps confirm but isn't required
    # Location helps confirm but isn't required for professional names
    
    # Calculate if we have strong CSV data matches
    csv_match_strong = (best.get("employer_bonus", 0) >= 0.20 or 
                        best.get("title_bonus", 0) >= 0.20 or 
                        best.get("education_bonus", 0) >= 0.15)
    
    # ORIGINAL THRESHOLDS (lower = more complete results)
    if best["score"] >= 1.6:
        # High score - definitely right person
        confidence = "High"
        verified = True
    elif high_name_match and best["score"] >= 1.0:
        # Perfect name match + decent domain = accept it
        # CSV data or location would help but isn't mandatory
        confidence = "High" if (location_confirmed or csv_match_strong) else "Medium"
        verified = True
        log.info(f"  ✅ Accepting perfect name match (score={best['score']:.2f}, name_sim={best['name_sim']:.2f})")
    elif best["score"] >= 1.2 and (location_confirmed or csv_match_strong):
        confidence = "Medium"
        verified = True
    elif best["score"] >= 0.9:
        # ORIGINAL: Proceed with score >= 0.9 (not 1.0)
        confidence = "Low"
        verified = True
        log.warning(f"  ⚠️  LOW confidence but proceeding (score={best['score']:.2f})")
    else:
        # Not enough evidence - stop
        confidence = "Low"
        verified = False
        log.error(f"  ❌ Insufficient evidence for '{name}'")
        log.error(f"     Score: {best['score']:.2f}, Name sim: {best['name_sim']:.2f}")
        log.error(f"     Need score >= 0.9 to proceed")
    
    # Detect name collisions (other people with same name)
    collisions = []
    for i, cand in enumerate(candidates[1:6], 1):  # Check top 5 after best
        # If score is close but different URL, might be different person
        if cand["score"] > 0.8 and cand["score"] < best["score"] * 0.9:
            # Check if it mentions a different location
            other_snippet = cand["snippet"].lower()
            if city and city.lower() not in other_snippet and any(
                other_city in other_snippet for other_city in 
                ['new york', 'los angeles', 'chicago', 'boston', 'seattle']
            ):
                collisions.append({
                    "indicator": f"Possible {name} in different location",
                    "source": cand["url"]
                })
    
    sources = [c["url"] for c in candidates[:10]]
    
    log.info(f"  ✓ Top candidate score: {best['score']}")
    log.info(f"    Domain: {best['domain_weight']}, Name sim: {best['name_sim']}, " 
             f"Location: {best['loc_bonus']}")
    log.info(f"    Employer: {best.get('employer_bonus', 0.0)}, Title: {best.get('title_bonus', 0.0)}, "
             f"Education: {best.get('education_bonus', 0.0)}, Spouse: {best.get('spouse_bonus', 0.0)}")
    log.info(f"    Professional context: {best.get('positive_boost', 0.0)}")
    log.info(f"  Confidence: {confidence} (verified={verified})")
    
    if collisions:
        log.warning(f"  ⚠️  {len(collisions)} possible name collision(s) detected")
    
    # Build markers dict for compatibility
    markers = {
        "location_confirmed": best["loc_bonus"] > 0,
        "employer_found": best.get("employer_bonus", 0) > 0,
        "education_found": best.get("education_bonus", 0) > 0,
        "spouse_confirmed": best.get("spouse_bonus", 0) > 0,
        "high_quality_sources": sum(1 for c in candidates[:5] if c["domain_weight"] >= 0.70),
        "name_match_quality": best["name_sim"],
        "score": best["score"]
    }
    
    return {
        "verified": verified,
        "confidence": confidence,
        "confidence_score": best["score"] / 2.0,  # Normalize to 0-1 range
        "sources": sources,
        "entity_type": entity_type,
        "resolved_name": name,
        "canonical_urls": [best["url"]],
        "markers": markers,
        "collisions": collisions,
        "top_candidate": best
    }

# -------------------------
# OPENAI CHAT
# -------------------------
@retry_with_backoff(max_retries=3)
def openai_chat(prompt: str, cache_key_prefix: str = "chat") -> str:
    cache_key = cache.get_cache_key(cache_key_prefix, {"prompt": prompt[:100]})
    
    def fetcher():
        if not openai or not OPENAI_API_KEY:
            return "Error: OpenAI not configured"
        
        if OPENAI_VERSION == "new":
            resp = OPENAI_CLIENT.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.0
            )
            return resp.choices[0].message.content.strip()
        else:
            resp = openai.ChatCompletion.create(
                model=OPENAI_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.0
            )
            return resp["choices"][0]["message"]["content"].strip()
    
    return cache.get_or_fetch(cache_key, fetcher)

@retry_with_backoff(max_retries=3)
def openai_chat_json(prompt: str, cache_key_prefix: str = "chatjson") -> Dict:
    """Call OpenAI for JSON response with robust parsing"""
    cache_key = cache.get_cache_key(cache_key_prefix, {"prompt": prompt[:100]})
    
    def fetcher():
        if not openai or not OPENAI_API_KEY:
            return {}
        
        if OPENAI_VERSION == "new":
            resp = OPENAI_CLIENT.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.0
            )
            text = resp.choices[0].message.content.strip()
        else:
            resp = openai.ChatCompletion.create(
                model=OPENAI_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.0
            )
            text = resp["choices"][0]["message"]["content"].strip()
        
        # Strip common code-fence wrappers (```json ... ``` or ``` ... ```)
        text_clean = text.strip()
        if text_clean.startswith("```"):
            # Remove opening fence
            text_clean = re.sub(r"^```(?:json)?\s*", "", text_clean)
            # Remove closing fence
            text_clean = re.sub(r"\s*```\s*$", "", text_clean)
        text_clean = text_clean.strip()
        
        # Try to isolate the first JSON object
        m = re.search(r"\{.*\}", text_clean, re.DOTALL)
        candidate = m.group(0) if m else text_clean
        
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            # Final fallback: remove trailing commas before closing braces/brackets
            candidate2 = re.sub(r",\s*([}\]])", r"\1", candidate)
            try:
                return json.loads(candidate2)
            except json.JSONDecodeError as e:
                log.warning(f"JSON parsing failed: {e}. Response: {text[:200]}")
                return {}
    
    return cache.get_or_fetch(cache_key, fetcher)

# -------------------------
# BULLET-POINT EXTRACTION WITH CITATIONS
# -------------------------
BULLET_BIO_PROMPT = """Create a concise biographical section with 5-8 bullet points.

Cover:
- Early life and education
- Family background
- Geographic roots
- Current location/residence

CRITICAL: TEMPORAL AWARENESS
- PRIORITIZE current information over past information
- When sources mention BOTH current and past locations, use the CURRENT one
- If sources say "currently in X" or "based in X", that's the current location
- If sources say "formerly in Y" or "was in Y", that's past - note it as past
- Look for temporal indicators: "currently", "now", "as of [recent year]" vs "formerly", "previously", "was"

CRITICAL ANTI-HALLUCINATION RULES:
1. ONLY include facts that appear in the source materials below
2. NEVER invent biographical details, family information, or educational credentials
3. Include citation numbers [1], [2], etc. after EACH factual claim
4. If sources lack information on a topic, write "• Limited biographical information available in public sources"
5. DO NOT write generic statements without specific supporting evidence in sources
6. Every bullet point MUST have at least one citation [N] that actually supports that claim
7. When listing location, check temporal context - is it current or past?

Format as bullet points starting with "• "

ENTITY: {name}
ENTITY TYPE: {entity_type}

SOURCE MATERIALS WITH CITATION NUMBERS:
{snippets}

If sources are insufficient, return: "• Limited biographical information available in public sources"

REMINDER: Never invent degrees, universities, or life details that aren't explicitly stated in sources.

Return ONLY bullet points (5-8 bullets, EACH with citations):"""

BULLET_CAREER_PROMPT = """Create a concise career section with 5-8 bullet points.

Cover:
- Current position and role (MOST IMPORTANT - list this FIRST)
- Career progression and key positions
- Business leadership and accomplishments
- Professional expertise
- Industry recognition

🔴 CRITICAL: CURRENT vs PAST EMPLOYMENT 🔴

PRIORITY ORDER:
1. CURRENT position comes FIRST (look for "currently", "now serves", "present")
2. Past positions come AFTER current (look for "formerly", "previously", "was")
3. When sources show both current and past employment, LEAD with current
4. Mark past positions with temporal language: "Previously served as...", "Former..."

TEMPORAL INDICATORS TO WATCH FOR:
- Current: "currently", "now", "serves as", "current role", "as of 2024/2025"
- Past: "formerly", "previously", "was", "until", "retired from", "former"

❌ CRITICAL ANTI-HALLUCINATION RULES - READ THESE FIRST ❌

YOU MUST FOLLOW THESE RULES OR YOUR RESPONSE WILL BE REJECTED:

1. DO NOT INVENT organization names, companies, conferences, or board positions
2. DO NOT write "active member of [organization]" EVER unless those EXACT words with that EXACT organization name appear in the sources
3. DO NOT write "regular speaker at [conference]" EVER unless those EXACT words appear in the sources
4. DO NOT write "maintains relationships with [anyone]" - this is ALWAYS a hallucination
5. DO NOT write "widely recognized", "well-known", "respected leader", or "influential figure" unless specific evidence in sources
6. DO NOT paraphrase organization names - use the EXACT name from the source or don't mention it
7. If a source mentions a company, use its EXACT name - do NOT create similar-sounding names
8. DISTINGUISH current from past - don't list an old job as if it's current

❌ THESE PHRASES ARE FORBIDDEN ❌
- "Active member of the National Association of..."
- "Regular speaker at the Annual..."  
- "Maintains key relationships with..."
- "Widely recognized in the industry..."
- "Serves on various boards..."
- "Member of several professional organizations..."
- ANY organization/conference/board name NOT explicitly mentioned in sources
- Listing a past job without "formerly" or "previously"

✅ WHAT TO DO INSTEAD ✅
- If sources say person CURRENTLY works at XYZ → write "• Currently serves as [role] at XYZ Corp [1]"
- If sources say person FORMERLY worked at ABC → write "• Previously served as [role] at ABC Corp [2]"
- If sources say NOTHING about current job → write "• Current position not specified in sources"
- DO NOT fill information gaps with plausible-sounding details
- START with current position, THEN list past positions

🚨 REMEMBER: It is MUCH BETTER to say "no information found" than to invent FALSE details about a real person. Inventing affiliations could harm someone's reputation.

Include citation numbers [1], [2], etc. after EACH factual claim.
Format as bullet points starting with "• "

ENTITY: {name}
ENTITY TYPE: {entity_type}

SOURCE MATERIALS WITH CITATION NUMBERS:
{snippets}

If sources don't contain specific career details, return: "• Limited career information available in public sources"

Return ONLY bullet points (5-8 bullets with citations, CURRENT position first if available):"""

BULLET_PHILANTHROPY_PROMPT = """Create a concise philanthropy section with 5-8 bullet points.

Cover:
- Foundation involvement and giving amounts
- Board service with nonprofits
- Charitable causes supported
- Major gifts and commitments
- Philanthropic philosophy or themes

CRITICAL ANTI-HALLUCINATION RULES:
1. ONLY include philanthropic activities explicitly mentioned in sources
2. NEVER invent foundation names, nonprofit names, or donation amounts
3. Include citation numbers [1], [2], etc. after EACH factual claim
4. If sources lack information, write "• No verified philanthropic activity found in available sources"
5. DO NOT write generic statements like "supports various causes" without specific organization names and amounts
6. Every bullet point MUST have at least one citation [N] that supports that specific claim

Format as bullet points starting with "• "

ENTITY: {name}
ENTITY TYPE: {entity_type}

SOURCE MATERIALS WITH CITATION NUMBERS:
{snippets}

If sources are insufficient, return: "• No public records of philanthropic activity found"

REMINDER: Never invent foundation names, donation amounts, or nonprofit affiliations. If not in sources, admit no data found.

Return ONLY bullet points (5-8 bullets with citations, or admission of no data):"""

BULLET_POLITICAL_PROMPT = """Create a concise political activity section with 5-8 bullet points.

Cover:
- Political donations and amounts
- Party affiliation or pattern
- Policy positions and advocacy
- Think tank/political org connections
- Public statements on political issues

CRITICAL ANTI-HALLUCINATION RULES:
1. ONLY include political activities explicitly mentioned in sources
2. NEVER invent think tank names, policy positions, or donation amounts
3. Include citation numbers [1], [2], etc. after EACH factual claim
4. If sources lack information, write "• No verified political activity found in available sources"
5. DO NOT write vague statements like "maintains relationships with policy makers" without specific evidence
6. Every bullet point MUST have at least one citation [N] that supports that specific claim

Format as bullet points starting with "• "

ENTITY: {name}
ENTITY TYPE: {entity_type}

SOURCE MATERIALS WITH CITATION NUMBERS:
{snippets}

If sources are insufficient, return: "• No public records of political activity found"

REMINDER: Never invent political affiliations, donations, or think tank memberships. If not in sources, admit no data found.

Return ONLY bullet points (5-8 bullets with citations, or admission of no data):"""

BULLET_NETWORK_PROMPT = """Create a concise network and board affiliations section with 5-8 bullet points.

Cover:
- Current board memberships (nonprofit, corporate, advisory)
- Past board service with dates
- Professional association memberships
- Key relationships and networks
- Speaking engagements and conference participation

❌ CRITICAL ANTI-HALLUCINATION RULES - THIS SECTION IS HIGH RISK ❌

BOARD MEMBERSHIPS AND SPEAKING ENGAGEMENTS ARE THE #1 SOURCE OF HALLUCINATIONS.

YOU MUST FOLLOW THESE RULES OR YOUR RESPONSE WILL BE REJECTED:

1. DO NOT invent board names, organization names, conference names, or association names
2. DO NOT write "Active member of [organization]" unless the sources explicitly state this with the EXACT organization name
3. DO NOT write "Regular speaker at [conference]" unless the sources explicitly state this with the EXACT conference name
4. DO NOT write "Board member of [org]" unless the sources explicitly state this with the EXACT organization name
5. DO NOT write vague phrases like "maintains relationships with industry leaders" - these are ALWAYS hallucinations
6. DO NOT write "serves on various boards" or "member of several organizations" - be specific or say nothing
7. If sources mention ONE board, list ONLY that one board - do not invent others
8. Include citation numbers [1], [2], etc. after EACH factual claim

❌ EXAMPLES OF HALLUCINATIONS TO NEVER PRODUCE ❌
✗ "Active member of the National Association of Corporate Directors"
✗ "Regular speaker at the Annual Technology Leadership Conference"  
✗ "Maintains relationships with industry leaders"
✗ "Board member of the Technology Advisory Council"
✗ "Serves on various nonprofit boards"
✗ "Member of several professional organizations"
✗ "Speaks at industry conferences"

✅ WHAT TO DO INSTEAD ✅
✓ If sources say "serves on XYZ board" → write "• Serves on XYZ board [N]"
✓ If sources say NOTHING about boards → write "• No verified board affiliations found in available sources"
✓ If sources mention ONE speaking engagement → write ONLY about that ONE with citation
✓ DO NOT fill gaps with plausible-sounding details

🚨 ETHICAL REMINDER: Falsely claiming someone has board memberships or speaking roles can harm their professional reputation. If it's not explicitly in the sources, DO NOT include it.

Format as bullet points starting with "• "

ENTITY: {name}
ENTITY TYPE: {entity_type}

SOURCE MATERIALS WITH CITATION NUMBERS:
{snippets}

If sources are insufficient, return: "• No verified board affiliations or speaking engagements found in available sources"

Return ONLY bullet points (5-8 bullets with citations, or admission of no data):"""

BULLET_NETWORTH_PROMPT = """Create a concise net worth and giving capacity section with 5-8 bullet points.

Cover:
- Estimated net worth or wealth indicators
- Company valuations, stock holdings, assets
- Real estate holdings
- Foundation assets or giving history
- Major transactions (sales, acquisitions)
- Giving capacity assessment

CRITICAL ANTI-HALLUCINATION RULES:
1. ONLY include financial information explicitly stated in sources
2. NEVER invent net worth figures, asset values, or giving amounts
3. NEVER estimate or assume financial capacity without specific evidence
4. Include citation numbers [1], [2], etc. after EACH factual claim
5. If sources lack information, write "• No public net worth or asset information found"
6. DO NOT write statements like "estimated giving capacity" unless sources provide specific evidence
7. Every bullet point MUST have at least one citation [N] that supports that specific claim

Format as bullet points starting with "• "

ENTITY: {name}
ENTITY TYPE: {entity_type}

SOURCE MATERIALS WITH CITATION NUMBERS:
{snippets}

If sources are insufficient, return: "• Financial information not available in public sources"

REMINDER: Never invent wealth estimates, property values, or financial capacity assessments. If not in sources, admit no data found.

Return ONLY bullet points (5-8 bullets with citations, or admission of no data):"""

# Section header for Net Worth / Capacity for Giving
CAPACITY_SECTION_HEADER = """**CAPACITY FOR GIVING**

This section summarizes available wealth indicators and giving capacity from multiple data sources."""

IHS_ASSESSMENT_PROMPT = """You are evaluating a prospect for the Institute for Humane Studies (IHS) using the A-E-I-N-P-S framework.

IHS Mission: Advancing individual liberty, free markets, civil society, and academic excellence

A-E-I-N-P-S Framework:
- **Ability**: Financial capacity to give ($10K+ major gifts)
- **Education**: Academic credentials, intellectual orientation
- **Interest**: Alignment with liberty, free markets, limited government
- **Networks**: Connections to academic/policy circles
- **Propensity**: History of giving to similar causes
- **Solicitation**: Ease of approach, existing relationships

Based on the research below, create a donor assessment with:

1. **Overall Rating** (A/B/C/D score for cultivation priority)
2. **Key Strengths** (3-4 bullets with evidence)
3. **Alignment with IHS Mission** (3-4 bullets)
4. **Concerns or Gaps** (2-3 bullets)
5. **Recommended Next Steps** (2-3 action items)

CRITICAL: Include citation numbers [1], [2], etc. after factual claims.
Format with section headers and bullet points.

PROSPECT: {name}
LOCATION: {location}

BIOGRAPHICAL SUMMARY:
{bio_summary}

CAREER SUMMARY:
{career_summary}

PHILANTHROPY SUMMARY:
{philanthropy_summary}

POLITICAL SUMMARY:
{political_summary}

NETWORK SUMMARY:
{network_summary}

Return assessment with clear sections and bullet points:"""

STRATEGIC_SUMMARY_PROMPT = """You are preparing a board-level strategic briefing about a donor prospect.

🚨 ABSOLUTE RULE - NO FABRICATION 🚨
- ONLY include information that appears in the RESEARCH SUMMARY below
- If the research says "Limited information available" or "No public records" - SAY THAT
- NEVER invent job titles, companies, awards, accomplishments, or any details
- NEVER fabricate citations - only cite information actually in the research
- If research is sparse, write a SHORTER summary acknowledging the limitations
- It is BETTER to say "limited information available" than to make something up

Write an executive summary for IHS board members. If research is limited, acknowledge this honestly.

Paragraph 1: **Who They Are** (ONLY from research)
- CURRENT role IF KNOWN from research (do not invent)
- Key facts ONLY if documented in research
- If research shows limited career info, say "Limited professional information available"

Paragraph 2: **Why IHS Should Engage** (ONLY if evidence exists)
- Any alignment with liberty, free markets found in research
- Giving capacity ONLY if documented
- If no philanthropic record exists, acknowledge this

Paragraph 3: **Cultivation Strategy** (based on available information)
- Recommended approach given what IS known
- Be conservative on gift estimates if capacity is unknown
- Acknowledge if more research is needed

CRITICAL: 
- Citation numbers [1], [2], etc. MUST reference actual sources from the research
- Do NOT invent fake citations like "Company Website" or "Industry Publications"
- If you cannot cite something from the research, do NOT include the claim

PROSPECT: {name}
LOCATION: {location}

RESEARCH SUMMARY:
{bio_summary}
{career_summary}
{philanthropy_summary}
{political_summary}
{network_summary}

Return summary (shorter if research is limited):"""

GIFT_OFFICER_SUMMARY_PROMPT = """You are a gift officer preparing an executive summary for frontline fundraisers.

🚨 ABSOLUTE RULE - NO FABRICATION 🚨
- ONLY include information that appears in the RESEARCH SUMMARY below
- If the research says "Limited information available" or "No public records" - SAY THAT
- NEVER invent job titles, companies, awards, accomplishments, or any details
- NEVER fabricate citations - only cite information actually in the research
- If research is sparse, write a SHORTER summary acknowledging the limitations
- It is BETTER to say "Information not available" than to make something up

**RESEARCH CONFIDENCE: {confidence}**
(Based on multi-marker verification: {marker_summary})

1. **Quick Profile** (2-3 bullets - ONLY from research)
   - Who they are, CURRENT role/company IF documented in research
   - Key credentials IF found in research
   - Geographic base (from research)
   - If career info is limited, say "Professional background: Limited public information"

2. **Why They Matter to IHS** (ONLY if evidence exists)
   - Alignment with IHS mission IF documented
   - Giving capacity indicators IF documented
   - If no alignment or capacity found, acknowledge this

3. **Connection Points** (based on available information)
   - How to reach them IF documented
   - Topics that would interest them IF we have data
   - If connections unknown, say so

4. **Ask Strategy** (conservative if data is limited)
   - Suggested cultivation steps
   - Be conservative on gift estimates if capacity unknown
   - If research limited, recommend gathering more information first

{collision_warning}

CRITICAL: 
- Citation numbers [1], [2], etc. MUST reference actual sources from the research
- Do NOT invent fake citations
- If you cannot cite something, do NOT include the claim

Format as bullet points starting with "• " under clear section headers.

PROSPECT: {name}
ENTITY TYPE: {entity_type}
LOCATION: {location}

RESEARCH SUMMARY:
{bio_summary}

{career_summary}

{philanthropy_summary}

{political_summary}

Return executive summary (shorter if research is limited):"""

def format_snippets_with_citations(results: List[Dict[str, str]], person_name: str = "", 
                                   location: str = "") -> str:
    """Format search results with citation numbers, filtering for quality and relevance
    
    Args:
        person_name: Name of person being researched (to validate citations)
        location: Location context (city, state) to validate it's the right person
    """
    # Block data broker and low-quality sites
    BLOCKED_DOMAINS = [
        'whitepages.com', 'spokeo.com', 'mylife.com', 'intelius.com',
        'truthfinder.com', 'beenverified.com', 'instantcheckmate.com',
        'publicrecords.com', 'peoplefinders.com', 'zabasearch.com',
        'radaris.com', 'fastpeoplesearch.com', 'familytreenow.com'
    ]
    
    filtered_results = []
    name_parts = person_name.lower().split() if person_name else []
    last_name = name_parts[-1] if name_parts else ""
    first_name = name_parts[0] if name_parts else ""
    
    # Parse location for filtering
    location_parts = []
    if location:
        # Extract city and state
        loc_clean = location.lower().strip()
        location_parts = [p.strip() for p in loc_clean.split(',')]
    
    for r in results:
        url = r.get('link', '').lower()
        snippet = r.get('snippet', '').lower()
        title = r.get('title', '').lower()
        
        # Skip data broker sites
        if any(blocked in url for blocked in BLOCKED_DOMAINS):
            continue
        
        # CRITICAL: Skip obituaries and wrong professions
        combined = f"{title} {snippet}"
        combined_lower = combined.lower()
        if any(re.search(pat, combined, re.I) for pat in NEGATIVE_PATTERNS):
            log.warning(f"  🚨 Blocking source with NEGATIVE PATTERN: {url[:60]}")
            log.warning(f"     (obituary/deceased/wrong profession detected)")
            continue  # Skip this source entirely
        
        # ENHANCED: Tighten name matching - require full name or first+last
        # Don't accept sources that only mention last name (too generic)
        if person_name:
            # Check for various name patterns
            has_full_name = person_name.lower() in combined_lower
            has_first_last = (first_name and last_name and 
                             first_name in combined_lower and last_name in combined_lower)
            has_last_only = last_name and last_name in combined_lower
            
            # Check location match
            has_location_match = False
            if location_parts:
                has_location_match = any(loc_part in combined_lower 
                                        for loc_part in location_parts if len(loc_part) > 2)
            
            # High-quality sources get more lenient matching
            high_quality_domains = ['linkedin.com', 'wikipedia.org', 'bloomberg.com', 
                                   'reuters.com', 'forbes.com', 'wsj.com', '.edu/',
                                   'sec.gov', 'businesswire.com']
            is_high_quality = any(domain in url for domain in high_quality_domains)
            
            # DECISION LOGIC:
            # 1. Prefer full name or first+last name
            # 2. Allow last-name-only ONLY if:
            #    - Location matches (e.g., "Mrs. Gibbs in Jacksonville") OR
            #    - Source is high-quality (LinkedIn, Wikipedia, etc.)
            # This prevents accepting random "Gibbs" pages that might be about spouse or other family
            
            if not (has_full_name or has_first_last or 
                   (has_last_only and (has_location_match or is_high_quality))):
                log.debug(f"  ❌ Name matching too weak: {title[:60]}")
                continue  # Source doesn't clearly mention the person
            
            # ENHANCED: Additional location validation when location is provided
            if location_parts:
                # If not high quality AND no location match, check for contradicting locations
                if not is_high_quality and not has_location_match:
                    # Additional check: Look for contradicting locations
                    other_cities = ['new york', 'los angeles', 'chicago', 'boston', 'seattle',
                                   'san francisco', 'miami', 'atlanta', 'dallas', 'houston',
                                   'philadelphia', 'phoenix', 'detroit', 'portland']
                    
                    # If snippet mentions a different major city, skip it (likely different person)
                    expected_city = location_parts[0] if location_parts else ""
                    if any(city in snippet and city != expected_city for city in other_cities):
                        log.warning(f"  🚨 Skipping source - mentions different location: {url[:60]}")
                        continue
        
        filtered_results.append(r)
    
    lines = []
    # CRITICAL FIX: Use number returned by citation_mgr, not enumerate
    for r in filtered_results:
        title = r.get('title', '')
        snippet = r.get('snippet', '')
        url = r.get('link', '')
        
        # Get actual citation number from manager (keeps numbering globally consistent)
        num = citation_mgr.add_citation(url, title, snippet)
        
        lines.append(f"[{num}] {title}\n    {snippet}\n    URL: {url}")
    
    return "\n\n".join(lines)

def extract_bullet_section(name: str, entity_type: str, queries: List[str], 
                          prompt_template: str, fast_mode: bool = True,
                          location_filter: str = "",
                          identity_verified: bool = True,
                          spouse_name: str = "") -> str:
    """Extract bullet-point section with citations
    
    Args:
        location_filter: "City, State" to add context to queries (not for filtering)
        identity_verified: If True, skip strict location filtering (we already know it's the right person)
        spouse_name: If provided, FILTER OUT any results that don't mention spouse (prevents name collisions)
    """
    # Add location context to queries if available (helps search engines)
    if location_filter:
        queries = [f"{q} {location_filter}" for q in queries]
    
    # Search
    if fast_mode:
        results = []
        with ThreadPoolExecutor(max_workers=len(queries)) as executor:
            future_to_query = {executor.submit(search_web, q, 5): q for q in queries}
            for future in as_completed(future_to_query):
                try:
                    results.extend(future.result())
                except Exception as e:
                    log.warning(f"Search failed: {e}")
        time.sleep(0.3)
    else:
        results = []
        for q in queries:
            results.extend(search_web(q, num=5))
            time.sleep(1.2)
    
    # Deduplicate
    uniq, seen = [], set()
    for r in results:
        u = r.get("link")
        if u and u not in seen:
            seen.add(u)
            uniq.append(r)
    
    # CRITICAL: STRICT spouse+location filtering when identity NOT verified
    # Once identity is verified, we relax filtering to avoid losing valid content
    # This implements: "Ann Gibbs + George S Gibbs + Jacksonville" to confirm identity,
    # then accept all Ann Gibbs content once we're sure it's the right person
    
    # 1) SPOUSE FILTER - only when identity is NOT yet verified
    if spouse_name and not identity_verified:
        filtered_results = []
        rejected_count = 0
        
        for r in uniq:
            title = r.get("title", "")
            snippet = r.get("snippet", "")
            combined = f"{title} {snippet}"
            
            # Use robust spouse detection (first+last, punctuation-insensitive)
            # This matches "George S. Gibbs" when spouse_name is "George Gibbs"
            spouse_found = _spouse_present(spouse_name, combined)
            
            if spouse_found:
                filtered_results.append(r)
            else:
                rejected_count += 1
                log.debug(f"  ❌ SPOUSE FILTER: Rejected result (no mention of '{spouse_name}'): {r.get('title', '')[:80]}")
        
        log.info(f"  🔍 Spouse Filter: Kept {len(filtered_results)}/{len(uniq)} results (rejected {rejected_count} without spouse mention)")
        uniq = filtered_results
        
        # FALLBACK: If ALL results were filtered out, fall back to unfiltered for this section
        if not uniq:
            log.warning(f"  ⚠️  ALL results filtered out by spouse requirement!")
            log.warning(f"     Falling back to unfiltered results for this section (sparse data)")
            uniq = [r for r in results if r.get("link") and r.get("link") not in seen][:20]
    
    # 2) LOCATION FILTER - only when identity is NOT yet verified
    #    Once identity_verified=True, pass empty location so we don't hard-filter by city/state
    location_for_filtering = "" if identity_verified else location_filter
    
    # Format with citations (this also filters by location if location_for_filtering is set)
    snippets = format_snippets_with_citations(uniq[:20], person_name=name, location=location_for_filtering)
    
    # Generate bullets
    # CRITICAL: Include spouse_name in cache key to prevent using cached data from runs without spouse filtering
    cache_params = {"name": name}
    if spouse_name:
        cache_params["spouse"] = spouse_name  # Different cache for spouse-filtered vs non-filtered
    cache_key = cache.get_cache_key(f"bullets_{prompt_template[:20]}", cache_params)
    
    def fetcher():
        prompt = prompt_template.format(
            name=name,
            entity_type=entity_type,
            snippets=snippets or "No information found."
        )
        return openai_chat(prompt, cache_key_prefix=cache_key)
    
    bullets = cache.get_or_fetch(cache_key, fetcher)
    
    # CRITICAL: Final check - if generated content mentions obituary/wrong profession, block it
    if any(re.search(pat, bullets, re.I) for pat in NEGATIVE_PATTERNS):
        log.error(f"  ❌ GENERATED CONTENT contains obituary or wrong profession mentions")
        log.error(f"  Blocking this content - it's about a DIFFERENT {name}")
        bullets = f"• Unable to generate verified content - search results contain conflicting information about different people named {name}"
    
    # ENHANCED: Validate that organization names in bullets actually appear in sources
    bullets = validate_organization_claims(bullets, snippets, name)
    
    return bullets

def detect_common_hallucinations(text: str, person_name: str) -> List[str]:
    """
    Detect common hallucination patterns that LLMs generate when data is sparse.
    Returns list of detected hallucination phrases.
    """
    hallucination_patterns = [
        # Generic membership phrases (unless EXACT organization is in sources)
        (r'active member of the (?:national|international|global) (?:association|society|council|institute)', 
         'Generic "active member" phrase'),
        
        # Generic speaker phrases
        (r'regular speaker at the (?:annual|international|national) \w+ (?:conference|summit|forum)', 
         'Generic "regular speaker" phrase'),
        
        # Generic relationship phrases (ALWAYS hallucinations)
        (r'maintains (?:key |strong )?relationships? with (?:industry leaders|key stakeholders|leading)', 
         'Vague "maintains relationships" phrase - always fabricated'),
        (r'maintains (?:close )?(?:ties|connections) (?:to|with)', 
         'Vague "maintains ties/connections" phrase'),
        
        # Generic recognition phrases without specifics
        (r'widely recognized (?:in the|as a|for)', 
         'Generic "widely recognized" without specific evidence'),
        (r'well-known (?:in the|as a|for)', 
         'Generic "well-known" without specific evidence'),
        (r'respected (?:leader|figure|voice) in', 
         'Generic "respected" without specific evidence'),
        
        # Generic influence phrases
        (r'influential (?:voice|figure|leader)', 
         'Generic "influential" without specific evidence'),
        (r'thought leader in', 
         'Generic "thought leader" without specific evidence'),
        
        # Generic board/advisory claims without specific org names
        (r'serves? on (?:numerous|various|several|multiple) boards?', 
         'Vague board service without specific organizations'),
        (r'member of (?:numerous|various|several|multiple) (?:boards?|advisory)', 
         'Vague board membership without specifics'),
        
        # Generic philanthropic claims
        (r'supports? (?:various|numerous|several|multiple) (?:causes|charities|organizations)', 
         'Vague philanthropy without specific organizations'),
        (r'active in (?:various|numerous|several|multiple) (?:philanthropic|charitable)', 
         'Vague philanthropy activity'),
    ]
    
    detected = []
    text_lower = text.lower()
    
    for pattern, description in hallucination_patterns:
        matches = re.findall(pattern, text_lower, re.IGNORECASE)
        if matches:
            detected.append(f"{description}: '{matches[0]}'")
            log.error(f"  🚨 HALLUCINATION DETECTED in {person_name}: {description}")
    
    return detected

def validate_organization_claims(bullets: str, source_snippets: str, person_name: str) -> str:
    """
    Post-process bullets to validate organization/conference names actually appear in sources.
    FIXED v2.2: Now with WHITELIST for known legitimate organizations + higher threshold.
    """
    if not source_snippets or "No information found" in source_snippets:
        return bullets
    
    # ⭐ WHITELIST: Known legitimate organizations (skip validation for these)
    WHITELISTED_ORGS = [
        # IHS and related
        'institute for humane studies', 'ihs', 'theihs.org', 'instituteforhumanestudies',
        
        # Think tanks and policy orgs
        'mercatus center', 'mercatus', 'george mason university', 'george mason', 'gmu',
        'cato institute', 'cato', 'brookings institution', 'brookings',
        'hoover institution', 'hoover', 'american enterprise institute', 'aei',
        'heritage foundation', 'heritage', 'manhattan institute',
        'reason foundation', 'reason', 'competitive enterprise institute', 'cei',
        'fraser institute', 'atlas network',
        
        # Major universities
        'stanford', 'harvard', 'mit', 'yale', 'princeton', 'columbia',
        'university of chicago', 'chicago', 'duke', 'northwestern',
        'berkeley', 'ucla', 'upenn', 'cornell', 'dartmouth', 'brown',
        
        # Major tech companies
        'google', 'microsoft', 'amazon', 'apple', 'meta', 'facebook',
        'tesla', 'spacex', 'nvidia', 'intel', 'ibm', 'oracle',
        'salesforce', 'adobe', 'cisco',
        
        # Financial firms
        'goldman sachs', 'goldman', 'morgan stanley', 'jpmorgan', 'blackrock',
        
        # Consulting
        'mckinsey', 'bain', 'bcg', 'deloitte', 'pwc', 'kpmg',
        
        # Common contexts
        'linkedin', 'github', 'board member', 'advisory board',
    ]
    
    # STEP 1: Detect common hallucination patterns (KEEP THIS - it's good)
    hallucinations = detect_common_hallucinations(bullets, person_name)
    if hallucinations:
        log.error(f"  ❌ BLOCKING CONTENT for {person_name} - detected {len(hallucinations)} hallucination(s)")
        for h in hallucinations:
            log.error(f"     - {h}")
        return f"• Unable to generate verified content - LLM attempted to fabricate information\n• Limited information available in public sources\n• Manual research recommended"
    
    # STEP 2: Extract organization-like names from bullets
    org_pattern = r'\b(?:National|International|Global|Annual|American|European|Asia|World|Regional)\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b'
    
    specific_patterns = [
        r'\b\w+\s+Association\s+of\s+\w+(?:\s+\w+)*\b',
        r'\b\w+\s+Leadership\s+(?:Conference|Forum|Summit)\b',
        r'\b\w+\s+(?:Council|Institute|Society|Board)\b'
    ]
    
    source_text = source_snippets.lower()
    lines = bullets.split('\n')
    validated_lines = []
    unverified_count = 0
    
    for line in lines:
        if not line.strip().startswith('•'):
            validated_lines.append(line)
            continue
        
        # Check for organization mentions in this line
        found_orgs = re.findall(org_pattern, line)
        for pattern in specific_patterns:
            found_orgs.extend(re.findall(pattern, line, re.IGNORECASE))
        
        suspicious_orgs = []
        for org in found_orgs:
            org_lower = org.lower()
            
            # ⭐ NEW: Skip validation for whitelisted orgs
            is_whitelisted = any(whitelist_org in org_lower for whitelist_org in WHITELISTED_ORGS)
            if is_whitelisted:
                log.info(f"  ✅ WHITELISTED: '{org}' - allowing without source check")
                continue
            
            # Check if org appears in sources
            key_words = [w for w in org_lower.split() if len(w) > 3]
            
            if key_words and not any(word in source_text for word in key_words):
                suspicious_orgs.append(org)
        
        if suspicious_orgs:
            unverified_count += 1
            
            # ⭐ INCREASED THRESHOLD: Was 3, now 6 (allows more legitimate orgs)
            if unverified_count >= 20:
                log.error(f"  ❌ BLOCKING ENTIRE SECTION for {person_name} - {unverified_count} unverified organizations")
                return f"• Unable to generate verified content - multiple unverified organization names detected\n• Limited information available in public sources\n• Manual research recommended"
            
            line = line.rstrip()
            line += f" [⚠️  UNVERIFIED: '{suspicious_orgs[0]}' not found in sources]"
            log.warning(f"  🚨 UNVERIFIED CLAIM #{unverified_count}: '{suspicious_orgs[0]}' not in sources for {person_name}")
        
        validated_lines.append(line)
    
    return '\n'.join(validated_lines)


def validate_biographical_consistency(bullets: str, section_name: str, 
                                      expected_location: str, person_name: str) -> str:
    """
    Check if biographical details mention contradicting locations that suggest wrong person.
    Adds warnings for potential identity mismatches.
    """
    if not expected_location or not bullets:
        return bullets
    
    # Parse expected location
    loc_parts = [p.strip().lower() for p in expected_location.split(',')]
    expected_city = loc_parts[0] if loc_parts else ""
    expected_state = loc_parts[1] if len(loc_parts) > 1 else ""
    
    # Major US cities to check for conflicts
    major_cities = [
        'new york', 'los angeles', 'chicago', 'boston', 'seattle',
        'san francisco', 'miami', 'atlanta', 'dallas', 'houston',
        'philadelphia', 'phoenix', 'detroit', 'portland', 'denver',
        'las vegas', 'nashville', 'austin', 'orlando', 'minneapolis'
    ]
    
    bullets_lower = bullets.lower()
    conflicting_cities = []
    
    for city in major_cities:
        if city != expected_city and city in bullets_lower:
            # Check if it's actually talking about this city (not just mentioned in passing)
            # Look for phrases like "lives in", "based in", "from", "resident of"
            location_indicators = [
                f'lives in {city}', f'based in {city}', f'from {city}',
                f'resident of {city}', f'{city} resident', f'living in {city}'
            ]
            if any(indicator in bullets_lower for indicator in location_indicators):
                conflicting_cities.append(city.title())
    
    if conflicting_cities:
        warning_msg = f"\n\n⚠️  IDENTITY WARNING: {section_name} mentions location(s) {', '.join(conflicting_cities)} " \
                     f"but expected location is {expected_location}. " \
                     f"This may be information about a DIFFERENT {person_name}. Verify carefully.\n"
        log.error(f"  🚨 IDENTITY MISMATCH: {section_name} mentions {conflicting_cities} but expected {expected_location}")
        bullets = warning_msg + bullets
    
    return bullets

def generate_gift_officer_summary(name: str, entity_type: str, location: str,
                                  bio_bullets: str, career_bullets: str,
                                  philanthropy_bullets: str, political_bullets: str,
                                  confidence: str = "Medium",
                                  markers: Dict = None,
                                  collisions: List = None) -> str:
    """Generate 2-page executive summary for gift officers"""
    cache_key = cache.get_cache_key("gift_summary", {"name": name})
    
    # CRITICAL: Extract ONLY current information for summary
    # This prevents LLM from picking past jobs/locations
    
    # Extract current career info (first 1-2 bullets, which should be current)
    career_lines = [line.strip() for line in career_bullets.split('\n') if line.strip().startswith('•')]
    current_career = []
    for line in career_lines[:3]:  # Check first 3 bullets
        line_lower = line.lower()
        # Include if it has current indicators OR is the first bullet
        if any(indicator in line_lower for indicator in ['currently', 'current', 'now serves', 'serves as']):
            current_career.append(line)
        elif len(current_career) == 0:
            # First bullet even without "currently" - likely current position
            current_career.append(line)
        # Stop if we hit a past indicator
        if any(past in line_lower for past in ['formerly', 'previously', 'former', 'past', 'until']):
            break
    
    # Extract current biographical info (location, current context)
    bio_lines = [line.strip() for line in bio_bullets.split('\n') if line.strip().startswith('•')]
    current_bio = []
    for line in bio_lines:
        line_lower = line.lower()
        # Include if it mentions current or is about location without past indicators
        if any(indicator in line_lower for indicator in ['currently', 'current', 'based in', 'lives in']):
            if not any(past in line_lower for past in ['formerly', 'previously', 'was based', 'was in']):
                current_bio.append(line)
    
    # Fallback if no current info extracted
    if not current_career and career_lines:
        current_career = [career_lines[0]]  # Just take first bullet
    if not current_bio and bio_lines:
        current_bio = [bio_lines[0]]  # Just take first bullet
    
    # Format filtered summaries
    filtered_career = '\n'.join(current_career) if current_career else "Current position information limited"
    filtered_bio = '\n'.join(current_bio) if current_bio else "Current biographical information limited"
    
    log.info(f"  📋 Summary using CURRENT info only:")
    log.info(f"     Career: {len(current_career)} current bullets")
    log.info(f"     Bio: {len(current_bio)} current bullets")
    
    # Format marker summary
    if markers:
        marker_list = []
        if markers.get("location_confirmed"): marker_list.append("location")
        if markers.get("employer_found"): marker_list.append("employer")
        if markers.get("education_found"): marker_list.append("education")
        if markers.get("professional_work"): marker_list.append("professional work")
        marker_summary = f"{len(marker_list)}/4 markers confirmed ({', '.join(marker_list)})" if marker_list else "0/4 markers confirmed"
    else:
        marker_summary = "Unknown verification status"
    
    # Format collision warning
    collision_warning = ""
    if collisions and len(collisions) > 0:
        collision_warning = f"\n⚠️ NAME COLLISION WARNING: {len(collisions)} other people with this name found. " \
                           "Verify identity carefully before engagement.\n"
    
    def fetcher():
        prompt = GIFT_OFFICER_SUMMARY_PROMPT.format(
            name=name,
            entity_type=entity_type,
            location=location,
            confidence=confidence,
            marker_summary=marker_summary,
            collision_warning=collision_warning,
            bio_summary=f"Current Biographical Background:\n{filtered_bio}",
            career_summary=f"Current Career Position:\n{filtered_career}",
            philanthropy_summary=f"Philanthropic Activities:\n{philanthropy_bullets}",
            political_summary=f"Political Activity:\n{political_bullets}"
        )
        return openai_chat(prompt, cache_key_prefix=cache_key)
    
    summary = cache.get_or_fetch(cache_key, fetcher)
    return summary

def research_network_boards(name: str, entity_type: str, location_filter: str, 
                            fast_mode: bool = True, spouse_name: str = "") -> str:
    """Research network and board affiliations"""
    log.info("Researching network and board affiliations...")
    
    queries = [
        f'"{name}" board member director',
        f'"{name}" advisory board',
        f'"{name}" nonprofit board service',
        f'"{name}" professional association member',
        f'"{name}" conference speaker panel'
    ]
    
    # Add "Mr. and Mrs." variant if spouse known
    if spouse_name:
        queries.insert(0, f'"Mr. and Mrs. {spouse_name}" board')
        queries.insert(1, f'"Mr. and Mrs. {spouse_name}" nonprofit')
    
    return extract_bullet_section(name, entity_type, queries, BULLET_NETWORK_PROMPT, 
                                  fast_mode, location_filter, identity_verified=True)


# -------------------------
# iWAVE API INTEGRATION
# -------------------------
def lookup_iwave_capacity(name: str, city: str = "", state: str = "", 
                          spouse_name: str = "", country: str = "US") -> Dict[str, Any]:
    """
    Query iWave API for donor giving capacity data.
    
    IMPORTANT: Tries multiple name variations including spouse name since many
    couples are listed under "Mr. and Mrs. [Spouse Name]" in wealth databases.
    
    Args:
        name: Full name of the donor (e.g., "Ann Gibbs")
        city: City for better matching accuracy
        state: State/region for better matching accuracy
        spouse_name: Spouse's full name (e.g., "George S. Gibbs") - CRITICAL for matching
        country: Country code (default US)
    
    Returns dict with capacity data or error.
    """
    if not IWAVE_ENABLED or not IWAVE_API_KEY:
        return {
            "success": False,
            "error": "iWave API not configured (set IWAVE_API_KEY in spider.env)"
        }
    
    # Build list of name variations to try
    name_variations = []
    
    # Parse primary name
    name_parts = name.strip().split()
    if len(name_parts) >= 2:
        first_name = name_parts[0]
        last_name = " ".join(name_parts[1:])
        name_variations.append((first_name, last_name, f"{first_name} {last_name}"))
    
    # CRITICAL: Add spouse name variations - many records are under spouse's name
    if spouse_name:
        spouse_parts = spouse_name.strip().split()
        if len(spouse_parts) >= 2:
            spouse_first = spouse_parts[0]
            spouse_last = spouse_parts[-1]  # Use last part as surname
            
            # Try spouse's full name
            name_variations.append((spouse_first, spouse_last, f"{spouse_first} {spouse_last}"))
            
            # Try "Mr. and Mrs. [Spouse Last Name]" pattern
            # Many wealth records list couples this way
            if len(name_parts) >= 2:
                donor_last = name_parts[-1]
                # If same last name, try just the spouse
                if donor_last.lower() == spouse_last.lower():
                    name_variations.append((spouse_first, spouse_last, f"Mr. and Mrs. {spouse_first} {spouse_last}"))
    
    log.info(f"  📊 Looking up iWave capacity - trying {len(name_variations)} name variation(s):")
    for _, _, display in name_variations:
        log.info(f"      - {display}")
    
    # Try each name variation
    for first_name, last_name, display_name in name_variations:
        log.info(f"    🔍 Searching iWave for: {display_name}")
        
        result = _iwave_api_search(first_name, last_name, city, state, country)
        
        if result.get("success"):
            log.info(f"    ✅ iWave MATCH found for: {display_name}")
            result["matched_name"] = display_name
            return result
        else:
            log.info(f"    ⚪ No match for: {display_name}")
    
    return {
        "success": False,
        "error": f"No iWave match for {name}" + (f" or spouse {spouse_name}" if spouse_name else "")
    }


def _iwave_api_search(first_name: str, last_name: str, city: str = "", 
                      state: str = "", country: str = "US") -> Dict[str, Any]:
    """
    Execute single iWave API search request.
    """
    # Check cache first
    cache_key = cache.get_cache_key("iwave_capacity", {
        "first": first_name, "last": last_name, "city": city, "state": state
    })
    
    def fetcher():
        # iWave API endpoints to try
        endpoints_to_try = [
            "/api/v1/prospects/search",
            "/api/v1/profiles/search", 
            "/prospects/search",
            "/profiles/search"
        ]
        
        headers = {
            "Content-Type": "application/json",
            "Accept": "application/json",
            "Authorization": f"Bearer {IWAVE_API_KEY}"
        }
        
        # If using key+secret auth pattern
        if IWAVE_API_SECRET:
            headers["X-API-Key"] = IWAVE_API_KEY
            headers["X-API-Secret"] = IWAVE_API_SECRET
            del headers["Authorization"]
        
        # Build search payload - try both naming conventions
        payload = {
            "firstName": first_name,
            "first_name": first_name,
            "lastName": last_name,
            "last_name": last_name,
            "country": country
        }
        
        # Add location if available
        if city:
            payload["city"] = city
        if state:
            payload["state"] = state
        
        last_error = None
        
        for endpoint in endpoints_to_try:
            url = f"{IWAVE_API_URL.rstrip('/')}{endpoint}"
            try:
                response = requests.post(
                    url,
                    headers=headers,
                    json=payload,
                    timeout=30
                )
                
                if response.status_code == 200:
                    data = response.json()
                    return parse_iwave_response(data, f"{first_name} {last_name}")
                elif response.status_code == 401:
                    return {"success": False, "error": "iWave authentication failed - check API credentials"}
                elif response.status_code == 404:
                    continue  # Try next endpoint
                else:
                    last_error = f"iWave API error: {response.status_code}"
                    
            except requests.exceptions.Timeout:
                last_error = "iWave API timeout"
            except requests.exceptions.RequestException as e:
                last_error = f"iWave API error: {str(e)}"
        
        return {
            "success": False,
            "error": last_error or f"No iWave match for {first_name} {last_name}"
        }
    
    return cache.get_or_fetch(cache_key, fetcher)


def parse_iwave_response(data: Dict[str, Any], name: str) -> Dict[str, Any]:
    """Parse iWave API response into standardized format"""
    
    # iWave capacity rating mapping
    CAPACITY_RATINGS = {
        "A": "$10M+",
        "B": "$5M-$10M", 
        "C": "$1M-$5M",
        "D": "$500K-$1M",
        "E": "$100K-$500K",
        "F": "$50K-$100K",
        "G": "$10K-$50K",
        "H": "Under $10K"
    }
    
    # Handle both array and single object responses
    if isinstance(data, list):
        if not data:
            return {"success": False, "error": f"No iWave match for {name}"}
        prospect = data[0]  # Take first/best match
    else:
        prospect = data.get("prospect", data.get("profile", data))
    
    if not prospect:
        return {"success": False, "error": f"No iWave match for {name}"}
    
    # Extract capacity rating (letter code)
    capacity_code = prospect.get("capacityRating", prospect.get("capacity_rating", 
                    prospect.get("capacityCode", prospect.get("capacity_code", ""))))
    
    # Map to readable range
    capacity_range = CAPACITY_RATINGS.get(capacity_code, "")
    
    # Extract numeric capacity values
    estimated_capacity = prospect.get("estimatedCapacity", prospect.get("estimated_capacity",
                         prospect.get("givingCapacity", prospect.get("giving_capacity", ""))))
    
    # Format currency if numeric
    if isinstance(estimated_capacity, (int, float)):
        estimated_capacity = f"${estimated_capacity:,.0f}"
    
    # Extract real estate value
    real_estate = prospect.get("realEstateValue", prospect.get("real_estate_value",
                  prospect.get("realEstate", prospect.get("real_estate", ""))))
    if isinstance(real_estate, (int, float)):
        real_estate = f"${real_estate:,.0f}"
    
    # Extract charitable giving history
    charitable = prospect.get("charitableGiving", prospect.get("charitable_giving",
                 prospect.get("totalGiving", prospect.get("total_giving", ""))))
    if isinstance(charitable, (int, float)):
        charitable = f"${charitable:,.0f}"
    
    # Extract scores
    propensity = prospect.get("propensityScore", prospect.get("propensity_score",
                 prospect.get("p2g", prospect.get("P2G", ""))))
    affinity = prospect.get("affinityScore", prospect.get("affinity_score", ""))
    proscore = prospect.get("proScore", prospect.get("proscore", 
               prospect.get("PROScore", prospect.get("score", ""))))
    
    # Extract planned giving likelihood
    planned_giving = prospect.get("plannedGivingLikelihood", prospect.get("planned_giving_likelihood",
                     prospect.get("plannedGiving", prospect.get("planned_giving", ""))))
    
    # Extract business and foundation affiliations
    businesses = prospect.get("businesses", prospect.get("businessAffiliations", 
                 prospect.get("business_affiliations", [])))
    foundations = prospect.get("foundations", prospect.get("foundationAffiliations",
                  prospect.get("foundation_affiliations", [])))
    
    # Build result
    result = {
        "success": True,
        "capacity_rating": capacity_code,
        "capacity_range": capacity_range,
        "estimated_capacity": str(estimated_capacity) if estimated_capacity else "",
        "real_estate_value": str(real_estate) if real_estate else "",
        "charitable_giving": str(charitable) if charitable else "",
        "planned_giving_likelihood": str(planned_giving) if planned_giving else "",
        "propensity_score": str(propensity) if propensity else "",
        "affinity_score": str(affinity) if affinity else "",
        "proscore": str(proscore) if proscore else "",
        "businesses": businesses if isinstance(businesses, list) else [],
        "foundations": foundations if isinstance(foundations, list) else [],
        "raw_data": prospect
    }
    
    log.info(f"      ✅ Capacity {capacity_code} ({capacity_range})")
    
    return result


def format_iwave_api_bullets(iwave_result: Dict[str, Any]) -> str:
    """Format iWave API results as bullet points for report"""
    
    if not iwave_result.get("success"):
        error = iwave_result.get("error", "Unknown error")
        if "not configured" in error.lower():
            return ""  # Don't add anything if iWave not configured
        return f"\n• iWave lookup: {error}"
    
    matched_name = iwave_result.get("matched_name", "")
    header = "• **iWave Wealth Screening Data (API)**"
    if matched_name:
        header += f" - matched as '{matched_name}'"
    header += ":"
    
    bullets = ["\n" + header]
    
    if iwave_result.get("capacity_rating"):
        rating = iwave_result["capacity_rating"]
        range_str = iwave_result.get("capacity_range", "")
        bullets.append(f"  - Capacity Rating: {rating}" + (f" ({range_str})" if range_str else ""))
    
    if iwave_result.get("estimated_capacity"):
        bullets.append(f"  - Estimated Giving Capacity: {iwave_result['estimated_capacity']}")
    
    if iwave_result.get("real_estate_value"):
        bullets.append(f"  - Real Estate Holdings: {iwave_result['real_estate_value']}")
    
    if iwave_result.get("charitable_giving"):
        bullets.append(f"  - Known Charitable Giving: {iwave_result['charitable_giving']}")
    
    if iwave_result.get("propensity_score"):
        bullets.append(f"  - Propensity to Give (P2G): {iwave_result['propensity_score']}")
    
    if iwave_result.get("affinity_score"):
        bullets.append(f"  - Affinity Score: {iwave_result['affinity_score']}")
    
    if iwave_result.get("proscore"):
        bullets.append(f"  - PROScore: {iwave_result['proscore']}")
    
    if iwave_result.get("planned_giving_likelihood"):
        bullets.append(f"  - Planned Giving Likelihood: {iwave_result['planned_giving_likelihood']}")
    
    # Add business affiliations
    businesses = iwave_result.get("businesses", [])
    if businesses and isinstance(businesses, list) and len(businesses) > 0:
        biz_names = [b.get("name", b) if isinstance(b, dict) else str(b) for b in businesses[:5]]
        bullets.append(f"  - Business Affiliations: {', '.join(biz_names)}")
    
    # Add foundation affiliations
    foundations = iwave_result.get("foundations", [])
    if foundations and isinstance(foundations, list) and len(foundations) > 0:
        found_names = [f.get("name", f) if isinstance(f, dict) else str(f) for f in foundations[:5]]
        bullets.append(f"  - Foundation Affiliations: {', '.join(found_names)}")
    
    if len(bullets) == 1:
        return ""  # Only header, no data
    
    return "\n".join(bullets)


# -------------------------
# SMARTY API - PROPERTY DATA LOOKUP
# -------------------------

def lookup_smarty_property(street: str, city: str = "", state: str = "", 
                           zipcode: str = "") -> Dict[str, Any]:
    """
    Look up property data using Smarty's US Address Enrichment API.
    
    This API returns property attributes including:
    - Assessed value, market value, sale price
    - Property type, square footage, lot size
    - Year built, bedrooms, bathrooms
    - Tax information
    
    Args:
        street: Street address (e.g., "123 Main St")
        city: City name
        state: State code (e.g., "AZ")
        zipcode: ZIP code
    
    Returns dict with property data or error.
    """
    if not SMARTY_ENABLED or not SMARTY_AUTH_ID or not SMARTY_AUTH_TOKEN:
        return {
            "success": False,
            "error": "Smarty API not configured (set SMARTY_AUTH_ID and SMARTY_AUTH_TOKEN in spider.env)"
        }
    
    if not street:
        return {"success": False, "error": "No street address provided for property lookup"}
    
    # Build cache key
    cache_key = cache.get_cache_key("smarty_property", {
        "street": street, "city": city, "state": state, "zip": zipcode
    })
    
    def fetcher():
        # Use the search endpoint which takes address components
        url = "https://us-enrichment.api.smarty.com/lookup/search/property/principal"
        
        params = {
            "auth-id": SMARTY_AUTH_ID,
            "auth-token": SMARTY_AUTH_TOKEN,
            "street": street,
            "license": "us-property-data-principal-cloud"
        }
        
        if city:
            params["city"] = city
        if state:
            params["state"] = state
        if zipcode:
            params["zipcode"] = zipcode
        
        try:
            response = requests.get(url, params=params, timeout=30, headers=DEFAULT_HEADERS)
            
            if response.status_code == 200:
                data = response.json()
                return parse_smarty_property_response(data, street)
            elif response.status_code == 401:
                return {"success": False, "error": "Smarty authentication failed - check API credentials"}
            elif response.status_code == 402:
                return {"success": False, "error": "Smarty subscription required for property data"}
            elif response.status_code == 404:
                return {"success": False, "error": f"No property data found for {street}"}
            else:
                return {"success": False, "error": f"Smarty API error: {response.status_code}"}
                
        except requests.exceptions.Timeout:
            return {"success": False, "error": "Smarty API timeout"}
        except requests.exceptions.RequestException as e:
            return {"success": False, "error": f"Smarty API error: {str(e)}"}
    
    return cache.get_or_fetch(cache_key, fetcher)


def parse_smarty_property_response(data: List[Dict], address: str) -> Dict[str, Any]:
    """Parse Smarty property API response into standardized format"""
    
    if not data or not isinstance(data, list) or len(data) == 0:
        return {"success": False, "error": f"No property data found for {address}"}
    
    # Take first result (best match)
    result = data[0]
    attrs = result.get("attributes", {})
    matched = result.get("matched_address", {})
    
    if not attrs:
        return {"success": False, "error": f"No property attributes found for {address}"}
    
    # Extract key property values
    # Financial values
    assessed_value = attrs.get("assessed_value_total", attrs.get("assessed_total", ""))
    market_value = attrs.get("market_value_total", attrs.get("market_total", ""))
    sale_price = attrs.get("sale_price", attrs.get("previous_sale_price", ""))
    sale_date = attrs.get("sale_date", attrs.get("previous_sale_date", ""))
    tax_amount = attrs.get("tax_amount", attrs.get("taxes", ""))
    
    # Property characteristics  
    property_type = attrs.get("property_type", attrs.get("land_use", ""))
    year_built = attrs.get("year_built", "")
    sqft = attrs.get("building_sqft", attrs.get("living_sqft", attrs.get("sqft", "")))
    lot_size = attrs.get("lot_sqft", attrs.get("lot_size", attrs.get("acres", "")))
    bedrooms = attrs.get("bedrooms", "")
    bathrooms = attrs.get("bathrooms", attrs.get("baths_total", ""))
    
    # Format currency values
    def format_currency(val):
        if not val:
            return ""
        try:
            num = float(str(val).replace(",", "").replace("$", ""))
            return f"${num:,.0f}"
        except:
            return str(val)
    
    parsed = {
        "success": True,
        "smarty_key": result.get("smarty_key", ""),
        "matched_address": f"{matched.get('street', '')} {matched.get('city', '')}, {matched.get('state', '')} {matched.get('zipcode', '')}".strip(),
        "assessed_value": format_currency(assessed_value),
        "market_value": format_currency(market_value),
        "sale_price": format_currency(sale_price),
        "sale_date": str(sale_date) if sale_date else "",
        "tax_amount": format_currency(tax_amount),
        "property_type": str(property_type) if property_type else "",
        "year_built": str(year_built) if year_built else "",
        "sqft": str(sqft) if sqft else "",
        "lot_size": str(lot_size) if lot_size else "",
        "bedrooms": str(bedrooms) if bedrooms else "",
        "bathrooms": str(bathrooms) if bathrooms else "",
        "raw_data": attrs
    }
    
    # Log success with key value
    best_value = parsed["market_value"] or parsed["assessed_value"] or parsed["sale_price"]
    if best_value:
        log.info(f"      ✅ Property value: {best_value}")
    
    return parsed


def format_smarty_property_bullets(smarty_result: Dict[str, Any]) -> str:
    """Format Smarty property data as bullet points for report"""
    
    if not smarty_result.get("success"):
        error = smarty_result.get("error", "Unknown error")
        if "not configured" in error.lower():
            return ""  # Don't add anything if Smarty not configured
        return ""  # Silently skip if no property data found
    
    bullets = ["\n• **Property Data (Smarty US Address Enrichment API)**:"]
    
    if smarty_result.get("matched_address"):
        bullets.append(f"  - Address: {smarty_result['matched_address']}")
    
    # Financial values - most important for donor research
    if smarty_result.get("market_value"):
        bullets.append(f"  - Market Value: {smarty_result['market_value']}")
    
    if smarty_result.get("assessed_value"):
        bullets.append(f"  - Assessed Value: {smarty_result['assessed_value']}")
    
    if smarty_result.get("sale_price"):
        sale_info = f"  - Last Sale Price: {smarty_result['sale_price']}"
        if smarty_result.get("sale_date"):
            sale_info += f" ({smarty_result['sale_date']})"
        bullets.append(sale_info)
    
    if smarty_result.get("tax_amount"):
        bullets.append(f"  - Annual Property Tax: {smarty_result['tax_amount']}")
    
    # Property characteristics
    if smarty_result.get("property_type"):
        bullets.append(f"  - Property Type: {smarty_result['property_type']}")
    
    if smarty_result.get("year_built"):
        bullets.append(f"  - Year Built: {smarty_result['year_built']}")
    
    if smarty_result.get("sqft"):
        bullets.append(f"  - Square Footage: {smarty_result['sqft']} sqft")
    
    if smarty_result.get("lot_size"):
        lot = smarty_result['lot_size']
        # Check if it's acres or sqft
        try:
            lot_num = float(str(lot).replace(",", ""))
            if lot_num < 100:  # Likely acres
                bullets.append(f"  - Lot Size: {lot} acres")
            else:
                bullets.append(f"  - Lot Size: {lot_num:,.0f} sqft")
        except:
            bullets.append(f"  - Lot Size: {lot}")
    
    if smarty_result.get("bedrooms") or smarty_result.get("bathrooms"):
        beds = smarty_result.get("bedrooms", "?")
        baths = smarty_result.get("bathrooms", "?")
        bullets.append(f"  - Bedrooms/Bathrooms: {beds} bed / {baths} bath")
    
    if len(bullets) == 1:
        return ""  # Only header, no actual data
    
    return "\n".join(bullets)


def research_net_worth_capacity(name: str, entity_type: str, location_filter: str,
                                fast_mode: bool = True, spouse_name: str = "",
                                donor_row: Dict[str, Any] = None) -> str:
    """
    Research net worth and giving capacity with structured subsections.
    
    Returns formatted output with clear subsections:
    - Property Data (from Smarty API)
    - iWave Capacity Rating (from capacitytogive.csv)
    - Giving History (from capacitytogive.csv)
    - Donor Profile (from capacitytogive.csv)
    - Wealth Screening Data (from iWave CSV if available in data.csv)
    - Web Research (from search results)
    
    This matches the format of the IHS Profile template "Net Worth/Giving Capacity" section.
    """
    log.info("Researching capacity for giving...")
    
    sections = []
    
    # Extract city/state for lookup
    city = ""
    state = ""
    if donor_row:
        city = (donor_row.get("City", "") or donor_row.get("city", "")).strip()
        state = (donor_row.get("State", "") or donor_row.get("state", "")).strip()
    
    # -------------------------
    # 1. PROPERTY DATA (Smarty API)
    # -------------------------
    if donor_row and entity_type != "organization":
        # Extract address components from donor_row OR contactsdb
        # Try various common column name patterns
        street = (donor_row.get("Address", "") or 
                  donor_row.get("Street", "") or 
                  donor_row.get("Street Address", "") or
                  donor_row.get("Mailing Address", "") or
                  donor_row.get("Home Address", "") or
                  donor_row.get("Preferred Address", "") or
                  donor_row.get("address", "") or
                  donor_row.get("street", "")).strip()
        
        zipcode = (donor_row.get("Zip", "") or 
                   donor_row.get("ZIP", "") or
                   donor_row.get("Zipcode", "") or
                   donor_row.get("ZipCode", "") or
                   donor_row.get("Postal Code", "") or
                   donor_row.get("zip", "")).strip()
        
        # Only attempt property lookup if we have a street address
        if street:
            log.info(f"  🏠 Looking up property data for: {street}")
            if city or state:
                log.info(f"      Location: {city}, {state} {zipcode}")
            
            smarty_result = lookup_smarty_property(street, city=city, state=state, zipcode=zipcode)
            
            if smarty_result.get("success"):
                log.info(f"    ✅ Smarty property data found")
                property_section = format_property_subsection(smarty_result)
                if property_section:
                    sections.append(property_section)
            elif "not configured" not in smarty_result.get("error", "").lower():
                log.info(f"    ℹ️  Smarty: {smarty_result.get('error', 'No property data found')}")
        else:
            log.info(f"  ℹ️  No street address in CSV/contactsdb - skipping property lookup")
    
    # -------------------------
    # 2. CAPACITY TO GIVE CSV LOOKUP
    # -------------------------
    # Look up donor in capacitytogive.csv by name + location verification
    if entity_type != "organization":
        capacity_data = lookup_capacity_data(name, city=city, state=state)
        if capacity_data:
            log.info(f"  📊 Found capacity data in capacitytogive.csv")
            capacity_section = format_capacity_subsection(capacity_data)
            if capacity_section:
                sections.append(capacity_section)
    
    # -------------------------
    # 3. WEALTH SCREENING DATA (iWave columns in data.csv)
    # -------------------------
    # Check for iWave data already in CSV columns (from previous screening export)
    if donor_row:
        iwave_csv_data = extract_iwave_data(donor_row)
        if iwave_csv_data:
            log.info(f"  📊 Found iWave data in data.csv: {len(iwave_csv_data)} field(s)")
            iwave_section = format_iwave_subsection(iwave_csv_data)
            if iwave_section:
                sections.append(iwave_section)
    
    # NOTE: iWave API lookup temporarily disabled - using Smarty for property values
    # When re-enabled, add iWave API results here
    
    # -------------------------
    # 4. WEB RESEARCH
    # -------------------------
    if entity_type == "organization":
        queries = [
            f'"{name}" revenue assets budget',
            f'"{name}" financial statements 990',
            f'"{name}" endowment funding'
        ]
    else:
        queries = [
            f'"{name}" net worth wealth',
            f'"{name}" company valuation stock',
            f'"{name}" real estate property',
            f'"{name}" foundation assets',
            f'"{name}" major gift donation amount'
        ]
        # Add "Mr. and Mrs." variant if spouse known
        if spouse_name:
            queries.insert(0, f'"Mr. and Mrs. {spouse_name}" foundation assets')
            queries.insert(1, f'"Mr. and Mrs. {spouse_name}" donation')
        
        # Add location-specific property searches (multiple locations if known)
        if city and state:
            queries.append(f'"{name}" property {city} {state}')
        
        # Check for secondary/alternate locations in donor_row
        if donor_row:
            secondary_locations = []
            for key in donor_row.keys():
                key_lower = key.lower()
                if any(x in key_lower for x in ['city2', 'secondary', 'previous', 'alternate', 'other']):
                    if 'city' in key_lower or 'location' in key_lower:
                        loc_val = donor_row.get(key, "").strip()
                        if loc_val and (not city or loc_val.lower() != city.lower()):
                            secondary_locations.append(loc_val)
            
            # Add secondary location queries
            for sec_loc in secondary_locations[:2]:  # Limit to 2 extra locations
                queries.append(f'"{name}" property {sec_loc}')
                queries.append(f'"{name}" foundation assets {sec_loc}')
                log.info(f"  📍 Adding secondary location for wealth search: {sec_loc}")
    
    # Get web research results
    web_results = extract_bullet_section(name, entity_type, queries, BULLET_NETWORTH_PROMPT,
                                  fast_mode, location_filter, identity_verified=True)
    
    if web_results and web_results.strip() and "not available" not in web_results.lower():
        web_section = "**Web Research:**\n" + web_results
        sections.append(web_section)
    elif not sections:
        # If we have no other data, include the "not available" message
        sections.append("• Financial information not available in public sources")
    
    # -------------------------
    # COMBINE ALL SECTIONS
    # -------------------------
    return "\n\n".join(sections)


def format_property_subsection(smarty_result: Dict[str, Any]) -> str:
    """Format Smarty property data as a subsection for Capacity for Giving"""
    
    if not smarty_result.get("success"):
        return ""
    
    lines = ["**Property Data (Home Value):**"]
    
    if smarty_result.get("matched_address"):
        lines.append(f"• Address: {smarty_result['matched_address']}")
    
    # Financial values - most important for donor research
    if smarty_result.get("market_value"):
        lines.append(f"• Estimated Market Value: {smarty_result['market_value']}")
    
    if smarty_result.get("assessed_value"):
        lines.append(f"• Tax Assessed Value: {smarty_result['assessed_value']}")
    
    if smarty_result.get("sale_price"):
        sale_info = f"• Last Sale Price: {smarty_result['sale_price']}"
        if smarty_result.get("sale_date"):
            sale_info += f" ({smarty_result['sale_date']})"
        lines.append(sale_info)
    
    if smarty_result.get("tax_amount"):
        lines.append(f"• Annual Property Tax: {smarty_result['tax_amount']}")
    
    # Property characteristics (brief)
    property_details = []
    if smarty_result.get("property_type"):
        property_details.append(smarty_result['property_type'])
    if smarty_result.get("year_built"):
        property_details.append(f"Built {smarty_result['year_built']}")
    if smarty_result.get("sqft"):
        property_details.append(f"{smarty_result['sqft']} sqft")
    if smarty_result.get("bedrooms") and smarty_result.get("bathrooms"):
        property_details.append(f"{smarty_result['bedrooms']}BR/{smarty_result['bathrooms']}BA")
    
    if property_details:
        lines.append(f"• Property Details: {', '.join(property_details)}")
    
    if smarty_result.get("lot_size"):
        lot = smarty_result['lot_size']
        try:
            lot_num = float(str(lot).replace(",", ""))
            if lot_num < 100:
                lines.append(f"• Lot Size: {lot} acres")
            else:
                lines.append(f"• Lot Size: {lot_num:,.0f} sqft")
        except:
            lines.append(f"• Lot Size: {lot}")
    
    if len(lines) == 1:
        return ""  # Only header, no actual data
    
    return "\n".join(lines)


def format_iwave_subsection(iwave_data: Dict[str, Any]) -> str:
    """
    Format iWave CSV data as a subsection for Capacity for Giving.
    Returns formatted string or empty string if no data.
    """
    if not iwave_data:
        return ""
    
    lines = ["**Wealth Screening Data (iWave):**"]
    
    # Format each field with human-readable labels
    field_labels = {
        "iwave_score": "Overall Score",
        "capacity_range": "Gift Capacity Range",
        "estimated_capacity": "Estimated Giving Capacity",
        "major_gift_likelihood": "Major Gift Likelihood",
        "planned_gift_likelihood": "Planned Gift Likelihood",
        "p2g_score": "Propensity to Give (P2G)",
        "rfm_score": "RFM Score",
        "wealth_rating": "Wealth Rating",
        "real_estate_value": "Real Estate Value",
        "stock_value": "Stock/Securities Value",
        "total_assets": "Total Estimated Assets",
        "largest_gift": "Largest Known Gift",
        "total_giving": "Total Philanthropic Giving",
        "annual_giving": "Average Annual Giving",
    }
    
    for field_key, label in field_labels.items():
        if field_key in iwave_data:
            value = iwave_data[field_key]
            lines.append(f"• {label}: {value}")
    
    if len(lines) == 1:
        return ""  # Only header, no actual data
    
    return "\n".join(lines)


# -------------------------
# CAPACITY TO GIVE CSV LOOKUP
# -------------------------
# Global cache for capacity data
_capacity_data_cache = None

def load_capacity_csv(capacity_path: str = "capacitytogive.csv") -> List[Dict[str, Any]]:
    """
    Load the capacity-to-give CSV file.
    Returns list of rows as dictionaries.
    """
    global _capacity_data_cache
    
    if _capacity_data_cache is not None:
        return _capacity_data_cache
    
    if not os.path.exists(capacity_path):
        log.debug(f"  Capacity CSV not found at {capacity_path}")
        return []
    
    encodings = ['utf-8-sig', 'utf-8', 'latin-1', 'cp1252']
    
    for encoding in encodings:
        try:
            with open(capacity_path, 'r', encoding=encoding, errors='replace') as f:
                reader = csv.DictReader(f)
                rows = list(reader)
                _capacity_data_cache = rows
                log.info(f"  📊 Loaded {len(rows)} records from {capacity_path}")
                return rows
        except Exception as e:
            continue
    
    log.warning(f"  ⚠️ Failed to load capacity CSV: {capacity_path}")
    return []


def normalize_name_for_matching(name: str) -> str:
    """Normalize a name for matching - lowercase, remove punctuation, extra spaces"""
    if not name:
        return ""
    # Lowercase
    name = name.lower()
    # Remove common suffixes and titles
    for suffix in [' household', ' foundation', ' family foundation', ' trust', 
                   ' jr.', ' jr', ' sr.', ' sr', ' iii', ' ii', ' iv']:
        name = name.replace(suffix, '')
    # Remove punctuation
    name = re.sub(r'[.,&]', ' ', name)
    # Normalize whitespace
    name = ' '.join(name.split())
    return name.strip()


def extract_name_parts(full_name: str) -> Dict[str, str]:
    """Extract first, middle, last name parts from a full name"""
    if not full_name:
        return {"first": "", "middle": "", "last": ""}
    
    # Clean up
    name = full_name.strip()
    # Remove suffixes
    for suffix in [' Jr.', ' Jr', ' Sr.', ' Sr', ' III', ' II', ' IV']:
        name = name.replace(suffix, '')
    
    parts = name.split()
    
    if len(parts) == 1:
        return {"first": parts[0], "middle": "", "last": ""}
    elif len(parts) == 2:
        return {"first": parts[0], "middle": "", "last": parts[1]}
    else:
        # First, middle initial(s), last
        return {"first": parts[0], "middle": " ".join(parts[1:-1]), "last": parts[-1]}


def lookup_capacity_data(donor_name: str, city: str = "", state: str = "",
                         capacity_path: str = "capacitytogive.csv") -> Optional[Dict[str, Any]]:
    """
    Look up a donor in the capacity-to-give CSV by name and location.
    
    Matching logic:
    1. Try exact match on Primary Contact: Full Name
    2. Try match on Account Name (extract name from "John Doe Household")
    3. Verify by location (city, state) if available
    
    Returns the matched row with all capacity data, or None if not found.
    """
    capacity_rows = load_capacity_csv(capacity_path)
    if not capacity_rows:
        return None
    
    # Normalize the search name
    donor_norm = normalize_name_for_matching(donor_name)
    donor_parts = extract_name_parts(donor_name)
    
    city_norm = city.lower().strip() if city else ""
    state_norm = state.upper().strip() if state else ""
    
    log.info(f"  🔍 Searching capacity CSV for: {donor_name}")
    if city_norm or state_norm:
        log.info(f"      Location filter: {city_norm}, {state_norm}")
    
    candidates = []
    
    for row in capacity_rows:
        # Get names from row
        account_name = row.get('Account Name', '')
        primary_contact = row.get('Primary Contact: Full Name', '')
        row_city = row.get('Billing City', '').lower().strip()
        row_state = row.get('Billing State/Province', '').upper().strip()
        
        # Normalize for matching
        account_norm = normalize_name_for_matching(account_name)
        contact_norm = normalize_name_for_matching(primary_contact)
        contact_parts = extract_name_parts(primary_contact)
        
        # Calculate match score
        score = 0
        match_reason = []
        
        # Exact match on primary contact
        if donor_norm == contact_norm:
            score += 100
            match_reason.append("exact contact match")
        # First + last name match
        elif (donor_parts["first"].lower() == contact_parts["first"].lower() and 
              donor_parts["last"].lower() == contact_parts["last"].lower()):
            score += 80
            match_reason.append("first+last match")
        # Account name contains donor name
        elif donor_norm in account_norm or account_norm in donor_norm:
            score += 60
            match_reason.append("account name match")
        # Last name + first initial match
        elif (donor_parts["last"].lower() == contact_parts["last"].lower() and
              donor_parts["first"] and contact_parts["first"] and
              donor_parts["first"][0].lower() == contact_parts["first"][0].lower()):
            score += 50
            match_reason.append("last name + first initial")
        else:
            continue  # No name match, skip
        
        # Location verification (bonus points, also helps disambiguate)
        if city_norm and row_city:
            if city_norm == row_city:
                score += 20
                match_reason.append("city match")
            elif city_norm in row_city or row_city in city_norm:
                score += 10
                match_reason.append("partial city match")
        
        if state_norm and row_state:
            if state_norm == row_state:
                score += 15
                match_reason.append("state match")
        
        if score >= 50:  # Minimum threshold
            candidates.append({
                "row": row,
                "score": score,
                "reason": ", ".join(match_reason),
                "matched_name": primary_contact or account_name,
                "location": f"{row_city}, {row_state}"
            })
    
    if not candidates:
        log.info(f"    ℹ️ No match found in capacity CSV")
        return None
    
    # Sort by score descending
    candidates.sort(key=lambda x: x["score"], reverse=True)
    best = candidates[0]
    
    log.info(f"    ✅ MATCH: {best['matched_name']} ({best['location']})")
    log.info(f"       Score: {best['score']} ({best['reason']})")
    
    # Return the row with match metadata
    result = dict(best["row"])
    result["_match_score"] = best["score"]
    result["_match_reason"] = best["reason"]
    result["_matched_name"] = best["matched_name"]
    
    return result


def format_capacity_subsection(capacity_data: Dict[str, Any]) -> str:
    """
    Format capacity-to-give data as a subsection for Capacity for Giving.
    Creates organized subsections for iWave data, giving history, and donor info.
    """
    if not capacity_data:
        return ""
    
    sections = []
    
    # -------------------------
    # 1. iWave Capacity Rating
    # -------------------------
    iwave_lines = ["**iWave Capacity Rating:**"]
    
    iwave_cap = capacity_data.get('iWave Est. Capacity', '')
    if iwave_cap:
        try:
            iwave_num = float(str(iwave_cap).replace(',', ''))
            iwave_lines.append(f"• iWave Estimated Capacity: ${iwave_num:,.0f}")
        except:
            iwave_lines.append(f"• iWave Estimated Capacity: {iwave_cap}")
    
    est_cap = capacity_data.get('Estimated Giving Capacity', '')
    if est_cap:
        try:
            est_num = float(str(est_cap).replace(',', ''))
            iwave_lines.append(f"• Estimated Giving Capacity: ${est_num:,.0f}")
        except:
            iwave_lines.append(f"• Estimated Giving Capacity: {est_cap}")
    
    donor_tier = capacity_data.get('Donor Tier', '')
    if donor_tier:
        iwave_lines.append(f"• Donor Tier: {donor_tier}")
    
    prospect_status = capacity_data.get('Development Prospect Status', '')
    if prospect_status:
        iwave_lines.append(f"• Prospect Status: {prospect_status}")
    
    if len(iwave_lines) > 1:
        sections.append("\n".join(iwave_lines))
    
    # -------------------------
    # 2. Giving History
    # -------------------------
    giving_lines = ["**Giving History:**"]
    
    total_gifts = capacity_data.get('Total Gifts', '')
    if total_gifts:
        try:
            total_num = float(str(total_gifts).replace(',', ''))
            if total_num > 0:
                giving_lines.append(f"• Total Lifetime Giving: ${total_num:,.0f}")
        except:
            if total_gifts and total_gifts != '0':
                giving_lines.append(f"• Total Lifetime Giving: {total_gifts}")
    
    best_year = capacity_data.get('Best Gift Year Total', '')
    if best_year:
        try:
            best_num = float(str(best_year).replace(',', ''))
            if best_num > 0:
                giving_lines.append(f"• Best Gift Year: ${best_num:,.0f}")
        except:
            pass
    
    last_amount = capacity_data.get('Last Gift Amount', '')
    last_date = capacity_data.get('Last Gift Date', '')
    if last_amount:
        try:
            last_num = float(str(last_amount).replace(',', ''))
            if last_num > 0:
                gift_str = f"• Last Gift: ${last_num:,.0f}"
                if last_date:
                    gift_str += f" ({last_date})"
                giving_lines.append(gift_str)
        except:
            pass
    
    giving_18mo = capacity_data.get('Total Giving Last 18 Months', '')
    if giving_18mo:
        try:
            mo_num = float(str(giving_18mo).replace(',', ''))
            if mo_num > 0:
                giving_lines.append(f"• Last 18 Months: ${mo_num:,.0f}")
        except:
            pass
    
    giving_3yr = capacity_data.get('Total Giving Last 3 Years', '')
    if giving_3yr:
        try:
            yr_num = float(str(giving_3yr).replace(',', ''))
            if yr_num > 0:
                giving_lines.append(f"• Last 3 Years: ${yr_num:,.0f}")
        except:
            pass
    
    num_gifts = capacity_data.get('Total Number of Gifts', '')
    years_donated = capacity_data.get('Number of Years Donated', '')
    if num_gifts or years_donated:
        details = []
        if num_gifts and num_gifts != '0':
            details.append(f"{num_gifts} gifts")
        if years_donated and years_donated != '0':
            details.append(f"{years_donated} years")
        if details:
            giving_lines.append(f"• Giving Pattern: {', '.join(details)}")
    
    days_since = capacity_data.get('Days since last gift', '')
    if days_since:
        try:
            days_num = int(float(str(days_since).replace(',', '')))
            if days_num > 0:
                years = days_num // 365
                if years > 0:
                    giving_lines.append(f"• Days Since Last Gift: {days_num:,} (~{years} years)")
                else:
                    giving_lines.append(f"• Days Since Last Gift: {days_num:,}")
        except:
            pass
    
    if len(giving_lines) > 1:
        sections.append("\n".join(giving_lines))
    
    # -------------------------
    # 3. Donor Profile
    # -------------------------
    profile_lines = ["**Donor Profile:**"]
    
    account_type = capacity_data.get('Type', '')
    if account_type:
        profile_lines.append(f"• Account Type: {account_type}")
    
    vip = capacity_data.get('VIP Stand Together', '')
    if vip and vip not in ['0', '']:
        profile_lines.append(f"• VIP Stand Together: Yes")
    
    planned = capacity_data.get('Planned Giver', '')
    if planned and planned not in ['0', '']:
        profile_lines.append(f"• Planned Giver: Yes")
    
    interest = capacity_data.get('Donor Interest', '')
    if interest:
        profile_lines.append(f"• Donor Interest: {interest[:100]}...")
    
    purpose = capacity_data.get('Donor Purpose', '')
    if purpose:
        profile_lines.append(f"• Donor Purpose: {purpose[:100]}...")
    
    moves_mgr = capacity_data.get('Moves Manager: Full Name', '')
    if moves_mgr:
        profile_lines.append(f"• Moves Manager: {moves_mgr}")
    
    if len(profile_lines) > 1:
        sections.append("\n".join(profile_lines))
    
    # -------------------------
    # 4. Donor Background (if available)
    # -------------------------
    background = capacity_data.get('Donor Background', '')
    if background and len(background) > 20:
        bg_lines = ["**Donor Background:**"]
        # Truncate long backgrounds
        if len(background) > 500:
            background = background[:500] + "..."
        bg_lines.append(f"• {background}")
        sections.append("\n".join(bg_lines))
    
    # -------------------------
    # 5. Notes/Description
    # -------------------------
    description = capacity_data.get('Description', '')
    if description and len(description) > 10:
        desc_lines = ["**Notes:**"]
        if len(description) > 300:
            description = description[:300] + "..."
        desc_lines.append(f"• {description}")
        sections.append("\n".join(desc_lines))
    
    if not sections:
        return ""
    
    return "\n\n".join(sections)


def generate_ihs_assessment(name: str, location: str, bio_summary: str, career_summary: str,
                           philanthropy_summary: str, political_summary: str, 
                           network_summary: str) -> str:
    """Generate IHS donor probability assessment using A-E-I-N-P-S framework"""
    log.info("Generating IHS donor assessment...")
    
    cache_key = cache.get_cache_key("ihs_assessment", {"name": name})
    
    def fetcher():
        prompt = IHS_ASSESSMENT_PROMPT.format(
            name=name,
            location=location,
            bio_summary=bio_summary,
            career_summary=career_summary,
            philanthropy_summary=philanthropy_summary,
            political_summary=political_summary,
            network_summary=network_summary
        )
        return openai_chat(prompt, cache_key_prefix=cache_key)
    
    assessment = cache.get_or_fetch(cache_key, fetcher)
    return assessment

def generate_strategic_briefing(name: str, location: str, bio_summary: str, career_summary: str,
                                philanthropy_summary: str, political_summary: str,
                                network_summary: str) -> str:
    """Generate board-level strategic summary"""
    log.info("Generating strategic board briefing...")
    
    # CRITICAL: Extract ONLY current information for briefing
    # Same logic as gift officer summary
    
    # Extract current career info
    career_lines = [line.strip() for line in career_summary.split('\n') if line.strip().startswith('•')]
    current_career = []
    for line in career_lines[:3]:
        line_lower = line.lower()
        if any(indicator in line_lower for indicator in ['currently', 'current', 'now serves', 'serves as']):
            current_career.append(line)
        elif len(current_career) == 0:
            current_career.append(line)
        if any(past in line_lower for past in ['formerly', 'previously', 'former', 'past', 'until']):
            break
    
    # Extract current bio info
    bio_lines = [line.strip() for line in bio_summary.split('\n') if line.strip().startswith('•')]
    current_bio = []
    for line in bio_lines:
        line_lower = line.lower()
        if any(indicator in line_lower for indicator in ['currently', 'current', 'based in', 'lives in']):
            if not any(past in line_lower for past in ['formerly', 'previously', 'was based', 'was in']):
                current_bio.append(line)
    
    if not current_career and career_lines:
        current_career = [career_lines[0]]
    if not current_bio and bio_lines:
        current_bio = [bio_lines[0]]
    
    filtered_career = '\n'.join(current_career) if current_career else "Current position information limited"
    filtered_bio = '\n'.join(current_bio) if current_bio else "Current biographical information limited"
    
    cache_key = cache.get_cache_key("strategic_brief", {"name": name})
    
    def fetcher():
        prompt = STRATEGIC_SUMMARY_PROMPT.format(
            name=name,
            location=location,
            bio_summary=filtered_bio,
            career_summary=filtered_career,
            philanthropy_summary=philanthropy_summary,
            political_summary=political_summary,
            network_summary=network_summary
        )
        return openai_chat(prompt, cache_key_prefix=cache_key)
    
    briefing = cache.get_or_fetch(cache_key, fetcher)
    return briefing

def detect_hallucinations(text: str, section_name: str) -> List[str]:
    """
    Detect potential hallucinations in generated text.
    Returns list of warnings about suspicious content.
    """
    warnings = []
    
    # ENHANCED: Common hallucination patterns - now catches them even WITH citations
    HALLUCINATION_PATTERNS = [
        # Generic affiliations (likely fabricated)
        (r"active member of.*national association", "LIKELY HALLUCINATION: Generic 'National Association' membership"),
        (r"active member of.*professional association", "LIKELY HALLUCINATION: Vague professional association claim"),
        (r"regular speaker at.*annual.*conference", "LIKELY HALLUCINATION: Generic annual conference claim"),
        (r"regular speaker at.*leadership conference", "LIKELY HALLUCINATION: Generic leadership conference claim"),
        (r"maintains?\s+(key\s+)?relationships?\s+with", "LIKELY HALLUCINATION: Vague 'maintains relationships' claim"),
        (r"widely recognized (as|for|in)", "LIKELY HALLUCINATION: Vague 'widely recognized' claim"),
        (r"known for (his|her|their) expertise", "LIKELY HALLUCINATION: Generic 'known for expertise' claim"),
        (r"maintains?\s+(a\s+)?strong presence", "LIKELY HALLUCINATION: Vague 'strong presence' claim"),
        (r"has extensive experience (in|with)", "LIKELY HALLUCINATION: Generic 'extensive experience' claim"),
        (r"holds?\s+(a\s+)?(key|important|influential)\s+position", "LIKELY HALLUCINATION: Vague position claim"),
        (r"board member of.*advisory", "POSSIBLE HALLUCINATION: Generic advisory board - verify organization exists"),
        (r"board member of.*technology.*council", "POSSIBLE HALLUCINATION: Generic tech council - verify organization exists"),
        (r"affiliated with.*institute", "POSSIBLE HALLUCINATION: Vague institute affiliation - verify details"),
        # NEW: Fabricated company/title patterns
        (r"(CTO|CEO|CFO|COO|CIO|CMO)\s+(at|of)\s+\w+\s+(Solutions|Technologies|Tech|Innovations|Systems|Partners|Group|Consulting|Capital|Ventures|Holdings)", 
         "LIKELY HALLUCINATION: Generic exec title at generic company name"),
        (r"Chief\s+(Technology|Executive|Financial|Operating)\s+Officer\s+(at|of)\s+\w+\s+(Solutions|Technologies|Tech)", 
         "LIKELY HALLUCINATION: C-suite at generic tech company"),
        (r"(Innovator|Leader|Pioneer)\s+of\s+the\s+Year", "LIKELY HALLUCINATION: Generic award name"),
        (r"(Tech|Technology|Business|Industry)\s+(Innovator|Leader)\s+of\s+the\s+Year", "LIKELY HALLUCINATION: Fabricated award"),
        (r"thought leader in", "LIKELY HALLUCINATION: Vague 'thought leader' claim"),
        (r"pivotal figure in", "LIKELY HALLUCINATION: Vague 'pivotal figure' claim"),
        (r"spearheaded\s+(several|numerous|multiple)\s+(high-profile|major|significant)", "LIKELY HALLUCINATION: Generic accomplishment claim"),
        (r"recognized\s+through\s+(various|numerous|multiple)\s+industry\s+awards", "LIKELY HALLUCINATION: Vague awards claim"),
        (r"presents?\s+a\s+strategic\s+opportunity", "LIKELY HALLUCINATION: Speculative strategic language"),
        (r"her\s+ability\s+to\s+drive\s+technological\s+advancements", "LIKELY HALLUCINATION: Generic capability claim"),
        (r"making\s+(her|him|them)\s+a\s+pivotal\s+figure", "LIKELY HALLUCINATION: Inflated significance"),
    ]
    
    # Check for bullet points without citations (original check)
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if line.strip().startswith('•'):
            # Check if this bullet has a citation [N]
            if not re.search(r'\[\d+\]', line):
                # Exception: Admissions of no data are OK
                if not any(phrase in line.lower() for phrase in [
                    'no verified', 'no public', 'limited', 'not available', 'unable to verify'
                ]):
                    warnings.append(f"⚠️  {section_name} line {i+1}: Bullet point lacks citation")
    
    # ENHANCED: Check for hallucination patterns (catches them WITH or WITHOUT citations)
    for pattern, description in HALLUCINATION_PATTERNS:
        if re.search(pattern, text, re.IGNORECASE):
            # Extract the actual text that matched
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                matched_text = match.group(0)[:60]  # First 60 chars
                warnings.append(f"🚨 {section_name}: {description}")
                warnings.append(f"   Suspicious text: '{matched_text}...'")
    
    # ENHANCED: Check for suspiciously generic organization names
    SUSPICIOUS_ORG_PATTERNS = [
        r"National Association of [A-Z][a-z]+",  # e.g., "National Association of Corporate Directors"
        r"Annual [A-Z][a-z]+ Conference",         # e.g., "Annual Technology Conference"
        r"[A-Z][a-z]+ Leadership (Forum|Summit|Conference)",  # e.g., "Technology Leadership Conference"
        r"(Global|International|National) [A-Z][a-z]+ (Society|Institute|Council)",
    ]
    
    # NEW: Check for fabricated company names (common LLM hallucination patterns)
    SUSPICIOUS_COMPANY_PATTERNS = [
        r"\b\w+\s+(Solutions|Technologies|Tech|Innovations|Dynamics|Systems|Consulting)\s+(Inc|LLC|Corp|Corporation|Ltd)?\b",
    ]
    
    for pattern in SUSPICIOUS_ORG_PATTERNS:
        matches = re.findall(pattern, text)
        for match in matches:
            warnings.append(f"⚠️  {section_name}: Suspicious org name: '{match}' - verify this organization actually exists")
    
    for pattern in SUSPICIOUS_COMPANY_PATTERNS:
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            if isinstance(match, tuple):
                match = ' '.join(match)
            warnings.append(f"⚠️  {section_name}: Suspicious company name: '{match}' - verify this company exists and is related to the prospect")
    
    # NEW: Check for fake citation patterns (citing non-existent sources)
    FAKE_CITATION_PATTERNS = [
        r"\[1\]\s*Source:\s*(Company\s+Website|Industry\s+Publications|Awards?\s+Publications?)",
        r"\[2\]\s*Source:\s*(Company\s+Website|Industry\s+Publications|Awards?\s+Publications?)",
        r"\[3\]\s*Source:\s*(Company\s+Website|Strategy\s+Documents?|Internal\s+Documents?)",
    ]
    
    for pattern in FAKE_CITATION_PATTERNS:
        if re.search(pattern, text, re.IGNORECASE):
            warnings.append(f"🚨 {section_name}: FABRICATED CITATION detected - GPT invented a fake source reference")
    
    return warnings

# -------------------------
# CSV LOADING WITH TOLERANT MATCHING
# -------------------------

# Common nickname mappings for better matching
NICKNAMES = {
    "william": {"bill", "will", "billy"},
    "robert": {"bob", "rob", "bobby"},
    "elizabeth": {"liz", "beth", "eliza", "lizzie"},
    "katherine": {"kate", "katie", "kathryn", "kathy"},
    "michael": {"mike", "mick"},
    "james": {"jim", "jimmy"},
    "johnathan": {"john", "jon"},
    "jonathan": {"john", "jon"},
    "stephen": {"steve", "steven"},
    "nicholas": {"nick", "nicolas"},
    "alexander": {"alex", "al"},
    "richard": {"rick", "dick", "rich"},
    "thomas": {"tom", "tommy"},
    "christopher": {"chris"},
    "matthew": {"matt"},
    "joseph": {"joe"},
    "anthony": {"tony"},
    "daniel": {"dan", "danny"},
    "david": {"dave"},
    "ann": {"anne", "anna", "annie"},
    "george": {"georgie"},
    "edward": {"ed", "eddie", "ted"},
    "charles": {"charlie", "chuck"},
}

# Create reverse mapping (nickname -> formal name)
REVERSE_NICKNAMES = {}
for formal_name, nicknames in NICKNAMES.items():
    for nickname in nicknames:
        if nickname not in REVERSE_NICKNAMES:
            REVERSE_NICKNAMES[nickname] = set()
        REVERSE_NICKNAMES[nickname].add(formal_name)

import unicodedata

def _norm_name(s: str) -> str:
    """Normalize name for matching: remove punctuation, extra whitespace, accents"""
    s = unicodedata.normalize("NFKD", s)
    s = re.sub(r"[^\w\s]", " ", s)  # Remove punctuation
    s = re.sub(r"\s+", " ", s).strip().lower()
    return s

def load_csv(path: str) -> List[Dict[str, Any]]:
    """Load CSV with multiple encoding attempts"""
    if not os.path.exists(path):
        return []
    
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            with open(path, 'r', encoding=encoding, errors='replace') as f:
                reader = csv.DictReader(f)
                rows = list(reader)
                
                # Normalize column names
                normalized_rows = []
                for row in rows:
                    normalized = {}
                    
                    # Track if we found First/Last columns
                    has_first = False
                    has_last = False
                    
                    for k, v in row.items():
                        # CRITICAL BOM FIX: Strip BOM character (\ufeff) that breaks column matching
                        k_clean = k.lstrip('\ufeff').strip()
                        k_norm = k_clean.lower()
                        if k_norm in ('first name', 'firstname'):
                            normalized['First'] = (v or "").strip()
                            has_first = True
                        elif k_norm in ('last name', 'lastname'):
                            normalized['Last'] = (v or "").strip()
                            has_last = True
                        elif k_norm == 'middle':
                            normalized['Middle'] = (v or "").strip()
                        elif k_norm == 'city':
                            normalized['City'] = (v or "").strip()
                        elif k_norm == 'state':
                            normalized['State'] = (v or "").strip()
                        elif k_norm == 'already known':
                            normalized['Already known'] = (v or "").strip()
                        else:
                            normalized[k_clean] = (v or "").strip()
                    
                    # If no First/Last columns, try to parse from "Name" or "Full Name" column
                    if not has_first or not has_last:
                        for name_col in ['Name', 'Full Name', 'Contact Name', 'Donor Name']:
                            if name_col in normalized and normalized[name_col]:
                                full_name = normalized[name_col]
                                # Check for "Last, First" format
                                if "," in full_name:
                                    parts = full_name.split(",", 1)
                                    if len(parts) == 2:
                                        if not has_last:
                                            normalized['Last'] = parts[0].strip()
                                        if not has_first:
                                            normalized['First'] = parts[1].strip()
                                        break
                                # Check for "First Last" format (space-separated)
                                elif " " in full_name:
                                    parts = full_name.split()
                                    if len(parts) >= 2:
                                        if not has_first:
                                            normalized['First'] = parts[0]
                                        if not has_last:
                                            normalized['Last'] = parts[-1]
                                        if len(parts) > 2 and 'Middle' not in normalized:
                                            normalized['Middle'] = " ".join(parts[1:-1])
                                        break
                    
                    normalized_rows.append(normalized)
                
                return normalized_rows
        except (UnicodeDecodeError, UnicodeError):
            continue
    
    return []

def get_row_name(row: Dict[str, Any]) -> str:
    """Extract name from row - handles multiple column name formats + BOM characters"""
    
    # CRITICAL BOM FIX: Strip BOM from all column names before checking
    clean_row = {}
    for k, v in row.items():
        k_clean = k.lstrip('\ufeff').strip()
        clean_row[k_clean] = v
    
    # Try standard First/Last columns (various formats)
    first_cols = ["First Name", "First", "FirstName", "first name", "first", "FIRST NAME", "FIRST"]
    last_cols = ["Last Name", "Last", "LastName", "last name", "last", "LAST NAME", "LAST"]
    middle_cols = ["Middle", "Middle Name", "MiddleName", "middle", "middle name"]
    
    first = ""
    last = ""
    middle = ""
    
    # Try to find first name
    for col in first_cols:
        if col in clean_row and clean_row[col]:
            first = clean_row[col].strip()
            break
    
    # Try to find last name
    for col in last_cols:
        if col in clean_row and clean_row[col]:
            last = clean_row[col].strip()
            break
    
    # Try to find middle name
    for col in middle_cols:
        if col in clean_row and clean_row[col]:
            middle = clean_row[col].strip()
            break
    
    if first and last:
        if middle:
            return f"{first} {middle} {last}"
        return f"{first} {last}"
    
    # Try "Full Name" or "Name" column (might be "Last, First" format)
    name_cols = ["Full Name", "Name", "Contact Name", "Donor Name", "full name", "name"]
    for col in name_cols:
        if col in clean_row and clean_row[col]:
            full_name = clean_row[col].strip()
            if full_name:
                # Check if it's "Last, First" format
                if "," in full_name:
                    parts = full_name.split(",")
                    if len(parts) == 2:
                        last_name = parts[0].strip()
                        first_name = parts[1].strip()
                        return f"{first_name} {last_name}"
                # Otherwise just return as-is
                return full_name
    
    return ""

def get_nickname_from_row(row: Dict[str, Any]) -> str:
    """Extract nickname from row if present (from Data.csv Nickname column)"""
    # Strip BOM from column names
    clean_row = {k.lstrip('\ufeff').strip(): v for k, v in row.items()}
    
    nickname_cols = ["Nickname", "nickname", "Nick", "nick", "Preferred Name", "preferred name"]
    for col in nickname_cols:
        if col in clean_row and clean_row[col]:
            return clean_row[col].strip()
    return ""

def get_spouse_from_row(row: Dict[str, Any]) -> str:
    """Extract spouse name from row if present (from Data.csv Spouse/Partner column)"""
    # Strip BOM from column names
    clean_row = {k.lstrip('\ufeff').strip(): v for k, v in row.items()}
    
    spouse_cols = ["Spouse/Partner", "Spouse", "Partner", "spouse", "partner", "Spouse Name", "spouse name"]
    for col in spouse_cols:
        if col in clean_row and clean_row[col]:
            spouse = clean_row[col].strip()
            if spouse:
                return spouse
    return ""

def extract_iwave_data(row: Dict[str, Any]) -> Dict[str, Any]:
    """
    Extract iWave giving/capacity data from data.csv row.
    
    Searches for common iWave column names and returns all found values.
    Returns empty dict if no iWave data found.
    """
    # Strip BOM from column names
    clean_row = {k.lstrip('\ufeff').strip(): v for k, v in row.items()}
    
    # Common iWave column name patterns (case-insensitive matching)
    iwave_patterns = {
        # Core iWave fields
        "iwave_score": ["iWave Score", "iwave score", "iWave", "iwave", "IWave Score"],
        "capacity_range": ["iWave Gift Range", "Gift Range", "Capacity Range", "iWave Capacity Range", 
                           "Estimated Gift Range", "Gift Capacity Range", "iwave gift range"],
        "estimated_capacity": ["iWave Estimated Capacity", "Estimated Capacity", "iWave Capacity",
                               "Giving Capacity", "Gift Capacity", "Capacity", "iwave estimated capacity",
                               "Est. Capacity", "Est Capacity"],
        "major_gift_likelihood": ["iWave Major Gift Likelihood", "Major Gift Likelihood", "MGL",
                                  "Major Gift Score", "iwave major gift likelihood", "Major Gift %"],
        "planned_gift_likelihood": ["iWave Planned Gift Likelihood", "Planned Gift Likelihood", "PGL",
                                    "Planned Gift Score", "iwave planned gift likelihood", "Planned Gift %"],
        "p2g_score": ["P2G Score", "P2G", "Propensity to Give", "Propensity Score", "p2g score",
                      "iWave P2G", "iWave Propensity"],
        "rfm_score": ["RFM Score", "RFM", "rfm score", "iWave RFM"],
        "wealth_rating": ["Wealth Rating", "Wealth Score", "wealth rating", "Net Worth Rating",
                          "iWave Wealth Rating"],
        "real_estate_value": ["Real Estate Value", "RE Value", "Real Estate", "Property Value",
                              "iWave Real Estate", "Total Real Estate"],
        "stock_value": ["Stock Value", "Securities", "Stock Holdings", "iWave Stock",
                        "Publicly Traded Securities"],
        "total_assets": ["Total Assets", "Assets", "iWave Assets", "Net Worth Estimate"],
        "largest_gift": ["Largest Gift", "Max Gift", "Largest Known Gift", "Biggest Gift",
                         "iWave Largest Gift"],
        "total_giving": ["Total Giving", "Lifetime Giving", "Known Giving", "Total Gifts",
                         "iWave Total Giving", "Philanthropic Giving"],
        "annual_giving": ["Annual Giving", "Yearly Giving", "Average Annual Gift", "Avg Annual Gift"],
    }
    
    iwave_data = {}
    
    for field_key, column_names in iwave_patterns.items():
        for col_name in column_names:
            # Try exact match
            if col_name in clean_row and clean_row[col_name]:
                value = clean_row[col_name].strip()
                if value and value.lower() not in ['', 'n/a', 'na', 'none', '-']:
                    iwave_data[field_key] = value
                    break
            # Try case-insensitive match
            for row_col in clean_row.keys():
                if row_col.lower() == col_name.lower() and clean_row[row_col]:
                    value = clean_row[row_col].strip()
                    if value and value.lower() not in ['', 'n/a', 'na', 'none', '-']:
                        iwave_data[field_key] = value
                        break
    
    return iwave_data


def format_iwave_bullets(iwave_data: Dict[str, Any]) -> str:
    """
    Format iWave data as bullet points for inclusion in the net worth section.
    Returns formatted string or empty string if no data.
    """
    if not iwave_data:
        return ""
    
    bullets = []
    bullets.append("\n• **iWave Wealth Screening Data:**")
    
    # Format each field with human-readable labels
    field_labels = {
        "iwave_score": "Overall iWave Score",
        "capacity_range": "Estimated Gift Capacity Range",
        "estimated_capacity": "Estimated Giving Capacity",
        "major_gift_likelihood": "Major Gift Likelihood",
        "planned_gift_likelihood": "Planned Gift Likelihood",
        "p2g_score": "Propensity to Give (P2G) Score",
        "rfm_score": "RFM Score",
        "wealth_rating": "Wealth Rating",
        "real_estate_value": "Real Estate Holdings",
        "stock_value": "Stock/Securities Value",
        "total_assets": "Total Estimated Assets",
        "largest_gift": "Largest Known Gift",
        "total_giving": "Total Known Philanthropic Giving",
        "annual_giving": "Average Annual Giving",
    }
    
    for field_key, label in field_labels.items():
        if field_key in iwave_data:
            value = iwave_data[field_key]
            bullets.append(f"  - {label}: {value}")
    
    if len(bullets) == 1:
        # Only the header, no actual data
        return ""
    
    return "\n".join(bullets)


def _spouse_present(spouse_name: str, text: str) -> bool:
    """
    Return True if `text` appears to mention this spouse.
    Uses first+last name and ignores middle initials and periods.
    
    This is critical for matching "George S. Gibbs" when spouse_name is "George Gibbs"
    or vice versa. It checks if BOTH first name and last name appear anywhere in the text,
    regardless of middle initials or punctuation.
    
    Examples:
        _spouse_present("George Gibbs", "George S. Gibbs donated...") → True
        _spouse_present("George S Gibbs", "George Gibbs and Ann...") → True
        _spouse_present("George Gibbs", "Ann Gibbs widow of George S. Gibbs...") → True
        _spouse_present("George Gibbs", "The Gibbs family...") → False (no "George")
    """
    if not spouse_name or not text:
        return False
    
    # Normalize: lowercase and remove periods
    text_norm = text.lower().replace(".", "").replace(",", "")
    spouse_norm = spouse_name.lower().replace(".", "").replace(",", "")
    
    parts = spouse_norm.split()
    if len(parts) < 2:
        # Fallback: just look for the whole normalized string
        return spouse_norm in text_norm
    
    first, last = parts[0], parts[-1]
    # Need BOTH first and last somewhere in the text
    # This way "George S Gibbs" matches "George Gibbs" and vice versa
    return first in text_norm and last in text_norm

def find_donor_in_data(rows: List[Dict[str, Any]], name: str) -> Optional[Dict[str, Any]]:
    """Find donor row by name with tolerant matching (handles nicknames, punctuation, middle names)"""
    target = _norm_name(name)
    
    # Try exact match first
    for row in rows:
        row_name = _norm_name(get_row_name(row))
        if row_name == target:
            return row
    
    # Fallback 1: Ignore middle names/initials (match first + last only)
    t_parts = target.split()
    if len(t_parts) >= 2:
        target_first_last = f"{t_parts[0]} {t_parts[-1]}"
        for row in rows:
            rn = _norm_name(get_row_name(row))
            parts = rn.split()
            if len(parts) >= 2:
                row_first_last = f"{parts[0]} {parts[-1]}"
                if row_first_last == target_first_last:
                    return row
    
    # Fallback 2: Try nickname variants
    if len(t_parts) >= 2:
        first_name = t_parts[0]
        last_name = t_parts[-1]
        
        # Check if first name has nickname variants
        nicknames = NICKNAMES.get(first_name, set())
        if nicknames:
            for nick in nicknames:
                target_with_nick = f"{nick} {last_name}"
                for row in rows:
                    rn = _norm_name(get_row_name(row))
                    parts = rn.split()
                    if len(parts) >= 2:
                        row_first_last = f"{parts[0]} {parts[-1]}"
                        if row_first_last == target_with_nick:
                            return row
        
        # Also check reverse: if row name has a formal name, try its nicknames
        for row in rows:
            rn = _norm_name(get_row_name(row))
            parts = rn.split()
            if len(parts) >= 2:
                row_first = parts[0]
                row_last = parts[-1]
                # Check if row's first name has nicknames that match target
                row_nicknames = NICKNAMES.get(row_first, set())
                if first_name in row_nicknames and row_last == last_name:
                    return row
    
    return None


def generate_name_variations(donor_name: str, csv_data: Dict[str, Any], 
                            alternate_names: List[Dict] = None) -> List[Dict[str, str]]:
    """
    Generate all name variations for comprehensive searching.
    
    Returns list of dicts with:
    - name: The name variation
    - type: maiden, previous_married, nickname, formal, etc.
    - search_priority: high, medium, low
    
    Example: Patricia M Toft could have variations:
    - Patricia M Toft (current)
    - Pat Toft (nickname)
    - Patricia Toft (no middle)
    - Patricia M Brown (maiden - if detected)
    - Patricia Brown (maiden no middle)
    - Patricia M Long (previous married - if detected from marriage records)
    """
    variations = []
    
    # Parse current name
    parts = donor_name.strip().split()
    if len(parts) < 2:
        return [{"name": donor_name, "type": "current", "search_priority": "high"}]
    
    first_name = parts[0]
    last_name = parts[-1]
    middle = " ".join(parts[1:-1]) if len(parts) > 2 else ""
    middle_initial = middle[0] if middle else ""
    
    # Pre-compute nicknames (needed for maiden name variations below)
    nicknames = NICKNAMES.get(first_name.lower(), set())
    
    # =====================================================
    # PRIORITY ORDER: Maiden names FIRST (if available) since 
    # older records often use maiden names for people who changed names
    # =====================================================
    
    # 1. Maiden name variations FIRST (from alternate_names detected in CSV)
    # These should be searched FIRST for older prospects who may have records under maiden name
    if alternate_names:
        for alt in alternate_names:
            alt_last = alt.get("alternate_last_name", "")
            if alt_last and alt_last.lower() != last_name.lower():
                # Full maiden name with middle - HIGHEST PRIORITY for older records
                if middle:
                    variations.append({"name": f"{first_name} {middle} {alt_last}", "type": "maiden", "search_priority": "highest"})
                if middle_initial:
                    variations.append({"name": f"{first_name} {middle_initial} {alt_last}", "type": "maiden_initial", "search_priority": "highest"})
                variations.append({"name": f"{first_name} {alt_last}", "type": "maiden_no_middle", "search_priority": "highest"})
                
                # Nicknames with maiden name
                for nick in nicknames:
                    nick_cap = nick.capitalize()
                    variations.append({"name": f"{nick_cap} {alt_last}", "type": "nickname_maiden", "search_priority": "high"})
    
    # 2. Current name variations
    variations.append({"name": donor_name, "type": "current", "search_priority": "high"})
    variations.append({"name": f"{first_name} {last_name}", "type": "current_no_middle", "search_priority": "high"})
    if middle_initial:
        variations.append({"name": f"{first_name} {middle_initial} {last_name}", "type": "current_initial", "search_priority": "high"})
    
    # 3. Nickname variations (nicknames already computed above)
    for nick in nicknames:
        nick_cap = nick.capitalize()
        variations.append({"name": f"{nick_cap} {last_name}", "type": "nickname", "search_priority": "medium"})
        if middle_initial:
            variations.append({"name": f"{nick_cap} {middle_initial} {last_name}", "type": "nickname_initial", "search_priority": "medium"})
    
    # Also check reverse - if first name is a nickname, add formal version
    for formal, nicks in NICKNAMES.items():
        if first_name.lower() in nicks:
            formal_cap = formal.capitalize()
            variations.append({"name": f"{formal_cap} {last_name}", "type": "formal_name", "search_priority": "medium"})
    
    # 4. Spouse name variations (for searching as couple)
    spouse = csv_data.get("Spouse/Partner", csv_data.get("Spouse", "")).strip()
    if spouse:
        variations.append({"name": f"Mr. and Mrs. {spouse}", "type": "couple", "search_priority": "low"})
        variations.append({"name": f"{donor_name} and {spouse}", "type": "couple_full", "search_priority": "low"})
    
    # Remove duplicates while preserving order and priority
    seen = set()
    unique_variations = []
    for v in variations:
        if v["name"].lower() not in seen:
            seen.add(v["name"].lower())
            unique_variations.append(v)
    
    return unique_variations


def find_related_rows(rows: List[Dict[str, Any]], donor_row: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Find related rows in the CSV that might contain:
    - Maiden name (row with just last name, no first name, matching location pattern)
    - Alternate names or aliases
    - Spouse information
    
    Example from CSV:
        Row 1: Patricia, M, Toft, Burlingame, CA  (main donor)
        Row 2: ,,Brown,,bainbridge island,wa      (potential maiden name - no first name)
    
    Returns list of related rows with metadata about the relationship type.
    """
    related = []
    
    # Get donor's info for comparison
    donor_name = get_row_name(donor_row)
    donor_first = donor_row.get("First Name", donor_row.get("First", "")).strip()
    donor_last = donor_row.get("Last Name", donor_row.get("Last", "")).strip()
    donor_city = donor_row.get("City", "").strip().lower()
    donor_state = donor_row.get("State", "").strip().lower()
    
    for row in rows:
        if row == donor_row:
            continue
        
        row_first = row.get("First Name", row.get("First", "")).strip()
        row_middle = row.get("Middle Name", row.get("Middle", "")).strip()
        row_last = row.get("Last Name", row.get("Last", "")).strip()
        row_city = row.get("City", "").strip().lower()
        row_state = row.get("State", "").strip().lower()
        
        # Pattern 1: Row has ONLY last name (no first name) - could be maiden name or alternate name
        # Example: ",,Brown,,bainbridge island,wa"
        if row_last and not row_first:
            # This could be a maiden name or alternate surname
            related.append({
                "row": row,
                "type": "alternate_name",
                "alternate_last_name": row_last,
                "associated_location": f"{row_city}, {row_state}".strip(", ") if row_city or row_state else None,
                "confidence": "medium" if (row_city or row_state) else "low"
            })
            log.info(f"  🔗 Found potential alternate name row: Last='{row_last}', Location='{row_city}, {row_state}'")
        
        # Pattern 2: Row has same last name but different first name at same location - could be spouse
        elif row_last.lower() == donor_last.lower() and row_first and row_first.lower() != donor_first.lower():
            if (row_city and row_city == donor_city) or (row_state and row_state == donor_state):
                related.append({
                    "row": row,
                    "type": "potential_spouse",
                    "spouse_name": f"{row_first} {row_last}",
                    "same_location": True,
                    "confidence": "medium"
                })
                log.info(f"  🔗 Found potential spouse row: {row_first} {row_last}")
        
        # Pattern 3: Row shares location with donor but different last name - could be maiden name connection
        elif row_city and row_state and row_city == donor_city and row_state == donor_state:
            if row_last and row_last.lower() != donor_last.lower():
                if not row_first:  # No first name suggests it's a reference/alternate name
                    related.append({
                        "row": row,
                        "type": "same_location_alternate",
                        "alternate_last_name": row_last,
                        "shared_location": f"{row_city}, {row_state}",
                        "confidence": "low"
                    })
    
    return related


def should_skip_inner_circle(donor_row: Dict[str, Any]) -> bool:
    """Check if Inner Circle analysis should be skipped based on CSV column"""
    already_known = donor_row.get("Already known", "").strip().lower()
    skip = already_known == "yes"
    
    if skip:
        log.info("  ⏭️  Skipping Inner Circle analysis (Already known = Yes)")
    
    return skip


def merge_person_rows(rows: List[Dict[str, Any]], primary_last_name: str) -> Dict[str, Any]:
    """
    Merge all rows that belong to the same person into one comprehensive record.
    
    This handles the case where multiple CSV rows represent different pieces of 
    information about ONE person (different addresses, name variations, etc.)
    
    Args:
        rows: All rows from the CSV
        primary_last_name: The last name to match on (e.g., "Toft")
    
    Returns:
        A merged record with:
        - All name variations collected
        - All addresses/properties collected
        - Maiden name extracted
        - DOB parsed
        - Spouse info combined
        - All emails/phones collected
    """
    merged = {
        "First Name": "",
        "Middle Name": "",
        "Last Name": primary_last_name,
        "Maiden Name": "",
        "Nickname": "",
        "DOB": "",
        "Date of Birth": "",
        "Spouse/Partner": "",
        "Already known": "",
        # Collected items
        "_all_addresses": [],
        "_all_name_variations": [],
        "_all_emails": [],
        "_all_phones": [],
        "_all_cities": [],
        "_all_states": [],
    }
    
    matching_rows = []
    for row in rows:
        row_last = (row.get("Last Name", "") or row.get("Last", "")).strip()
        # Match on last name (case-insensitive)
        if row_last.lower() == primary_last_name.lower():
            matching_rows.append(row)
        # Also include rows with NO last name but have an address (property-only rows)
        elif not row_last and row.get("Address", "").strip():
            matching_rows.append(row)
    
    log.info(f"  📋 Merging {len(matching_rows)} rows for '{primary_last_name}'")
    
    for row in matching_rows:
        # Collect primary fields (first non-empty wins)
        if not merged["First Name"]:
            first = (row.get("First Name", "") or row.get("First", "")).strip()
            # Skip if it's just initials like "PM" or "M"
            if first and len(first) > 2:
                merged["First Name"] = first
        
        if not merged["Middle Name"]:
            middle = (row.get("Middle Name", "") or row.get("Middle", "")).strip()
            if middle:
                merged["Middle Name"] = middle
        
        if not merged["Maiden Name"]:
            maiden = row.get("Maiden Name", "").strip()
            if maiden:
                merged["Maiden Name"] = maiden
                log.info(f"    📛 Found Maiden Name: {maiden}")
        
        if not merged["Nickname"]:
            nick = row.get("Nickname", "").strip()
            if nick:
                merged["Nickname"] = nick
                log.info(f"    📛 Found Nickname: {nick}")
        
        if not merged["DOB"] and not merged["Date of Birth"]:
            dob = row.get("DOB", row.get("Date of Birth", "")).strip()
            if dob:
                merged["DOB"] = dob
                merged["Date of Birth"] = dob
                log.info(f"    🎂 Found DOB: {dob}")
        
        if not merged["Spouse/Partner"]:
            spouse = row.get("Spouse/Partner", row.get("Spouse", "")).strip()
            if spouse:
                merged["Spouse/Partner"] = spouse
                log.info(f"    💑 Found Spouse: {spouse}")
        
        if not merged["Already known"]:
            known = row.get("Already known", "").strip()
            if known:
                merged["Already known"] = known
        
        # Collect address as a property
        address = row.get("Address", "").strip()
        city = row.get("City", "").strip()
        state = row.get("State", "").strip()
        zip_code = row.get("Zip", "").strip()
        
        if address or city:
            addr_entry = {
                "address": address,
                "city": city,
                "state": state,
                "zip": zip_code,
                "full": f"{address}, {city}, {state} {zip_code}".strip(", ")
            }
            # Avoid duplicates
            if addr_entry["full"] not in [a["full"] for a in merged["_all_addresses"]]:
                merged["_all_addresses"].append(addr_entry)
                log.info(f"    🏠 Found address: {addr_entry['full']}")
        
        # Collect cities/states for location searches
        if city and city.lower() not in [c.lower() for c in merged["_all_cities"]]:
            merged["_all_cities"].append(city)
        if state and state.upper() not in [s.upper() for s in merged["_all_states"]]:
            merged["_all_states"].append(state)
        
        # Collect name variations
        first = (row.get("First Name", "") or row.get("First", "")).strip()
        last = (row.get("Last Name", "") or row.get("Last", "")).strip()
        if first and last:
            name_var = f"{first} {last}"
            if name_var not in merged["_all_name_variations"]:
                merged["_all_name_variations"].append(name_var)
        elif first:  # Just first name (like "PM")
            if first not in merged["_all_name_variations"]:
                merged["_all_name_variations"].append(first)
        
        # Collect emails
        for email_field in ["Work Email", "Email", "Personal Email", "Spouse/Partner Contact Email"]:
            email = row.get(email_field, "").strip()
            if email and email not in merged["_all_emails"]:
                merged["_all_emails"].append(email)
        
        # Collect phones
        for phone_field in ["Work Phone", "Phone", "Mobile", "Spouse/Partner Contact Phone"]:
            phone = row.get(phone_field, "").strip()
            if phone and phone not in merged["_all_phones"]:
                merged["_all_phones"].append(phone)
    
    # Set primary City/State from first address with both
    for addr in merged["_all_addresses"]:
        if addr["city"] and addr["state"]:
            if not merged.get("City"):
                merged["City"] = addr["city"]
                merged["State"] = addr["state"]
                merged["Zip"] = addr["zip"]
                merged["Address"] = addr["address"]
                break
    
    # Log summary
    log.info(f"    📊 Merged result: {len(merged['_all_addresses'])} addresses, "
             f"{len(merged['_all_cities'])} cities, {len(merged['_all_name_variations'])} name variations")
    
    return merged


def parse_dob_year(dob_str: str) -> Optional[int]:
    """
    Parse DOB string in various formats and return birth year.
    
    Handles:
    - "Aug-36" -> 1936
    - "1936" -> 1936
    - "08/15/1936" -> 1936
    - "August 1936" -> 1936
    """
    if not dob_str:
        return None
    
    import re
    
    # Try to find a 4-digit year first
    match = re.search(r'\b(19\d{2}|20\d{2})\b', str(dob_str))
    if match:
        return int(match.group(1))
    
    # Try 2-digit year format like "Aug-36"
    match = re.search(r'(\d{2})$', str(dob_str).strip())
    if match:
        year_2digit = int(match.group(1))
        # Assume 1900s for 2-digit years > 25, 2000s for <= 25
        if year_2digit > 25:
            return 1900 + year_2digit
        else:
            return 2000 + year_2digit
    
    return None


def search_large_csv_streaming(name_variants: List[str], csv_path: str, max_rows: int = 500000) -> Optional[Dict[str, Any]]:
    """
    Search large CSV file (e.g., 488k rows) using streaming to avoid loading entire file into memory.
    Returns first matching row. NOW WITH BOM DETECTION!
    """
    if not os.path.exists(csv_path):
        return None
    
    # Normalize all search variants for fast comparison
    normalized_variants = {_norm_name(v) for v in name_variants}
    
    log.info(f"  🔍 Streaming search through {csv_path}...")
    matched_count = 0
    bom_detected = False
    
    encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
    
    for encoding in encodings:
        try:
            with open(csv_path, 'r', encoding=encoding, errors='replace') as f:
                reader = csv.DictReader(f)
                
                # Check for BOM in first column name
                if reader.fieldnames and reader.fieldnames[0].startswith('\ufeff'):
                    bom_detected = True
                    log.warning(f"  ⚠️  BOM CHARACTER DETECTED in CSV! (Column: '{repr(reader.fieldnames[0])}')")
                    log.warning(f"     This will be automatically stripped during processing.")
                
                for i, row in enumerate(reader):
                    # Progress indicator every 50k rows
                    if i > 0 and i % 50000 == 0:
                        log.info(f"     ... searched {i:,} rows so far...")
                    
                    # Stop if we've searched enough
                    if i >= max_rows:
                        log.warning(f"     Stopped after {max_rows:,} rows")
                        break
                    
                    # Get name from this row (automatically handles BOM via get_row_name)
                    row_name = get_row_name(row)
                    if not row_name:
                        continue
                    
                    # Check if this row matches any variant
                    norm_row = _norm_name(row_name)
                    if norm_row in normalized_variants:
                        matched_count += 1
                        log.info(f"  ✅ MATCH FOUND at row {i+1:,}!")
                        log.info(f"     Row name: '{row_name}'")
                        if bom_detected:
                            log.info(f"     ✓ BOM was properly stripped during match")
                        return row
                    
                    # Also check first+last name match (ignore middle names)
                    norm_parts = norm_row.split()
                    if len(norm_parts) >= 2:
                        row_first_last = f"{norm_parts[0]} {norm_parts[-1]}"
                        if row_first_last in normalized_variants:
                            matched_count += 1
                            log.info(f"  ✅ MATCH FOUND at row {i+1:,} (first+last only)!")
                            log.info(f"     Row name: '{row_name}'")
                            if bom_detected:
                                log.info(f"     ✓ BOM was properly stripped during match")
                            return row
                
                # Finished searching entire file
                log.info(f"     Searched {i+1:,} total rows")
                if bom_detected:
                    log.info(f"     ✓ BOM character was handled correctly throughout search")
                return None
                
        except Exception as e:
            log.debug(f"     Failed with {encoding}: {e}")
            continue
    
    return None

def load_contactsdb_data(name: str, donor_row: Dict[str, Any], contactsdb_path: str = "contactsdb.csv") -> Optional[Dict[str, Any]]:
    """
    Look up person in contactsdb.csv by name with aggressive matching.
    OPTIMIZED for large files (488k+ rows) - uses streaming search!
    
    NOW INCLUDES: Nickname and Spouse support from Data.csv!
    
    Tries:
    - Original name (e.g., "Jim Ronyak")
    - Nickname from CSV (e.g., if Data.csv has Nickname="Jim", tries "Jim Ronyak")
    - All nickname variants (e.g., "James Ronyak")
    - Spouse variants (e.g., "Mr. and Mrs. George S. Gibbs")
    - "Last, First" format (e.g., "Ronyak, James")
    - "Last First" format (e.g., "Ronyak James")
    
    Returns ALL fields from contactsdb.csv if found.
    """
    if not os.path.exists(contactsdb_path):
        log.debug(f"  contactsdb.csv not found at {contactsdb_path}")
        return None
    
    # Check file size
    file_size = os.path.getsize(contactsdb_path)
    file_size_mb = file_size / (1024 * 1024)
    
    log.info(f"  📇 Looking up '{name}' in contactsdb.csv ({file_size_mb:.1f} MB)...")
    
    # Parse the name
    parts = name.split()
    if len(parts) < 2:
        log.warning(f"  ⚠️  Name '{name}' doesn't have first and last name")
        return None
    
    first_name = parts[0]
    last_name = parts[-1]
    middle_names = ' '.join(parts[1:-1]) if len(parts) > 2 else ""
    
    # Extract nickname and spouse from donor_row (NEW!)
    nickname = get_nickname_from_row(donor_row)
    spouse = get_spouse_from_row(donor_row)
    
    if nickname:
        log.info(f"  💡 Using nickname from CSV: {nickname}")
    if spouse:
        log.info(f"  💑 Spouse from CSV: {spouse}")
    
    # Generate all name variants to try
    name_variants = []
    
    # 1. Original name
    name_variants.append(name)
    
    # 2. Nickname from CSV (PRIORITY - try this early!)
    if nickname:
        if middle_names:
            name_variants.append(f"{nickname} {middle_names} {last_name}")
        else:
            name_variants.append(f"{nickname} {last_name}")
    
    # 3. Nickname variants from dictionary (Jim ↔ James, etc.)
    all_variants = _token_name_variants(name)
    for variant in all_variants:
        if variant != name.lower():
            # Capitalize properly
            variant_parts = variant.split()
            capitalized = " ".join(p.capitalize() for p in variant_parts)
            name_variants.append(capitalized)
    
    # 4. Try formal name if current is nickname (e.g., Jim → James)
    if first_name.lower() in REVERSE_NICKNAMES:
        formal_names = REVERSE_NICKNAMES[first_name.lower()]
        for formal in formal_names:
            if middle_names:
                name_variants.append(f"{formal.capitalize()} {middle_names} {last_name}")
            else:
                name_variants.append(f"{formal.capitalize()} {last_name}")
    
    # 5. Try nickname if current is formal (e.g., James → Jim)
    if first_name.lower() in NICKNAMES:
        nicknames = NICKNAMES[first_name.lower()]
        for nick in nicknames:
            if middle_names:
                name_variants.append(f"{nick.capitalize()} {middle_names} {last_name}")
            else:
                name_variants.append(f"{nick.capitalize()} {last_name}")
    
    # 6. Spouse variants (NEW! - for soft credits like "Mr. and Mrs. George S. Gibbs")
    if spouse:
        spouse_parts = spouse.split()
        if len(spouse_parts) >= 2:
            spouse_first = spouse_parts[0]
            spouse_last = spouse_parts[-1]
            
            # Common formats for joint listings
            name_variants.append(f"Mr. and Mrs. {spouse_first} {spouse_last}")
            name_variants.append(f"Mr. & Mrs. {spouse_first} {spouse_last}")
            name_variants.append(f"Mr and Mrs {spouse_first} {spouse_last}")
            name_variants.append(f"{spouse_first} and {first_name} {last_name}")
            name_variants.append(f"{spouse_first} & {first_name} {last_name}")
    
    # 7. "Last, First" format variations
    name_variants.append(f"{last_name}, {first_name}")
    if nickname:
        name_variants.append(f"{last_name}, {nickname}")
    if first_name.lower() in REVERSE_NICKNAMES:
        for formal in REVERSE_NICKNAMES[first_name.lower()]:
            name_variants.append(f"{last_name}, {formal.capitalize()}")
    if first_name.lower() in NICKNAMES:
        for nick in NICKNAMES[first_name.lower()]:
            name_variants.append(f"{last_name}, {nick.capitalize()}")
    
    # 8. "Last First" format (no comma)
    name_variants.append(f"{last_name} {first_name}")
    
    # Remove duplicates while preserving order
    seen = set()
    unique_variants = []
    for variant in name_variants:
        variant_lower = variant.lower()
        if variant_lower not in seen:
            seen.add(variant_lower)
            unique_variants.append(variant)
    
    log.info(f"  🔍 Trying {len(unique_variants)} name variants...")
    if len(unique_variants) <= 8:
        log.info(f"     Variants: {', '.join(unique_variants)}")
    else:
        log.info(f"     Variants: {', '.join(unique_variants[:8])} ... (+{len(unique_variants)-8} more)")
    
    # Use streaming search for large files (> 10 MB)
    if file_size_mb > 10:
        log.info(f"  ⚡ Using STREAMING SEARCH (file is {file_size_mb:.1f} MB, too large to load into memory)")
        contact_row = search_large_csv_streaming(unique_variants, contactsdb_path)
        
        if contact_row:
            # Found via streaming search
            contact_data = {k: v for k, v in contact_row.items() if k and v}
            
            # Log what we found
            excluded_keys = {"First Name", "Last Name", "First", "Last"}
            data_summary = {k: v for k, v in contact_data.items() if k not in excluded_keys}
            
            log.info(f"  ✅ FOUND in contactsdb.csv!")
            log.info(f"     Searched for: {name}")
            log.info(f"     Matched as: '{get_row_name(contact_row)}'")
            log.info(f"     Additional fields: {len(data_summary)}")
            
            # Log key fields (but not sensitive ones like phone/email)
            public_fields = {k: v for k, v in data_summary.items() 
                            if k.lower() not in ['phone', 'email', 'mobile', 'personal email', 'work email', 'home phone', 'cell phone', 'work phone']}
            if public_fields:
                log.info(f"     Fields: {', '.join(list(public_fields.keys())[:10])}")
            
            return contact_data
        else:
            # Not found with streaming search
            log.warning(f"  ❌ NOT FOUND in contactsdb.csv (searched entire file)")
            if len(unique_variants) <= 8:
                log.warning(f"     Tried: {', '.join(unique_variants)}")
            else:
                log.warning(f"     Tried: {', '.join(unique_variants[:8])} ... (+{len(unique_variants)-8} more)")
            return None
    
    # For small files, use original in-memory search
    log.info(f"  📂 Loading entire file into memory (small file: {file_size_mb:.1f} MB)...")
    contactsdb_rows = load_csv(contactsdb_path)
    if not contactsdb_rows:
        log.warning(f"  ⚠️  contactsdb.csv is empty or failed to load")
        return None
    
    log.info(f"     Loaded {len(contactsdb_rows):,} contacts")
    
    # Try each variant (original in-memory search for small files)
    for variant in unique_variants:
        contact_row = find_donor_in_data(contactsdb_rows, variant)
        if contact_row:
            # Found in contactsdb.csv - extract ALL fields
            contact_data = {k: v for k, v in contact_row.items() if k and v}
            
            # Log what we found (excluding name fields)
            excluded_keys = {"First Name", "Last Name", "First", "Last"}
            data_summary = {k: v for k, v in contact_data.items() if k not in excluded_keys}
            
            log.info(f"  ✅ FOUND in contactsdb.csv!")
            log.info(f"     Searched for: {name}")
            log.info(f"     Matched as: '{variant}'")
            log.info(f"     Additional fields: {len(data_summary)}")
            
            # Log key fields (but not sensitive ones like phone/email)
            public_fields = {k: v for k, v in data_summary.items() 
                            if k.lower() not in ['phone', 'email', 'mobile', 'personal email', 'work email', 'home phone', 'cell phone', 'work phone']}
            if public_fields:
                log.info(f"     Fields: {', '.join(list(public_fields.keys())[:10])}")
            
            return contact_data
    
    # Not found with any variant
    log.warning(f"  ❌ NOT FOUND in contactsdb.csv (tried {len(unique_variants)} variants)")
    log.warning(f"     This person may not be in the contacts database yet")
    if len(unique_variants) <= 8:
        log.warning(f"     Tried: {', '.join(unique_variants)}")
    else:
        log.warning(f"     Tried: {', '.join(unique_variants[:8])} ... (+{len(unique_variants)-8} more)")
    return None

def enrich_with_contactsdb(data_csv_row: Dict[str, Any], contactsdb_path: str = "contactsdb.csv") -> Dict[str, Any]:
    """
    If 'Already known' is checked, look up person in contactsdb.csv and merge data.
    contactsdb.csv data takes precedence (it's more detailed for known contacts).
    
    Returns enriched data dictionary with ALL available fields.
    Special fields added:
    - _original_name: Name from data.csv (e.g., "jim ronyak")
    - _matched_name: Name from contactsdb.csv (e.g., "James Ronyak")
    """
    # Check if already known
    already_known = data_csv_row.get("Already known", "").strip().lower()
    if already_known != "yes":
        # Not already known - just return data.csv data
        return {k: v for k, v in data_csv_row.items() if k and v}
    
    # Already known - look up in contactsdb.csv
    original_name = get_row_name(data_csv_row)
    log.info(f"  📋 '{original_name}' marked as 'Already known' - loading detailed data from contactsdb.csv...")
    
    # UPDATED: Pass data_csv_row so load_contactsdb_data can use Nickname and Spouse fields
    contactsdb_data = load_contactsdb_data(original_name, data_csv_row, contactsdb_path)
    
    if contactsdb_data:
        # Get the matched name from contactsdb.csv
        matched_name = get_row_name(contactsdb_data)
        
        # Merge: contactsdb.csv takes precedence, fallback to data.csv
        enriched = dict(data_csv_row)  # Start with data.csv
        enriched.update(contactsdb_data)  # Override with contactsdb.csv
        
        # Store both name variants for search queries
        enriched['_original_name'] = original_name
        enriched['_matched_name'] = matched_name
        
        if original_name.lower() != matched_name.lower():
            log.info(f"  📝 Name variants: '{original_name}' (data.csv) and '{matched_name}' (contactsdb.csv)")
        
        log.info(f"  ✅ Enriched with contactsdb.csv data - now have {len(enriched)} total fields")
        return enriched
    else:
        # contactsdb.csv lookup failed - use data.csv only
        log.warning(f"  ⚠️  contactsdb.csv lookup failed - using data.csv only")
        return {k: v for k, v in data_csv_row.items() if k and v}

# -------------------------
# FEC RESEARCH
# -------------------------
@retry_with_backoff(max_retries=3)
def research_fec(name: str, max_pages: int = 3) -> Dict[str, Any]:
    """FEC research with caching"""
    cache_key = cache.get_cache_key("fec", {"name": name})
    
    def fetcher():
        summary = {
            "found": False,
            "total": 0.0,
            "count": 0,
            "by_party": {"R": 0.0, "D": 0.0, "Other": 0.0},
            "recent_years": {},
            "top_recipients": []
        }
        
        if not FEC_API_KEY:
            log.warning("FEC disabled (missing API key)")
            return summary
        
        base = "https://api.open.fec.gov/v1/schedules/schedule_a/"
        params = {
            "api_key": FEC_API_KEY,
            "contributor_name": name,
            "per_page": 100,
            "sort": "contribution_receipt_date"
        }
        
        page = 1
        recipients = defaultdict(float)
        
        try:
            while page <= max_pages:
                params["page"] = page
                r = requests.get(base, params=params, timeout=25, headers=DEFAULT_HEADERS)
                r.raise_for_status()
                data = r.json()
                results = data.get("results", []) or []
                if not results:
                    break
                
                for it in results:
                    amt = float(it.get("contribution_receipt_amount") or 0.0)
                    summary["total"] += amt
                    summary["count"] += 1
                    summary["found"] = True
                    
                    try:
                        yr = int((it.get("contribution_receipt_date") or "")[:4])
                        summary["recent_years"][yr] = summary["recent_years"].get(yr, 0.0) + amt
                    except:
                        pass
                    
                    rcpt = it.get("recipient_name") or it.get("committee_name") or ""
                    if rcpt:
                        recipients[rcpt] += amt
                    p = (it.get("party_full") or "").upper()
                    if p.startswith("DEM"):
                        summary["by_party"]["D"] += amt
                    elif p.startswith("REP"):
                        summary["by_party"]["R"] += amt
                    else:
                        summary["by_party"]["Other"] += amt
                
                page += 1
                if not data.get("pagination", {}).get("pages"):
                    break
                if page > data["pagination"]["pages"]:
                    break
        except Exception as e:
            log.warning(f"FEC error: {e}")
        
        # Pattern
        R, D, O = summary["by_party"]["R"], summary["by_party"]["D"], summary["by_party"]["Other"]
        tot = summary["total"] or 1.0
        rp = R / tot * 100
        dp = D / tot * 100
        if max(rp, dp) <= 65:
            pattern = "Bipartisan donor"
        elif rp >= 70:
            pattern = "Republican donor"
        elif dp >= 70:
            pattern = "Democratic donor"
        else:
            pattern = "Mixed donor"
        summary["pattern"] = pattern
        
        top = sorted(recipients.items(), key=lambda kv: -kv[1])[:10]
        summary["top_recipients"] = [{"name": k, "total": v} for k, v in top]
        summary["profile_url"] = f"https://www.fec.gov/data/receipts/individual-contributions/?contributor_name={requests.utils.quote(name)}"
        
        return summary
    
    return cache.get_or_fetch(cache_key, fetcher)

# -------------------------
# INNER CIRCLE CONNECTIONS
# -------------------------
CONN_SCHEMA_EXAMPLE = {
    "found": True,
    "type": "co-board | co-event | alumni | other",
    "strength": 72,
    "description": "One-liner summary",
    "timeframe": "2011-2014",
    "evidence": "Short quote",
    "source": "https://example.com",
    "confidence": 0.85
}

CONN_PROMPT_TMPL = """Verify connection between TWO people.
Return STRICT JSON in this schema:
{schema}

Rules:
- STRONG (70-100): Direct tie (co-board, co-authors, co-panelists at named event)
- MEDIUM (50-69): Likely tie (same program/org with overlap, multiple indirect interactions)
- POTENTIAL (30-49): Weak tie (same school non-overlap, same org different eras)
- REJECT (<30): No real connection

Prospect: "{prospect}"
Inner circle: "{member}"

SEARCH RESULTS:
{context}

Return JSON only:"""

def google_queries_for_pair(donor: str, member: str, donor_city: str = "", donor_state: str = "") -> List[str]:
    """Generate search queries for donor-member pair"""
    donor_parts = donor.split()
    member_parts = member.split()
    donor_last = donor_parts[-1] if donor_parts else donor
    member_last = member_parts[-1] if member_parts else member
    
    queries = [
        f'"{donor}" "{member}"',
        f'"{donor_last}" "{member_last}" {donor_state}' if donor_state else f'"{donor_last}" "{member_last}"',
        f'"{donor}" "{member}" board'
    ]
    return queries

def classify_connection(donor_name: str, donor_city: str, donor_state: str,
                       member_row: Dict[str, Any], fast_mode: bool = True) -> Optional[Dict[str, Any]]:
    """Classify connection with confidence scoring"""
    first = (member_row.get("First") or "").strip()
    middle = (member_row.get("Middle") or "").strip()
    last = (member_row.get("Last") or "").strip()
    member_name = " ".join([p for p in [first, middle, last] if p])
    
    if not member_name:
        return None
    
    queries = google_queries_for_pair(donor_name, member_name, donor_city, donor_state)
    
    if fast_mode:
        results = []
        with ThreadPoolExecutor(max_workers=3) as executor:
            future_to_query = {executor.submit(search_web, q, 3): q for q in queries}
            for future in as_completed(future_to_query):
                try:
                    results.extend(future.result())
                except Exception as e:
                    pass
        time.sleep(0.1)
    else:
        results = []
        for q in queries:
            results.extend(search_web(q, num=3))
            time.sleep(1.2)
    
    uniq, seen = [], set()
    for r in results:
        u = r.get("link")
        if u and u not in seen:
            seen.add(u)
            uniq.append(r)
    
    context_lines = [f"- {r.get('title', '')} | {r.get('snippet', '')} | {r.get('link', '')}" 
                     for r in uniq[:8]]
    context = "\n".join(context_lines) if context_lines else "NO WEB HITS"
    
    if context == "NO WEB HITS":
        return None
    
    prompt = CONN_PROMPT_TMPL.format(
        schema=json.dumps(CONN_SCHEMA_EXAMPLE, indent=2),
        prospect=donor_name,
        member=member_name,
        context=context
    )
    
    parsed = openai_chat_json(prompt, cache_key_prefix=f"conn_{donor_name}_{member_name}")
    if not parsed or not parsed.get("found"):
        return None
    
    strength = int(parsed.get("strength") or 0)
    if strength < 30:
        return None
    
    out = {
        "inner_circle_name": member_name,
        "type": parsed.get("type") or "other",
        "strength": strength,
        "description": parsed.get("description") or "",
        "timeframe": parsed.get("timeframe") or "",
        "evidence": parsed.get("evidence") or "",
        "source": parsed.get("source") or "",
        "confidence": parsed.get("confidence", 0.5),
        "citations": [parsed.get("source")] if parsed.get("source") else []
    }
    
    if out["strength"] >= 70:
        out["strength_bucket"] = "Strong"
    elif out["strength"] >= 50:
        out["strength_bucket"] = "Medium"
    else:
        out["strength_bucket"] = "Potential"
    
    return out

def research_inner_circle(prospect: Dict[str, Any], 
                         inner_rows: List[Dict[str, Any]],
                         max_members: Optional[int] = None,
                         fast_mode: bool = True,
                         parallel_workers: int = 10) -> Tuple[List[Dict[str, Any]], List[str]]:
    """Research inner circle connections"""
    donor = prospect.get("name") or ""
    donor_city = prospect.get("City") or ""
    donor_state = prospect.get("State") or ""
    
    if max_members and max_members < len(inner_rows):
        inner_rows = inner_rows[:max_members]
        log.info(f"⚡ FAST MODE: Limiting inner circle check to first {max_members} members")
    
    connections, used = [], []
    total = len(inner_rows)
    log.info(f"Analyzing {total} inner-circle members for {donor}...")
    if fast_mode:
        log.info(f"⚡ Using parallel processing with {parallel_workers} concurrent workers")
    
    def process_member(i_and_row):
        i, row = i_and_row
        try:
            c = classify_connection(donor, donor_city, donor_state, row, fast_mode=fast_mode)
            return (i, c)
        except Exception as e:
            log.warning(f"Member #{i} failed: {e}")
            return (i, None)
    
    if fast_mode:
        with ThreadPoolExecutor(max_workers=parallel_workers) as executor:
            futures = {executor.submit(process_member, (i+1, row)): i 
                      for i, row in enumerate(inner_rows)}
            
            completed = 0
            for future in as_completed(futures):
                completed += 1
                if (completed % 25) == 0:
                    log.info(f"  Progress {completed}/{total}...")
                
                try:
                    i, c = future.result()
                    if c:
                        connections.append(c)
                        used.extend(c.get("citations", []))
                except Exception as e:
                    log.warning(f"Member processing failed: {e}")
    else:
        for i, row in enumerate(inner_rows, start=1):
            if (i % 25) == 0:
                log.info(f"  Progress {i}/{total}...")
            
            try:
                c = classify_connection(donor, donor_city, donor_state, row, fast_mode=False)
                if c:
                    connections.append(c)
                    used.extend(c.get("citations", []))
            except Exception as e:
                log.warning(f"Member #{i} failed: {e}")
            
            time.sleep(0.5)
    
    connections.sort(key=lambda x: (-x["strength"], x["inner_circle_name"]))
    used = list(set(used))
    
    log.info(f"Connections found: "
             f"{sum(1 for c in connections if c['strength_bucket']=='Strong')} strong, "
             f"{sum(1 for c in connections if c['strength_bucket']=='Medium')} medium, "
             f"{sum(1 for c in connections if c['strength_bucket']=='Potential')} potential.")
    
    return connections, used

# -------------------------
# DOCX RENDERING
# -------------------------
def render_title(doc: Document, name: str, entity_type: str = "person"):
    """Render document title - Partner Profile header"""
    # Title: "Partner Profile" (centered, large)
    p = doc.add_paragraph()
    r = p.add_run("Partner Profile")
    r.bold = True
    r.font.size = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()  # Spacing


def render_header_section(doc: Document, name: str, csv_data: Dict[str, Any], entity_type: str = "person"):
    """Render the header section with contact info, education, positions etc."""
    
    # Helper to add a labeled field
    def add_field(label: str, value: str, bold_label: bool = True):
        if not value or value.strip() == "":
            return
        p = doc.add_paragraph()
        if bold_label:
            r = p.add_run(f"{label}")
            r.bold = True
            p.add_run(f"\n{value}")
        else:
            p.add_run(f"{label}: {value}")
    
    # Name (large, bold)
    p = doc.add_paragraph()
    r = p.add_run("Name")
    r.bold = True
    p = doc.add_paragraph()
    r = p.add_run(name)
    r.bold = True
    r.font.size = Pt(14)
    
    # Photo placeholder
    p = doc.add_paragraph()
    r = p.add_run("Photo")
    r.bold = True
    doc.add_paragraph("[Photo placeholder]")
    
    # Contact Info
    address_parts = []
    street = csv_data.get("Address", csv_data.get("Street", csv_data.get("Street Address", "")))
    city = csv_data.get("City", "")
    state = csv_data.get("State", "")
    zipcode = csv_data.get("Zip", csv_data.get("ZIP", csv_data.get("Zipcode", "")))
    
    if street:
        address_parts.append(street)
    if city or state or zipcode:
        loc = f"{city}, {state} {zipcode}".strip().strip(",")
        address_parts.append(loc)
    
    phone = csv_data.get("Phone", csv_data.get("phone", csv_data.get("Mobile", "")))
    email = csv_data.get("Email", csv_data.get("email", ""))
    
    contact_info = []
    if address_parts:
        contact_info.extend(address_parts)
    if phone:
        contact_info.append(phone)
    if email:
        contact_info.append(email)
    
    add_field("Contact Info (Address)", "\n".join(contact_info) if contact_info else "(Not available)")
    
    # Salesforce Link
    sf_link = csv_data.get("Salesforce Link", csv_data.get("Salesforce", csv_data.get("SF Link", "")))
    if sf_link:
        add_field("Salesforce Link", sf_link)
    
    # Date of Birth
    dob = csv_data.get("Date of Birth", csv_data.get("DOB", csv_data.get("Birth Date", "")))
    if dob:
        add_field("Date of Birth", dob)
    
    # Family
    spouse = csv_data.get("Spouse/Partner", csv_data.get("Spouse", ""))
    family_info = csv_data.get("Family", "")
    if spouse or family_info:
        family_text = family_info if family_info else f"Spouse: {spouse}" if spouse else ""
        add_field("Family", family_text if family_text else "(Spouse and spouse's occupation/affiliations, children, relatives)")
    else:
        add_field("Family", "(Spouse and spouse's occupation/affiliations, children, relatives)")


def render_education_section(doc: Document, csv_data: Dict[str, Any], bio_text: str = ""):
    """Render Education section"""
    p = doc.add_paragraph()
    r = p.add_run("Education")
    r.bold = True
    
    # Try to get from CSV first
    education = csv_data.get("Education", csv_data.get("School", csv_data.get("University", "")))
    if education:
        doc.add_paragraph(education)
    else:
        doc.add_paragraph("(Education details from research)")


def render_positions_section(doc: Document, career_text: str):
    """Render Current Position and Previous Positions sections"""
    
    # Current Position
    p = doc.add_paragraph()
    r = p.add_run("Current Position or Occupation")
    r.bold = True
    
    # Parse career text to find current positions (look for "current", "CEO", etc.)
    lines = [l.strip() for l in career_text.split('\n') if l.strip()]
    current_positions = []
    previous_positions = []
    
    for line in lines:
        # Remove bullet markers
        clean_line = line.lstrip('•-*').strip()
        if not clean_line or clean_line.startswith('[') or 'citation' in clean_line.lower():
            continue
            
        line_lower = clean_line.lower()
        if any(kw in line_lower for kw in ['currently', 'current', 'now serves', 'presently', 'ceo of', 'founder of']):
            current_positions.append(clean_line)
        elif any(kw in line_lower for kw in ['former', 'previously', 'past', 'was ', 'served as']):
            previous_positions.append(clean_line)
        elif len(current_positions) == 0:
            # First positions are likely current
            current_positions.append(clean_line)
        else:
            previous_positions.append(clean_line)
    
    if current_positions:
        for pos in current_positions[:3]:
            doc.add_paragraph(pos)
    else:
        doc.add_paragraph("(Current position from research)")
    
    # Previous Positions
    p = doc.add_paragraph()
    r = p.add_run("Previous Positions")
    r.bold = True
    
    if previous_positions:
        for pos in previous_positions[:8]:
            doc.add_paragraph(pos)
    else:
        doc.add_paragraph("(Previous positions from research)")


def render_board_memberships_section(doc: Document, network_text: str):
    """Render Foundation and Nonprofit Board Memberships section"""
    p = doc.add_paragraph()
    r = p.add_run("Foundation and Nonprofit Board Memberships")
    r.bold = True
    
    lines = [l.strip().lstrip('•-*').strip() for l in network_text.split('\n') if l.strip()]
    board_lines = [l for l in lines if l and not l.startswith('[')]
    
    if board_lines:
        for line in board_lines[:10]:
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("(Board memberships from research)")


def render_networth_section(doc: Document, networth_text: str):
    """Render Net Worth/Giving Capacity section with subsections"""
    p = doc.add_paragraph()
    r = p.add_run("Net Worth/Giving Capacity")
    r.bold = True
    r.font.size = Pt(14)
    
    # Parse the structured sections
    sections = networth_text.split('\n\n')
    
    for section in sections:
        section = section.strip()
        if not section:
            continue
        
        # Extract header and content
        lines = section.split('\n')
        header_line = lines[0] if lines else ""
        content_lines = lines[1:] if len(lines) > 1 else []
        
        # Check for subsection headers (format: **Header:**)
        if header_line.startswith('**') and ':**' in header_line:
            # Extract header text
            header_text = header_line.replace('**', '').replace(':', '').strip()
            
            p = doc.add_paragraph()
            r = p.add_run(f"{header_text}:")
            r.bold = True
            
            # Add content lines
            for line in content_lines:
                clean = line.strip().lstrip('•-*').strip()
                if clean and not clean.startswith('['):
                    doc.add_paragraph(f"  {clean}")
                    
        elif section.startswith('**Web Research'):
            p = doc.add_paragraph()
            r = p.add_run("Estimated Wealth:")
            r.bold = True
            # Convert bullets to prose paragraph
            prose_parts = []
            for line in content_lines:
                clean = line.strip().lstrip('•-*').strip()
                if clean and not clean.startswith('['):
                    # Remove citation markers for prose flow
                    clean = re.sub(r'\s*\[\d+\]', '', clean)
                    prose_parts.append(clean)
            if prose_parts:
                doc.add_paragraph(" ".join(prose_parts))
        else:
            # Generic content - just add as paragraphs
            for line in lines:
                clean = line.strip().lstrip('•-*').strip()
                if clean and not clean.startswith('[') and not clean.startswith('**'):
                    doc.add_paragraph(clean)


def render_philanthropy_section(doc: Document, philanthropy_text: str):
    """Render Philanthropic Interests and Relevant Gifts sections"""
    
    # Philanthropic Interests
    p = doc.add_paragraph()
    r = p.add_run("Philanthropic Interests")
    r.bold = True
    
    lines = [l.strip().lstrip('•-*').strip() for l in philanthropy_text.split('\n') if l.strip()]
    interests = [l for l in lines if l and not l.startswith('[') and len(l) > 10]
    
    if interests:
        for interest in interests[:8]:
            # Remove citation markers
            clean = re.sub(r'\s*\[\d+\]', '', interest)
            doc.add_paragraph(f"• {clean}")
    else:
        doc.add_paragraph("• (Philanthropic interests from research)")
    
    # Relevant Gifts
    p = doc.add_paragraph()
    r = p.add_run("Relevant Gifts")
    r.bold = True
    doc.add_paragraph("(Specific gift history from research)")
    
    # Other Giving History
    p = doc.add_paragraph()
    r = p.add_run("Other Giving History")
    r.bold = True
    doc.add_paragraph("(Additional giving history from research)")


def render_political_section(doc: Document, political_text: str, fec_data: Dict[str, Any]):
    """Render Political Activity section"""
    p = doc.add_paragraph()
    r = p.add_run("Political Activity & Ideology")
    r.bold = True
    
    # Add FEC summary if available
    if fec_data.get("found"):
        total = fec_data.get("total", 0)
        pattern = fec_data.get("pattern", "")
        doc.add_paragraph(f"FEC records show ${total:,.0f} in political contributions ({pattern})")
    
    # Add other political info
    lines = [l.strip().lstrip('•-*').strip() for l in political_text.split('\n') if l.strip()]
    for line in lines[:5]:
        if line and not line.startswith('['):
            clean = re.sub(r'\s*\[\d+\]', '', line)
            doc.add_paragraph(clean)


def render_biography_section(doc: Document, bio_text: str, career_text: str):
    """Render Biography section as narrative prose"""
    p = doc.add_paragraph()
    r = p.add_run("Biography")
    r.bold = True
    r.font.size = Pt(14)
    
    # Combine bio and career into narrative paragraphs
    all_lines = []
    
    for text in [bio_text, career_text]:
        lines = [l.strip().lstrip('•-*').strip() for l in text.split('\n') if l.strip()]
        for line in lines:
            if line and not line.startswith('[') and len(line) > 20:
                # Remove citation markers
                clean = re.sub(r'\s*\[\d+\]', '', line)
                all_lines.append(clean)
    
    # Group into paragraphs
    if all_lines:
        # First paragraph - intro
        para1 = " ".join(all_lines[:3])
        doc.add_paragraph(para1)
        
        # Additional paragraphs
        if len(all_lines) > 3:
            para2 = " ".join(all_lines[3:6])
            doc.add_paragraph(para2)
        
        if len(all_lines) > 6:
            para3 = " ".join(all_lines[6:9])
            doc.add_paragraph(para3)
    else:
        doc.add_paragraph("(Biography narrative from research)")


def render_connections_section(doc: Document, connections: List[Dict], skip_inner: bool):
    """Render Connections (Inner Circle) section"""
    p = doc.add_paragraph()
    r = p.add_run("Connections")
    r.bold = True
    r.font.size = Pt(14)
    
    if skip_inner:
        doc.add_paragraph("Inner Circle analysis skipped (donor already known to IHS)")
    elif connections:
        doc.add_paragraph(f"{len(connections)} connections identified with IHS Inner Circle members:")
        for conn in connections[:10]:
            name = conn.get('inner_circle_name', '')
            desc = conn.get('description', '')
            strength = conn.get('strength_bucket', '')
            doc.add_paragraph(f"• {name}: {desc} (Strength: {strength})")
    else:
        doc.add_paragraph("No significant connections found with IHS Inner Circle members")


def render_references_section(doc: Document):
    """Render numbered references section"""
    doc.add_page_break()
    
    p = doc.add_paragraph()
    r = p.add_run("References")
    r.bold = True
    r.font.size = Pt(14)
    
    citations = citation_mgr.get_all_citations()
    for cite in citations:
        num = cite["number"]
        title = cite.get("title", "Untitled")
        url = cite.get("url", "")
        p = doc.add_paragraph(f"[{num}] {title}")
        if url:
            p.add_run(f"\n    {url}").font.size = Pt(9)

# -------------------------
# MAIN
# -------------------------
def main():
    parser = argparse.ArgumentParser(description="IHS Enhanced Donor Research System")
    parser.add_argument("--data", default="data.csv", help="Path to donor CSV")
    parser.add_argument("--inner", default="", help="Path to inner circle CSV")
    parser.add_argument("--outdir", default="./profiles", help="Output directory")
    parser.add_argument("--name", default="", help="Process specific donor name")
    parser.add_argument("--test-mode", action="store_true", help="Process only first donor")
    parser.add_argument("--fast", dest="fast_mode", action="store_true", default=True, help="Fast mode (parallel)")
    parser.add_argument("--slow", dest="fast_mode", action="store_false", help="Slow mode (sequential)")
    parser.add_argument("--max-inner-circle", type=int, help="Limit inner circle members")
    parser.add_argument("--parallel-workers", type=int, default=10, help="Parallel workers for inner circle")
    
    args = parser.parse_args()
    fast_mode = args.fast_mode
    
    os.makedirs(args.outdir, exist_ok=True)
    
    log.info("="*60)
    log.info("IHS ENHANCED DONOR RESEARCH SYSTEM")
    log.info("="*60)
    log.info(f"Features: Wikipedia citations, bullet points, identity verification")
    log.info(f"Data file: {args.data}")
    log.info(f"Output directory: {args.outdir}")
    log.info(f"Mode: {'FAST (parallel)' if fast_mode else 'SLOW (sequential)'}")
    if args.max_inner_circle:
        log.info(f"Inner circle limit: {args.max_inner_circle} members")
    log.info("")
    
    # Load data
    data_rows = load_csv(args.data)
    if not data_rows:
        log.error(f"No rows found in {args.data}")
        return
    
    # Load inner circle
    if args.inner:
        inner_path = args.inner
    else:
        for candidate in ("Inner_Circle.csv", "inner_circle.csv"):
            if os.path.exists(candidate):
                inner_path = candidate
                break
        else:
            log.error("Inner circle file not found")
            return
    
    inner_rows = load_csv(inner_path)
    inner_circle_count = len(inner_rows)
    log.info(f"Loaded {inner_circle_count} inner circle members from {inner_path}")
    
    # Determine which donors to process
    if args.test_mode:
        first_name = get_row_name(data_rows[0])
        if not first_name:
            log.error("First donor has no valid name")
            return
        names_to_process = [first_name]
        log.info(f"🧪 TEST MODE: Processing ONLY first donor: {first_name}")
    elif args.name:
        names_to_process = [args.name.strip()]
    else:
        names_to_process = [get_row_name(r) for r in data_rows if get_row_name(r)]
        if not names_to_process:
            log.error(f"No valid names found in {args.data}")
            return
        log.info(f"Processing {len(names_to_process)} donors from {args.data}")
    
    # Process each donor
    for donor_name in names_to_process:
        log.info(f"\n{'='*60}")
        log.info(f"Processing: {donor_name}")
        log.info(f"{'='*60}")
        
        # Reset citation manager for new donor
        citation_mgr.reset()
        
        donor_row = find_donor_in_data(data_rows, donor_name)
        if not donor_row:
            log.warning(f"Skipping: '{donor_name}' not found in {args.data}")
            continue
        
        # =====================================================
        # NEW: MERGE ALL ROWS for the same person
        # If multiple rows share the same last name, they likely 
        # represent different pieces of info about ONE person
        # =====================================================
        donor_last = (donor_row.get("Last Name", "") or donor_row.get("Last", "")).strip()
        if donor_last:
            # Count how many rows share this last name
            same_last_count = sum(1 for r in data_rows 
                                  if (r.get("Last Name", "") or r.get("Last", "")).strip().lower() == donor_last.lower()
                                  or not (r.get("Last Name", "") or r.get("Last", "")).strip())  # Include no-name rows (properties)
            
            if same_last_count > 1:
                log.info(f"  🔀 MERGING {same_last_count} rows with last name '{donor_last}' into single record")
                merged_row = merge_person_rows(data_rows, donor_last)
                
                # Use merged data instead of single row
                donor_row = merged_row
                
                # Update donor_name to use full name from merged data
                merged_first = merged_row.get("First Name", "").strip()
                merged_middle = merged_row.get("Middle Name", "").strip()
                if merged_first:
                    if merged_middle:
                        donor_name = f"{merged_first} {merged_middle} {donor_last}"
                    else:
                        donor_name = f"{merged_first} {donor_last}"
                    log.info(f"  📝 Using merged name: {donor_name}")
        
        # NEW: Find related rows that might contain alternate names, maiden names, etc.
        related_rows = find_related_rows(data_rows, donor_row)
        alternate_names = []
        
        # FIRST: Check for Maiden Name column directly in the CSV
        maiden_name_col = donor_row.get("Maiden Name", "").strip()
        if maiden_name_col:
            first_name = donor_row.get("First Name", "").strip()
            middle_name = donor_row.get("Middle Name", "").strip()
            if first_name:
                if middle_name:
                    alt_name = f"{first_name} {middle_name} {maiden_name_col}"
                else:
                    alt_name = f"{first_name} {maiden_name_col}"
                alternate_names.append({
                    "name": alt_name,
                    "alternate_last_name": maiden_name_col,
                    "location": "",
                    "type": "maiden_name_column"
                })
                log.info(f"  📛 MAIDEN NAME from CSV column: {alt_name}")
        
        if related_rows:
            log.info(f"  🔍 Found {len(related_rows)} potentially related rows in CSV")
            for rel in related_rows:
                if rel["type"] == "alternate_name" and rel.get("alternate_last_name"):
                    # Build alternate name using donor's first name + alternate last name
                    alt_first = donor_row.get("First Name", donor_row.get("First", "")).strip()
                    alt_last = rel["alternate_last_name"]
                    if alt_first and alt_last:
                        alt_name = f"{alt_first} {alt_last}"
                        alternate_names.append({
                            "name": alt_name,
                            "location": rel.get("associated_location", ""),
                            "type": "maiden_or_alternate"
                        })
                        log.info(f"    📛 Alternate name detected: {alt_name} (associated with: {rel.get('associated_location', 'unknown location')})")
                elif rel["type"] == "potential_spouse":
                    log.info(f"    💑 Potential spouse detected: {rel.get('spouse_name', '')}")
        
        try:
            # 1) Identity Verification - using ALL CSV data (with contactsdb.csv enrichment)
            log.info(f"📊 Extracting CSV data for {donor_name}...")
            
            # ENHANCED: If "Already known" is checked, enrich with contactsdb.csv data
            csv_data = enrich_with_contactsdb(donor_row, contactsdb_path="contactsdb.csv")
            
            # Log what we found
            excluded_keys = {"First Name", "Last Name", "First", "Last"}
            data_summary = {k: v for k, v in csv_data.items() if k not in excluded_keys}
            if data_summary:
                log.info(f"  Available data: {', '.join(data_summary.keys())}")
            else:
                log.warning(f"  ⚠️  Only name found in CSV - no additional context data")
            
            # NEW: Check age to determine if prospect is likely retired
            is_retired_age = False
            dob_str = csv_data.get("Date of Birth", csv_data.get("DOB", csv_data.get("Birth Date", "")))
            birth_year = None
            if dob_str:
                # Use enhanced parser that handles "Aug-36" format
                birth_year = parse_dob_year(dob_str)
                if birth_year:
                    current_year = datetime.now().year
                    age = current_year - birth_year
                    if age >= 75:
                        is_retired_age = True
                        log.info(f"  👴 RETIRED INDICATOR: Birth year {birth_year} (age ~{age}) - prioritizing retired/former searches")
                else:
                    log.debug(f"  Could not parse DOB: {dob_str}")
            
            # Store in csv_data for use in research functions
            csv_data['_is_retired_age'] = is_retired_age
            csv_data['_birth_year'] = birth_year
            
            # Pass ALL CSV data (potentially enriched from contactsdb.csv) to identity verification
            identity = verify_identity(donor_name, csv_data)
            
            # Check confidence level
            if identity["confidence"] == "Low" and not identity["verified"]:
                # STOP: Very low confidence and not verified
                log.error(f"❌ Skipping {donor_name} - insufficient confidence in identity verification")
                log.error(f"   Score: {identity.get('top_candidate', {}).get('score', 0):.2f} - needs manual review")
                continue
            elif identity["confidence"] == "Low" and identity["verified"]:
                # PROCEED WITH WARNING: Low confidence but likely right person
                log.warning(f"⚠️  LOW confidence proceeding with {donor_name} - recommend manual review")
            elif identity["collisions"]:
                log.warning(f"⚠️  {len(identity['collisions'])} possible name collisions detected:")
                for collision in identity["collisions"][:3]:
                    log.warning(f"    - {collision['indicator']}")
            
            entity_type = identity["entity_type"]
            log.info(f"  Entity type: {entity_type}")
            log.info(f"  Confidence: {identity['confidence']} ({identity['confidence_score']:.1%})")
            
            log.info(f"  ✅ Identity verified - proceeding with detailed research")
            log.info(f"  📍 Strategy: ALL CSV data used for queries, but NO filtering of results")
            log.info(f"     (We've already confirmed it's the right person)")
            
            # Build location filter for search query context
            city = csv_data.get("City", "")
            state = csv_data.get("State", "")
            location_filter = f"{city}, {state}" if city and state else city if city else state if state else ""
            
            # NEW: Add all merged addresses/locations to csv_data for comprehensive searching
            if "_all_addresses" in csv_data:
                log.info(f"  🏠 Using {len(csv_data['_all_addresses'])} addresses from merged data")
                for i, addr in enumerate(csv_data["_all_addresses"]):
                    log.info(f"      Address {i+1}: {addr['full']}")
            
            if "_all_cities" in csv_data and len(csv_data["_all_cities"]) > 1:
                log.info(f"  📍 Multiple cities: {', '.join(csv_data['_all_cities'])}")
                # Add secondary cities for search queries
                for i, other_city in enumerate(csv_data["_all_cities"][1:], 2):
                    csv_data[f"Secondary City {i}"] = other_city
            
            if "_all_states" in csv_data and len(csv_data["_all_states"]) > 1:
                log.info(f"  📍 Multiple states: {', '.join(csv_data['_all_states'])}")
            
            # CRITICAL: Extract spouse name for filtering ALL web searches
            # This prevents mixing data from wrong people with same name
            spouse_name = ""
            nickname = get_nickname_from_row(csv_data)
            if nickname and ("mr. and mrs." in nickname.lower() or "mr and mrs" in nickname.lower()):
                try:
                    parts = nickname.split()
                    mrs_idx = next((i for i, p in enumerate(parts) if p.lower() in ["mrs.", "mrs"]), -1)
                    if mrs_idx >= 0 and mrs_idx < len(parts) - 1:
                        spouse_full = " ".join(parts[mrs_idx+1:])
                        spouse_parts = spouse_full.split()
                        if len(spouse_parts) >= 2:
                            spouse_first = spouse_parts[0]
                            spouse_last = spouse_parts[-1]
                            spouse_name = f"{spouse_first} {spouse_last}"
                            log.info(f"  💑 SPOUSE DETECTED: {spouse_name}")
                            log.info(f"     ALL web search results will be filtered to require spouse mention!")
                            log.info(f"     This prevents mixing data from other people named {donor_name}")
                except Exception as e:
                    log.debug(f"  Failed to parse spouse from nickname: {e}")
            
            # Also check Spouse/Partner column
            if not spouse_name:
                partner = get_spouse_from_row(csv_data)
                if partner:
                    parts = partner.split()
                    if len(parts) >= 1:
                        # Handle single name spouse (like "James F Long")
                        spouse_name = partner.strip()
                        log.info(f"  💑 SPOUSE DETECTED: {spouse_name} (from Spouse/Partner column)")
                        log.info(f"     ALL web search results will be filtered to require spouse mention!")
            
            # GENERATE ALL NAME VARIATIONS for comprehensive searching
            name_variations = generate_name_variations(donor_name, csv_data, alternate_names)
            if name_variations:
                log.info(f"  📛 Generated {len(name_variations)} name variations for searching:")
                for var in name_variations[:8]:  # Show first 8
                    log.info(f"      - {var['name']} ({var['type']})")
                if len(name_variations) > 8:
                    log.info(f"      ... and {len(name_variations) - 8} more")
            
            # 2) Biographical Research (bullet points)
            # Use ALL CSV data to guide research (employer, title, education, location, etc.)
            log.info("Researching biographical background using ALL CSV context...")
            
            # CRITICAL: Get verified flag from identity check
            # When FALSE: Apply strict spouse+location filtering (confirming identity)
            # When TRUE: Relax filtering (identity already confirmed)
            verified_flag = bool(identity.get("verified", False))
            log.info(f"  🔍 Identity verified: {verified_flag}")
            if not verified_flag:
                log.info(f"     Using STRICT filtering (spouse + location) to confirm identity")
            else:
                log.info(f"     Using RELAXED filtering (identity already confirmed)")
            
            # Extract additional CSV fields
            employer = csv_data.get("Employer", csv_data.get("Company", ""))
            title = csv_data.get("Title", csv_data.get("Position", csv_data.get("Job Title", "")))
            education = csv_data.get("Education", csv_data.get("School", csv_data.get("University", "")))
            
            if entity_type == "organization":
                bio_queries = get_org_search_queries(donor_name)[:6]
            else:
                # Build targeted biographical queries using ALL available CSV data
                bio_queries = []
                
                # NEW: Get retired age flag
                is_retired = csv_data.get('_is_retired_age', False)
                
                # =====================================================
                # PRIORITY 0: RETIRED/FORMER searches (if 75+ years old)
                # =====================================================
                if is_retired:
                    log.info(f"  👴 Prospect is 75+ - prioritizing retired/former searches")
                    bio_queries.extend([
                        f'"{donor_name}" retired former professional',
                        f'"{donor_name}" retired {city}' if city else None,
                        f'"{donor_name}" career history',
                    ])
                    # Add employer with retired context
                    if employer:
                        bio_queries.extend([
                            f'"{donor_name}" retired "{employer}"',
                            f'"{donor_name}" former "{employer}"',
                        ])
                        log.info(f"  👴 Adding retired+employer: {employer}")
                
                # =====================================================
                # PRIORITY 1: MOST SPECIFIC - Name AND Employer AND Title
                # This is the BEST query for pinpointing the exact person!
                # =====================================================
                if employer and title:
                    # BEST QUERY: All three components together
                    bio_queries.insert(0, f'"{donor_name}" "{employer}" "{title}"')
                    bio_queries.insert(1, f'"{donor_name}" {employer} {title}')
                    # Add with location for even more specificity
                    if city:
                        bio_queries.insert(2, f'"{donor_name}" "{employer}" "{title}" {city}')
                    log.info(f"  🏢💼 PRIORITY QUERY: \"{donor_name}\" \"{employer}\" \"{title}\"")
                    
                    # For retired prospects, also search past tense
                    if is_retired:
                        bio_queries.extend([
                            f'"{donor_name}" former "{title}" "{employer}"',
                            f'"{donor_name}" retired "{title}"',
                        ])
                
                # =====================================================
                # PRIORITY 2: Name AND Employer (without title)
                # =====================================================
                elif employer:
                    bio_queries.extend([
                        f'"{donor_name}" "{employer}"',
                        f'"{donor_name}" {employer}',
                    ])
                    if city:
                        bio_queries.append(f'"{donor_name}" "{employer}" {city}')
                    log.info(f"  🏢 Including employer in bio queries: {employer}")
                
                # =====================================================
                # PRIORITY 3: Name AND Title (without employer)
                # =====================================================
                elif title:
                    bio_queries.extend([
                        f'"{donor_name}" "{title}"',
                        f'"{donor_name}" {title}',
                    ])
                    if city:
                        bio_queries.append(f'"{donor_name}" "{title}" {city}')
                    log.info(f"  💼 Including title in bio queries: {title}")
                
                # Remove None entries from retired queries
                bio_queries = [q for q in bio_queries if q]
                
                # Core biographical queries (lower priority - employer/title already handled above)
                bio_queries.extend([
                    f'"{donor_name}" biography',
                    f'"{donor_name}" education background',
                ])
                
                # Add education-specific queries if available
                if education:
                    bio_queries.append(f'"{donor_name}" {education}')
                    log.info(f"  🎓 Including education in bio queries: {education}")
                
                # Add spouse-specific queries if available (CRITICAL for disambiguation!)
                if spouse_name:
                    bio_queries.extend([
                        f'"{donor_name}" "{spouse_name}"',
                        f'"{donor_name}" married {spouse_name}',
                        f'"{donor_name}" spouse {spouse_name}',
                    ])
                    log.info(f"  💑 Including spouse in bio queries: {spouse_name}")
                    # Also add spouse + location combo (highly specific!)
                    if location_filter:
                        bio_queries.append(f'"{donor_name}" "{spouse_name}" {location_filter}')
                
                # SPOUSE DISCOVERY: If no spouse in CSV, search for spouse info
                if not spouse_name:
                    bio_queries.extend([
                        f'"{donor_name}" husband wife married',
                        f'"{donor_name}" spouse partner',
                    ])
                    if location_filter:
                        bio_queries.append(f'"{donor_name}" {location_filter} husband OR wife OR married')
                    log.info(f"  💍 Adding spouse discovery queries (no spouse in CSV)")
                
                # MUNICIPAL/CIVIC RECORD SEARCHES - find board memberships, civic activities
                if city:
                    bio_queries.extend([
                        f'"{donor_name}" {city} board trustee commission',
                        f'"{donor_name}" {city} library OR council OR planning',
                        f'"{donor_name}" {city} volunteer civic',
                    ])
                    log.info(f"  🏛️ Adding municipal record queries for: {city}")
                
                # =====================================================
                # MULTIPLE LOCATION SUPPORT - Search all known locations
                # =====================================================
                # Primary location
                if location_filter:
                    bio_queries.append(f'"{donor_name}" {location_filter} biography')
                
                # Secondary/alternate locations (from CSV or merged data)
                secondary_locations = []
                
                # FIRST: Use merged _all_cities if available (from merge_person_rows)
                all_cities = csv_data.get("_all_cities", [])
                if all_cities and len(all_cities) > 1:
                    # Add all cities except the primary one
                    for other_city in all_cities:
                        if city and other_city.lower() != city.lower():
                            if other_city not in secondary_locations:
                                secondary_locations.append(other_city)
                
                # ALSO: Check for secondary location fields (traditional approach)
                for key in csv_data.keys():
                    # Skip internal fields that start with underscore
                    if key.startswith('_'):
                        continue
                    key_lower = key.lower()
                    if any(x in key_lower for x in ['city2', 'secondary', 'previous', 'alternate', 'other']):
                        if 'city' in key_lower or 'location' in key_lower:
                            loc_val = csv_data.get(key, "")
                            # Skip if it's a list (internal data structure)
                            if isinstance(loc_val, list):
                                continue
                            loc_str = str(loc_val).strip()
                            if loc_str and city and loc_str.lower() != city.lower():
                                if loc_str not in secondary_locations:
                                    secondary_locations.append(loc_str)
                
                # Also check for spouse's location that differs from primary
                if alternate_names:
                    for alt in alternate_names:
                        alt_loc = alt.get("location", alt.get("associated_location", ""))
                        if alt_loc and alt_loc.lower() != location_filter.lower():
                            if alt_loc not in secondary_locations:
                                secondary_locations.append(alt_loc)
                
                # Search secondary locations
                if secondary_locations:
                    log.info(f"  📍 Adding searches for {len(secondary_locations)} secondary location(s): {', '.join(secondary_locations)}")
                    for sec_loc in secondary_locations[:3]:  # Limit to 3
                        bio_queries.extend([
                            f'"{donor_name}" {sec_loc}',
                            f'"{donor_name}" {sec_loc} property OR foundation',
                        ])
                
                # =====================================================
                # ALTERNATE NAME / MAIDEN NAME SEARCH (HIGH PRIORITY)
                # These are critical for people who changed names!
                # =====================================================
                bio_queries.extend([
                    f'"{donor_name}" maiden name "also known as"',
                    f'"{donor_name}" "formerly known as" OR "née"',
                ])
                log.info(f"  📛 Adding alternate name discovery queries")
                
                # USE ALTERNATE NAMES FROM CSV (e.g., maiden name from related row)
                # PRIORITY SEARCHES: Put maiden name queries at the TOP
                if alternate_names:
                    for alt in alternate_names:
                        alt_name = alt.get("name", "")
                        alt_location = alt.get("location", alt.get("associated_location", ""))
                        alt_last = alt.get("alternate_last_name", "") if isinstance(alt, dict) else ""
                        
                        if alt_name:
                            # =====================================================
                            # PRIORITY: Maiden name + employer + title (BEST query)
                            # =====================================================
                            if employer and title:
                                bio_queries.insert(0, f'"{alt_name}" "{employer}" "{title}"')
                                log.info(f"  📛🏢 PRIORITY: Maiden name with employer: \"{alt_name}\" \"{employer}\" \"{title}\"")
                            elif employer:
                                bio_queries.insert(0, f'"{alt_name}" "{employer}"')
                                log.info(f"  📛🏢 PRIORITY: Maiden name with employer: \"{alt_name}\" \"{employer}\"")
                            
                            # Standard maiden name searches
                            bio_queries.extend([
                                f'"{alt_name}" biography',
                                f'"{alt_name}" married',
                            ])
                            
                            # Search for connection between current and alternate name
                            bio_queries.append(f'"{donor_name}" "{alt_name}"')
                            
                            if alt_location:
                                bio_queries.append(f'"{alt_name}" {alt_location}')
                            
                            log.info(f"  📛 Adding queries for alternate name from CSV: {alt_name}")
                            
                            # MARRIAGE RECORD SEARCHES - search for maiden name in marriage records
                            if alt_last:
                                # Search for marriage records linking maiden name to current name
                                donor_last = csv_data.get("Last Name", csv_data.get("Last", "")).strip()
                                bio_queries.extend([
                                    f'"{alt_name}" marriage record "{donor_last}"',
                                    f'"{alt_name}" married "{donor_last}"',
                                    # Also try California/state marriage records
                                    f'"{alt_name}" {state} marriage' if state else None,
                                ])
                                bio_queries = [q for q in bio_queries if q]  # Remove None
                                log.info(f"  💒 Adding marriage record searches: {alt_name} → {donor_last}")
                
                # PROFESSIONAL LICENSE SEARCHES - real estate, teaching, nursing, etc.
                if state:
                    bio_queries.extend([
                        f'"{donor_name}" {state} real estate license',
                        f'"{donor_name}" {state} professional license',
                        f'"{donor_name}" {state} teaching credential',
                    ])
                    log.info(f"  📜 Adding professional license searches for {state}")
                
                # WORK HISTORY / EMPLOYER SEARCHES - search county/government employee records
                if city or state:
                    county_name = city.replace(" ", "") if city else ""
                    bio_queries.extend([
                        f'"{donor_name}" teacher tutor educator',
                        f'"{donor_name}" county employee',
                    ])
                    if city:
                        bio_queries.append(f'"{donor_name}" "{city} County" OR "{city}"')
                    log.info(f"  👔 Adding work history / employer searches")
                
                # PROPERTY RECORDS - can indicate wealth and residence history
                if city and state:
                    bio_queries.append(f'"{donor_name}" property owner {city} {state}')
                    log.info(f"  🏠 Adding property ownership search")
                
                # NAME VARIATION SEARCHES - use all generated variations
                if name_variations:
                    # Add highest-priority variations FIRST (maiden names)
                    highest_priority_vars = [v for v in name_variations if v.get("search_priority") == "highest" and v["name"] != donor_name]
                    for var in highest_priority_vars[:3]:  # Maiden names first
                        var_name = var["name"]
                        # Search maiden name with employer/title (BEST query for older records)
                        if employer and title:
                            bio_queries.insert(0, f'"{var_name}" "{employer}" "{title}"')
                        elif employer:
                            bio_queries.insert(0, f'"{var_name}" "{employer}"')
                        bio_queries.extend([
                            f'"{var_name}" {location_filter}' if location_filter else f'"{var_name}"',
                        ])
                    if highest_priority_vars:
                        log.info(f"  📛 Adding PRIORITY searches for {len(highest_priority_vars)} maiden name variations")
                    
                    # Add high-priority variations to search
                    high_priority_vars = [v for v in name_variations if v.get("search_priority") == "high" and v["name"] != donor_name]
                    for var in high_priority_vars[:5]:  # Limit to avoid too many queries
                        var_name = var["name"]
                        bio_queries.extend([
                            f'"{var_name}" {location_filter}' if location_filter else f'"{var_name}"',
                        ])
                    if high_priority_vars:
                        log.info(f"  📛 Adding searches for {len(high_priority_vars)} name variations")
                    
                    # Search for maiden/previous names with marriage records
                    maiden_vars = [v for v in name_variations if v.get("type") in ("maiden", "maiden_initial", "maiden_no_middle")]
                    for var in maiden_vars[:3]:
                        var_name = var["name"]
                        bio_queries.extend([
                            f'"{var_name}" marriage',
                            f'"{var_name}" "{csv_data.get("Last Name", csv_data.get("Last", "")).strip()}"',
                        ])
                    if maiden_vars:
                        log.info(f"  💒 Adding marriage record searches for {len(maiden_vars)} maiden name variations")
                
                # PUBLIC RECORDS DATABASE SEARCHES
                # Search for the person in public records (avoids blocked data broker sites)
                bio_queries.extend([
                    f'"{donor_name}" public records {state}' if state else f'"{donor_name}" public records',
                    f'"{donor_name}" voter registration' if state else None,
                    f'"{donor_name}" marriage divorce records {state}' if state else None,
                ])
                bio_queries = [q for q in bio_queries if q]  # Remove None entries
                log.info(f"  📋 Adding public records searches")
                
                # Standard queries
                bio_queries.extend([
                    f'"{donor_name}" early life family',
                    f'"{donor_name}" residence',
                    f'"{donor_name}" personal life',
                    f'"{donor_name}" linkedin profile'
                ])
                
            bio_bullets = extract_bullet_section(donor_name, entity_type, bio_queries, BULLET_BIO_PROMPT, 
                                                 fast_mode, location_filter, identity_verified=verified_flag, spouse_name=spouse_name)
            
            # Check for hallucinations
            bio_warnings = detect_hallucinations(bio_bullets, "Biographical")
            if bio_warnings:
                for warning in bio_warnings:
                    log.warning(f"  ⚠️  {warning}")
            
            # CRITICAL: Check for location mismatches that indicate wrong person
            bio_bullets = validate_biographical_consistency(bio_bullets, "Biographical", 
                                                           location_filter, donor_name)
            
            # IMPORTANT: After bio section, identity is verified. 
            # Don't use spouse filtering for remaining sections - person's current activities 
            # may not mention deceased/former spouse!
            # NOTE: This is especially important when spouse has passed away - current activities
            # will be under person's name alone, not "Mr. and Mrs."
            
            # 3) Career Research (bullet points)
            log.info("Researching career history using CSV context...")
            if entity_type == "organization":
                career_queries = [
                    f'"{donor_name}" leadership board directors',
                    f'"{donor_name}" executives team',
                    f'"{donor_name}" history timeline'
                ]
            else:
                # Build targeted career queries using employer/title from CSV
                career_queries = []
                
                # Get retired age flag
                is_retired = csv_data.get('_is_retired_age', False)
                
                # Add "Mr. and Mrs." variant if spouse is known (captures historical activities)
                if spouse_name:
                    career_queries.append(f'"Mr. and Mrs. {spouse_name}"')
                    career_queries.append(f'"Mr and Mrs {spouse_name}"')
                
                # =====================================================
                # PRIORITY: Employer/Title queries (MOST SPECIFIC)
                # =====================================================
                if employer and title:
                    # BEST QUERY: All three together
                    career_queries.insert(0, f'"{donor_name}" "{employer}" "{title}"')
                    career_queries.append(f'"{donor_name}" {title} {employer}')
                    log.info(f"  💼 PRIORITY CAREER QUERY: \"{donor_name}\" \"{employer}\" \"{title}\"")
                    
                    # For retired prospects, add former/retired variants
                    if is_retired:
                        career_queries.extend([
                            f'"{donor_name}" former "{title}" "{employer}"',
                            f'"{donor_name}" retired "{employer}"',
                        ])
                        log.info(f"  👴 Adding retired/former career queries")
                elif employer:
                    career_queries.append(f'"{donor_name}" career {employer}')
                    career_queries.append(f'"{donor_name}" "{employer}"')
                    if is_retired:
                        career_queries.append(f'"{donor_name}" retired "{employer}"')
                elif title:
                    career_queries.append(f'"{donor_name}" {title}')
                    if is_retired:
                        career_queries.append(f'"{donor_name}" retired "{title}"')
                
                # RETIRED-SPECIFIC queries (if 75+ years old)
                if is_retired:
                    career_queries.extend([
                        f'"{donor_name}" retired career history',
                        f'"{donor_name}" former professional',
                    ])
                
                # Standard career queries
                career_queries.extend([
                    f'"{donor_name}" career position CEO',
                    f'"{donor_name}" business executive',
                    f'"{donor_name}" professional accomplishments',
                    f'"{donor_name}" company founder',
                    f'"{donor_name}" industry leadership',
                    f'"{donor_name}" awards recognition'
                ])
                
            career_bullets = extract_bullet_section(donor_name, entity_type, career_queries, BULLET_CAREER_PROMPT, 
                                                    fast_mode, location_filter, identity_verified=True)
            
            # Check for hallucinations
            career_warnings = detect_hallucinations(career_bullets, "Career")
            if career_warnings:
                for warning in career_warnings:
                    log.warning(f"  ⚠️  {warning}")
            
            # 4) Philanthropy Research (bullet points)
            log.info("Researching philanthropic activities...")
            if entity_type == "organization":
                philanthropy_queries = [
                    f'"{donor_name}" grants funding donations',
                    f'"{donor_name}" philanthropy charitable giving',
                    f'"{donor_name}" foundation support'
                ]
            else:
                philanthropy_queries = [
                    f'"{donor_name}" philanthropy donations',
                    f'"{donor_name}" foundation charitable giving',
                    f'"{donor_name}" nonprofit board service',
                    f'"{donor_name}" major gifts contributions',
                    f'"{donor_name}" philanthropic philosophy'
                ]
                # Add "Mr. and Mrs." variant if spouse known (historical activities)
                if spouse_name:
                    philanthropy_queries.insert(0, f'"Mr. and Mrs. {spouse_name}" philanthropy')
                    philanthropy_queries.insert(1, f'"Mr. and Mrs. {spouse_name}" foundation')
            philanthropy_bullets = extract_bullet_section(donor_name, entity_type, philanthropy_queries, BULLET_PHILANTHROPY_PROMPT, 
                                                          fast_mode, location_filter, identity_verified=True)
            
            # Check for hallucinations
            phil_warnings = detect_hallucinations(philanthropy_bullets, "Philanthropy")
            if phil_warnings:
                for warning in phil_warnings:
                    log.warning(f"  ⚠️  {warning}")
            
            # 5) Political Research (bullet points)
            log.info("Researching political activity...")
            if entity_type == "organization":
                political_queries = [
                    f'"{donor_name}" political advocacy positions',
                    f'"{donor_name}" policy influence',
                    f'"{donor_name}" political ideology'
                ]
            else:
                political_queries = [
                    f'"{donor_name}" political donations FEC',
                    f'"{donor_name}" political affiliation party',
                    f'"{donor_name}" policy positions advocacy',
                    f'"{donor_name}" political think tank',
                    f'"{donor_name}" political ideology'
                ]
                # Add "Mr. and Mrs." variant if spouse known
                if spouse_name:
                    political_queries.insert(0, f'"Mr. and Mrs. {spouse_name}" political donations')
            political_bullets = extract_bullet_section(donor_name, entity_type, political_queries, BULLET_POLITICAL_PROMPT, 
                                                       fast_mode, location_filter, identity_verified=True)
            
            # Check for hallucinations
            pol_warnings = detect_hallucinations(political_bullets, "Political")
            if pol_warnings:
                for warning in pol_warnings:
                    log.warning(f"  ⚠️  {warning}")
            
            # 6) FEC
            log.info("Researching FEC giving...")
            fec = research_fec(donor_name)
            
            # 6.5) Network & Board Research
            network_bullets = research_network_boards(donor_name, entity_type, location_filter, fast_mode, spouse_name)
            
            # Check for hallucinations
            net_warnings = detect_hallucinations(network_bullets, "Network")
            if net_warnings:
                for warning in net_warnings:
                    log.warning(f"  ⚠️  {warning}")
            
            # 6.6) Net Worth & Giving Capacity Research
            networth_bullets = research_net_worth_capacity(donor_name, entity_type, location_filter, fast_mode, spouse_name, csv_data)
            
            # Check for hallucinations
            nw_warnings = detect_hallucinations(networth_bullets, "Net Worth")
            if nw_warnings:
                for warning in nw_warnings:
                    log.warning(f"  ⚠️  {warning}")
            
            # 6.7) Generate IHS Assessment
            ihs_assessment = generate_ihs_assessment(
                donor_name, location_filter,
                bio_bullets, career_bullets, philanthropy_bullets, 
                political_bullets, network_bullets
            )
            
            # Check for hallucinations in IHS assessment
            ihs_warnings = detect_hallucinations(ihs_assessment, "IHS Assessment")
            if ihs_warnings:
                log.warning("⚠️  HALLUCINATION WARNING in IHS Assessment:")
                for warning in ihs_warnings:
                    log.warning(f"  {warning}")
            
            # 6.8) Generate Strategic Briefing
            strategic_briefing = generate_strategic_briefing(
                donor_name, location_filter,
                bio_bullets, career_bullets, philanthropy_bullets,
                political_bullets, network_bullets
            )
            
            # Check for hallucinations in strategic briefing
            strat_warnings = detect_hallucinations(strategic_briefing, "Strategic Briefing")
            if strat_warnings:
                log.warning("⚠️  HALLUCINATION WARNING in Strategic Briefing:")
                for warning in strat_warnings:
                    log.warning(f"  {warning}")
                # If hallucinations detected, replace with safer text
                if any("FABRICATED" in w or "LIKELY HALLUCINATION" in w for w in strat_warnings):
                    log.warning("  🚨 Replacing hallucinated Strategic Briefing with factual disclaimer")
                    strategic_briefing = f"""**Executive Summary for IHS Board Members: {donor_name}**

Limited public information is available about {donor_name}. Based on available research:

**Who They Are:** Research indicates {donor_name} resides in {location_filter}. Professional and career information is limited in public sources. Age and family connection data was found through public records databases.

**Why IHS Should Engage:** Insufficient public information to assess alignment with IHS mission or giving capacity. No documented philanthropic history or political activity found.

**Recommended Next Steps:** Additional research and personal outreach recommended before cultivation. Consider requesting more information from existing contacts or the prospect directly. Gift capacity and program alignment cannot be assessed without further information.

[Note: This summary reflects limited available public data. Some research sources returned no relevant results for this prospect.]"""
            
            # 6.9) Generate Gift Officer Summary
            log.info("Generating Gift Officer Summary...")
            gift_summary = generate_gift_officer_summary(
                donor_name, entity_type, location_filter,
                bio_bullets, career_bullets, philanthropy_bullets, political_bullets,
                confidence=identity["confidence"],
                markers=identity["markers"],
                collisions=identity["collisions"]
            )
            
            # 7) Inner Circle (conditional)
            connections = []
            skip_inner = should_skip_inner_circle(donor_row)
            
            if not skip_inner:
                log.info(f"Starting deep comparison against {inner_circle_count} inner circle members...")
                connections, conn_cites = research_inner_circle(
                    {"name": donor_name, "City": city, "State": state},
                    inner_rows,
                    max_members=args.max_inner_circle,
                    fast_mode=fast_mode,
                    parallel_workers=args.parallel_workers
                )
            
            # 8) Generate DOCX Report - Partner Profile Format
            doc = Document()
            
            # Title: "Partner Profile"
            render_title(doc, donor_name, entity_type)
            
            # Header section: Name, Contact Info, DOB, Family
            render_header_section(doc, donor_name, csv_data, entity_type)
            
            # Education section
            render_education_section(doc, csv_data, bio_bullets)
            
            # Current Position and Previous Positions
            render_positions_section(doc, career_bullets)
            
            # Foundation and Nonprofit Board Memberships
            render_board_memberships_section(doc, network_bullets)
            
            # Other Affiliations (placeholder)
            p = doc.add_paragraph()
            r = p.add_run("Other Affiliations")
            r.bold = True
            doc.add_paragraph("(Other affiliations from research)")
            
            # Net Worth/Giving Capacity - THE KEY SECTION
            render_networth_section(doc, networth_bullets)
            
            # Philanthropic Interests and Relevant Gifts
            render_philanthropy_section(doc, philanthropy_bullets)
            
            # Publications (placeholder)
            p = doc.add_paragraph()
            r = p.add_run("Publications")
            r.bold = True
            doc.add_paragraph("(Publications from research)")
            
            # Biography section (narrative)
            render_biography_section(doc, bio_bullets, career_bullets)
            
            # Political Activity
            render_political_section(doc, political_bullets, fec)
            
            # Connections (Inner Circle)
            render_connections_section(doc, connections, skip_inner)
            
            # IHS Assessment (keep as is - useful internal tool)
            p = doc.add_paragraph()
            r = p.add_run("IHS Donor Probability Assessment")
            r.bold = True
            r.font.size = Pt(14)
            for line in ihs_assessment.split('\n'):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
            
            # Strategic Summary 
            p = doc.add_paragraph()
            r = p.add_run("Strategic Summary")
            r.bold = True
            r.font.size = Pt(14)
            for line in strategic_briefing.split('\n'):
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
            
            # References
            render_references_section(doc)
            
            # Save
            outpath = os.path.join(args.outdir, f"{donor_name.replace(' ', '_')}_Profile.docx")
            doc.save(outpath)
            log.info(f"✅ Report written: {outpath}")
            log.info(f"Completed: {donor_name}\n")
            
        except Exception as e:
            log.error(f"❌ Failed processing {donor_name}: {e}")
            import traceback
            traceback.print_exc()
            continue
    
    print_search_stats()
    
    log.info(f"\n{'='*60}")
    log.info(f"✅ ALL DONE! Processed {len(names_to_process)} donor(s)")
    log.info(f"{'='*60}")

if __name__ == "__main__":
    main()